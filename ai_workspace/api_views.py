import copy
import io
import json
import logging
import mimetypes
import os
import shutil
import zipfile,itertools
from datetime import datetime
from glob import glob
from urllib.parse import urlparse
from ai_auth.tasks import count_update,weighted_count_update
import django_filters
import docx2txt
import nltk
import requests
from delta import html
from django.conf import settings
from django.core.files import File as DJFile
from django.core.files.base import ContentFile
from django.db import IntegrityError
from django.db import transaction
from django.db.models import Q, Sum, Count
from django.http import Http404, HttpResponse
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, get_list_or_404
from django_filters.rest_framework import DjangoFilterBackend
from filesplit.split import Split
from google.cloud import speech_v1p1beta1 as speech
from google.cloud import storage
from htmldocx import HtmlToDocx
from indicnlp.tokenize.sentence_tokenize import sentence_split
from notifications.signals import notify
from pydub import AudioSegment
from rest_framework import permissions
from rest_framework import viewsets, status, generics, pagination
from rest_framework.decorators import api_view
from rest_framework.decorators import permission_classes
from rest_framework.filters import SearchFilter, OrderingFilter
from rest_framework.pagination import PageNumberPagination
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from filesplit.split import Split
from ai_auth.utils import authorize_list,filter_authorize, unassign_task
from django_oso.auth import authorize
logger = logging.getLogger('django')
from django.db import models
from django.db.models import Prefetch
from django.db.models.functions import Lower
from ai_auth.models import AiUser, UserCredits
from ai_auth.models import HiredEditors
from ai_auth.tasks import mt_only, text_to_speech_long_celery, transcribe_long_file_cel, project_analysis_property
from ai_auth.tasks import write_doc_json_file,record_api_usage
from ai_glex.serializers import GlossarySetupSerializer, GlossaryFileSerializer, GlossarySerializer
from ai_marketplace.models import ChatMessage
from ai_marketplace.serializers import ThreadSerializer
from ai_pay.api_views import po_modify
from ai_staff.models import LanguagesLocale, AilaysaSupportedMtpeEngines,AiCustomize
from ai_workspace import forms as ws_forms
from ai_workspace.excel_utils import WriteToExcel_lite
from ai_workspace.tbx_read import upload_template_data_to_db, user_tbx_write
from ai_workspace.utils import create_assignment_id
from ai_workspace_okapi.models import Document
from ai_workspace_okapi.utils import download_file, text_to_speech, text_to_speech_long, get_res_path
from ai_workspace_okapi.utils import get_translation, file_translate
from .models import AiRoleandStep, Project, Job, File, ProjectContentType, ProjectSubjectField, TempProject, TmxFile, ReferenceFiles, \
    Templangpair, TempFiles, TemplateTermsModel, TaskDetails, \
    TaskAssignInfo, TaskTranscriptDetails, TaskAssign, Workflows, Steps, WorkflowSteps, TaskAssignHistory, \
    ExpressProjectDetail
from .models import Task
from cacheops import cached
from operator import attrgetter
from .models import TbxFile, Instructionfiles, MyDocuments, ExpressProjectSrcSegment, ExpressProjectSrcMTRaw,\
                    ExpressProjectAIMT, WriterProject,DocumentImages,ExpressTaskHistory, TaskTranslatedFile
from .serializers import (ProjectContentTypeSerializer, \
                          ProjectSerializer, JobSerializer, FileSerializer, \
                          ProjectSubjectSerializer, TempProjectSetupSerializer, \
                          TaskSerializer, FileSerializerv2, TmxFileSerializer, \
                          PentmWriteSerializer, TbxUploadSerializer, ProjectQuickSetupSerializer, TbxFileSerializer, \
                          VendorDashBoardSerializer, ProjectSerializerV2, ReferenceFileSerializer,
                          TbxTemplateSerializer, TaskTranslatedFileSerializer,\
                          TaskAssignInfoSerializer, TaskDetailSerializer, ProjectListSerializer, \
                          GetAssignToSerializer, TaskTranscriptDetailSerializer, InstructionfilesSerializer,\
                          StepsSerializer, WorkflowsSerializer, ProjectSimpleSerializer,\
                          WorkflowsStepsSerializer, TaskAssignUpdateSerializer, ProjectStepsSerializer,\
                          ExpressProjectDetailSerializer,MyDocumentSerializer,ExpressProjectAIMTSerializer,\
                          WriterProjectSerializer,DocumentImagesSerializer,ExpressTaskHistorySerializer,MyDocumentSerializerNew) #ProjectCreationSerializer
from .utils import DjRestUtils
from django.utils import timezone
from .utils import get_consumable_credits_for_text_to_speech,\
                   get_consumable_credits_for_speech_to_text, progress_filter#,filter_status
import regex as re
spring_host = os.environ.get("SPRING_HOST")
from django.db.models import Case, When, F, Value, DateTimeField, ExpressionWrapper
from django.db.models.functions import Coalesce
from django.db.models.query import QuerySet
from ai_auth.utils import get_assignment_role
from django.views.decorators.cache import never_cache
from ai_canvas.serializers import CanvasDesignSerializer
from rest_framework.authentication import TokenAuthentication
from ai_auth.authentication import APIAuthentication
from rest_framework.decorators import authentication_classes
from .utils import merge_dict,split_file_by_size
from ai_auth.access_policies import IsEnterpriseUser
from datetime import date
import spacy, time



nlp = spacy.load("en_core_web_sm")

class IsCustomer(permissions.BasePermission):

    def has_permission(self, request, view):
        if request.user.user_permissions.filter(codename="user-attribute-exist").first():
            return True


class JobView(viewsets.ModelViewSet):
    ''' 
    Used in QuickSetupProject view internally to create and delete jobs. 
    '''
    serializer_class = JobSerializer

    def get_object(self, many=False):
        objs = []
        obj = None
        if not many:
            try:
                obj = get_object_or_404(Job.objects.all(),\
                    id=self.kwargs.get("pk"))
            except:
                raise Http404
            return  obj

        objs_ids_list =  self.kwargs.get("ids").split(",")

        for obj_id in objs_ids_list:
            try:
                objs.append(get_object_or_404(Job.objects.all(),\
                    id=obj_id))
            except:
                raise Http404
        return objs

    def get_queryset(self):
        return Job.objects.filter(project__ai_user=self.request.user)

    def create(self, request):
        serializer = self.serializer_class(data = request.data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            objs=authorize_list(objs,"delete",request.user) 
            for obj in objs: 
                obj.delete()
            return Response(status=204)
        obj=self.get_object(many=False) 
        authorize(request,resource=obj,action="delete",actor=self.request.user)
        return super().destroy(request, *args, **kwargs)


class ProjectSubjectView(viewsets.ModelViewSet):
    ''' 
    Used in QuickSetupProject view internally to create and delete subject fields.
    '''
    serializer_class = ProjectSubjectField

    def get_object(self, many=False):
        objs = []
        obj = None
        if not many:
            try:
                obj = get_object_or_404(ProjectSubjectField.objects.all(),\
                    id=self.kwargs.get("pk"))
            except:
                raise Http404
            return  obj

        objs_ids_list =  self.kwargs.get("ids").split(",")

        for obj_id in objs_ids_list:
            try:
                objs.append(get_object_or_404(ProjectSubjectField.objects.all(),\
                    id=obj_id))
            except:
                raise Http404
        return objs


    def get_queryset(self):
        project_id = self.request.query_params.get('project_id')
        return ProjectSubjectField.objects.filter(project__id=project_id)


    def list(self,request):
        queryset = self.get_queryset()
        serializer = ProjectSubjectSerializer(queryset, many=True, context={'request': request})
        return  Response(serializer.data)

    def create(self, request):
        serializer = self.serializer_class(data = request.data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            for obj in objs:
                obj.delete()
            return Response(status=204)
        return super().destroy(request, *args, **kwargs)


class ProjectContentTypeView(viewsets.ModelViewSet):
    ''' 
    Used in QuickSetupProject view internally to create and delete content_type.
    '''
    serializer_class = ProjectContentTypeSerializer

    def get_object(self, many=False):
        objs = []
        obj = None
        if not many:
            try:
                obj = get_object_or_404(ProjectContentType.objects.all(),\
                    id=self.kwargs.get("pk"))
            except:
                raise Http404
            return  obj

        objs_ids_list =  self.kwargs.get("ids").split(",")

        for obj_id in objs_ids_list:
            try:
                objs.append(get_object_or_404(ProjectContentType.objects.all(),\
                    id=obj_id))
            except:
                raise Http404
        return objs

    def get_queryset(self):
        project_id = self.request.query_params.get('project_id')
        return ProjectContentType.objects.filter(project__id=project_id)

    def list(self,request):
        queryset = self.get_queryset()
        # pagin_tc = self.paginate_queryset( queryset, request , view=self )
        serializer = ProjectContentTypeSerializer(queryset, many=True, context={'request': request})
        # response =self.get_paginated_response(serializer.data)
        return  Response(serializer.data)

    def create(self, request):
        serializer = self.serializer_class(data = request.data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data)

    def destroy(self, request, *args, **kwargs):
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            for obj in objs:
                obj.delete()
            return Response(status=204)
        return super().destroy(request, *args, **kwargs)


class FileView(viewsets.ModelViewSet):
    ''' 
    Used in QuickSetupProject view internally to create and delete files.
    '''
    serializer_class = FileSerializer
    parser_classes = [MultiPartParser, FormParser]

    def get_object(self, many=False):
        objs = []
        obj = None
        if not many:
            try:
                obj = get_object_or_404(File.objects.all(),\
                    id=self.kwargs.get("pk"))
            except:
                raise Http404
            return  obj

        objs_ids_list =  self.kwargs.get("ids").split(",")

        for obj_id in objs_ids_list:
            try:
                objs.append(get_object_or_404(File.objects.all(),\
                    id=obj_id))
            except:
                raise Http404
        return objs

    def get_queryset(self):
        return File.objects.filter(project__ai_user=self.request.user)

    def create(self, request):
        print(request.data)
        serializer = FileSerializer(data=request.data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data, status=201)

    def destroy(self, request, *args, **kwargs):
        from ai_workspace_okapi.api_views import DocumentViewByTask  
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            objs=authorize_list(objs,"delete",request.user)
            for obj in objs:
                tasks = obj.file_tasks_set.all()
                for i in tasks:
                    path = DocumentViewByTask.get_json_file_path(i)
                    if os.path.exists(path):
                        os.remove(path)
                os.remove(obj.file.path)
                obj.delete()
            return Response(status=204)
        file=self.get_object(many=False)
        authorize(request,resource=file,action="delete",actor=request.user)
        return super().destroy(request, *args, **kwargs)


def integrity_error(func):
    def decorator(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except IntegrityError as e:
            print("error---->", e)
            return Response({'message': "integrirty error"}, 409)

    return decorator


# convert text to file
def text_file_processing(text_data):
    name =  text_data.split()[0]+ ".txt" if len(text_data.split()[0])<=15 else text_data[:5]+ ".txt"
    f1 = open(name, 'w')
    f1.write(text_data)
    f1.close()
    f2 = open(name, 'rb')
    file_obj2 = DJFile(f2)
    return file_obj2,f2,name



class TempProjectSetupView(viewsets.ViewSet):
    ''' 
    Project creation before user logs in..  Not using now.
    '''
    serializer_class = TempProjectSetupSerializer
    parser_classes = [MultiPartParser, JSONParser]
    permission_classes = [AllowAny,]

    def get_queryset(self):
        return TempProject.objects.filter(ai_user=self.request.user)

    @integrity_error
    def create(self, request):
        text_data=request.POST.get('text_data')
        if text_data:
            if urlparse(text_data).scheme:
                return Response({"msg":"Url not Accepted"},status=406)
            file_obj2,f2,name = text_file_processing(text_data)
            serializer = TempProjectSetupSerializer(data={**request.POST.dict(),"tempfiles":[file_obj2]})
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                f2.close()
                os.remove(os.path.abspath(name))
                return Response(serializer.data, status=201)
            os.remove(os.path.abspath(name))
            return Response(serializer.errors, status=409)
        else:
            serializer = TempProjectSetupSerializer(data={**request.POST.dict(),
                "tempfiles":request.FILES.getlist('files')})
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                return Response(serializer.data, status=201)
            return Response(serializer.errors, status=409)


class Files_Jobs_List(APIView):
    permission_classes = [IsAuthenticated]

    ''' 
    Retrive all related objects associated with project.
    '''

    def get_queryset(self, project_id):
        project = get_object_or_404(Project.objects.all(), id=project_id)
        authorize(self.request, resource=project, actor=self.request.user, action="read")
        jobs = project.project_jobs_set.all()
        jobs = filter_authorize(self.request,jobs,"read",self.request.user)
        contents = project.proj_content_type.all()
        subjects = project.proj_subject.all()
        steps = project.proj_steps.all()
        try:gloss = project.glossary_project
        except:gloss = None
        files = project.project_files_set.filter(usage_type__use_type="source").all()
        glossary_files = project.project_files.all()
        return jobs, files, contents, subjects, steps, project, gloss, glossary_files

    def get(self, request, project_id):
        jobs, files, contents, subjects, steps, project, gloss, glossary_files = self.get_queryset(project_id)
        team_edit = False if project.assigned == True else True
        jobs = JobSerializer(jobs, many=True)
        files = FileSerializer(files, many=True)
        wc_selected = True if project.project.filter(glossary__project__project_type_id =10).exists() else False 
        glossary = GlossarySerializer(gloss).data if gloss else None
        glossary_files = GlossaryFileSerializer(glossary_files,many=True)
        contents = ProjectContentTypeSerializer(contents,many=True)
        subjects = ProjectSubjectSerializer(subjects,many=True)
        steps = ProjectStepsSerializer(steps,many=True)
        return Response({"files":files.data,"glossary_files":glossary_files.data,"glossary":glossary,"jobs": jobs.data, "subjects":subjects.data,\
                        "contents":contents.data, "steps":steps.data, "project_name": project.project_name, "team":project.get_team,"get_mt_by_page":project.get_mt_by_page,\
                         "team_edit":team_edit,"project_type_id":project.project_type.id,"mt_engine_id":project.mt_engine_id,'pre_translate':project.pre_translate,\
                         "project_deadline":project.project_deadline, "mt_enable": project.mt_enable, "revision_step_edit":project.PR_step_edit, "wc_selected":wc_selected}, status=200)



class TmxFilesOfProject(APIView):
    def get_queryset(self, project_id):
        project_qs = Project.objects.all()
        project = get_object_or_404(project_qs, id=project_id, ai_user=self.request.user)
        files = project.project_files_set.all()
        return files

    def post(self, request, project_id):
        files = self.get_queryset(project_id=project_id)
        res_paths = {"srx_file_path": "okapi_resources/okapi_default_icu4j.srx",
                     "fprm_file_path": None}
        data = []
        for file in files:
            params_data = FileSerializerv2(file).data
            res = requests.post(
                f"http://{spring_host}:8080/source/createTmx",
                data={
                    "doc_req_params": json.dumps(params_data),
                    "doc_req_res_params": json.dumps(res_paths)
                }
            )

            # if res.status_code in [200, 201]:
            data.append(res.text)
        return JsonResponse({"results":data}, safe=False)

# class ProjectReportAnalysis(APIView):
#     def get_queryset(self, project_id):
#         project_qs = Project.objects.all()
#         project = get_object_or_404(project_qs, id=project_id, ai_user=self.request.user)
#         files = project.project_files_set.all()
#         return files
#
#     def post(self, request, project_id):
#         data = dict(
#             pentm_path = "/home/langscape/Documents/ailaysa_github/Ai_TMS/media/u343460/u343460p1/.pentm/",
#             report_output_path = "/home/langscape/Documents/ailaysa_github/Ai_TMS/media/u343460/u343460p1/tt/report.html",
#             srx_file_path = "/home/langscape/Documents/ailaysa_github/Ai_TMS/okapi_resources/okapi_default_icu4j.srx"
#         )
#         files = self.get_queryset(project_id)
#         batches_data =  FileSerializerv3(files, many=True).data
#         data = {
#             **data,
#             **dict(batches=batches_data)
#         }
#         print("data---->", data)
#         res = requests.post(
#             f"http://{spring_host}:8080/project/report-analysis",
#             data = {"report_params": json.dumps(data)}
#         )
#         if res.status_code in [200, 201]:
#             return JsonResponse({"msg": res.text}, safe=False)
#         else:
#             return JsonResponse({"msg": "something went to wrong"}, safe=False)

class TmxFileView(viewsets.ViewSet):

    @staticmethod
    def TmxToPenseiveWrite(data):
        if len(data)==0:
            return
        project_id = data[0]["project"]
        project = Project.objects.get(id=project_id)
        data = PentmWriteSerializer(project).data

        res = requests.post(f"http://{spring_host}:8080/project/pentm/create",
                            data={"pentm_params": json.dumps(data)})
        if res.status_code == 200:
            for tmx_data in res.json():
                instance = project.project_tmx_files.filter(id=tmx_data.get('tmx_id','')).first()
                ser = TmxFileSerializer(instance, data=tmx_data, partial=True)
                if ser.is_valid(raise_exception=True):
                    ser.save()
            return JsonResponse(res.json(), safe=False)
        else:
            return JsonResponse({"msg": "Something wrong with file processing"}, status=res.status_code)

    def create(self, request):
        data = {**request.POST.dict(), "tmx_files": request.FILES.getlist("tmx_files")}
        ser_data = TmxFileSerializer.prepare_data(request,data)
        ser = TmxFileSerializer(data=ser_data, many=True)
        if ser.is_valid(raise_exception=True):
            ser.save()
            return self.TmxToPenseiveWrite(ser.data)

# class TbxUploadView(APIView):##currently not using
#     def post(self, request):
#         tbx_files = request.FILES.get('tbx_files')
#         project_id = request.POST.get('project', 0)
#         doc_id = request.POST.get('doc_id', 0)
#         if doc_id != 0:
#             job_id = Document.objects.get(id=doc_id).job_id
#             project_id = Job.objects.get(id=job_id).project_id
#         pro=get_object_or_404(Project,id=project_id)
#         authorize(request,resource=pro,action="create",actor=self.request.user)
#         serializer = TbxUploadSerializer(data={'tbx_files':tbx_files,'project':project_id})
#         if serializer.is_valid():
#             serializer.save()
#             return Response(serializer.data)
#         else:
#             return Response(serializer.errors)



def docx_save_pdf(pdf_obj):
    # conversion of pdf to docx
    from docx import Document
    from htmldocx import HtmlToDocx

    document = Document()
    new_parser = HtmlToDocx()
    new_parser.table_style = 'TableGrid'
    new_parser.add_html_to_document(pdf_obj.html_data, document)
    document.save(pdf_obj.docx_file_name)
    f2 = open(pdf_obj.docx_file_name, 'rb')
    file_obj = DJFile(f2)
    pdf_obj.docx_file_from_writer = file_obj
    pdf_obj.save()



def get_file_from_pdf(pdf_obj_id,pdf_task_id):
    # Get file object from pdf creation
    from ai_exportpdf.models import Ai_PdfUpload
    from ai_exportpdf.views import get_docx_file_path
    if pdf_obj_id:
        pdf_obj = Ai_PdfUpload.objects.get(id = pdf_obj_id)
    else:
        pdf_obj = Ai_PdfUpload.objects.filter(task_id = pdf_task_id).last() 
    
    if pdf_obj.pdf_api_use == "convertio":
        docx_file_path = get_docx_file_path(pdf_obj.id)
        file = open(docx_file_path,'rb')
        file_obj = ContentFile(file.read(),name= os.path.basename(docx_file_path))#name=docx_file_name
        pdf_obj.translation_task_created = True
        pdf_obj.save()
        
    else:
        file_obj = ContentFile(pdf_obj.docx_file_from_writer.file.read(),name= os.path.basename(pdf_obj.docx_file_from_writer.path))
    return file_obj


def get_field_type(field_name, queryset):
    stripped_field_name = field_name.lstrip('-')
    if stripped_field_name in queryset.query.annotations:
        return queryset.query.annotations[stripped_field_name].output_field
    return queryset.model._meta.get_field(stripped_field_name)

class CaseInsensitiveOrderingFilter(OrderingFilter):
    
    def filter_queryset(self, request, queryset, view):
        ordering = self.get_ordering(request, queryset, view)

        if ordering:
            new_ordering = []
            for field in ordering:
                if not isinstance(get_field_type(field, queryset), (models.CharField, models.TextField)):
                    new_ordering.append(field)
                elif field.startswith('-'):
                    new_ordering.append(Lower(field[1:]).desc())
                else:
                    new_ordering.append(Lower(field).asc())
            return queryset.order_by(*new_ordering)

        return queryset


from ai_exportpdf.models import Ai_PdfUpload
class ProjectFilter(django_filters.FilterSet):
    project = django_filters.CharFilter(field_name='project_name',lookup_expr='icontains')
    filter = django_filters.CharFilter(label='glossary or voice',method='filter_not_empty')
    team = django_filters.CharFilter(field_name='team__name',method='filter_team')#lookup_expr='isnull')
    type = django_filters.NumberFilter(field_name='project_type_id')
    assign_status = django_filters.CharFilter(method='filter_status')
    #assign_to = django_filters.CharFilter(method='filter_assign_to')


    class Meta:
        model = Project
        fields = ('project', 'team','type','assign_status')#,'assign_to')


    def filter_team(self, queryset, name, value):
        #it checks for the user is having team or not
        if value=="None":
            lookup = '__'.join([name, 'isnull'])
            return queryset.filter(**{lookup: True})
        else:
            lookup = '__'.join([name, 'icontains'])
            return queryset.filter(**{lookup: value})

    def filter_status(self, queryset, name, value):
        #It is to check for the assigned user work status
        user = self.request.user
        assign_to = self.request.query_params.get('assign_to')
        if user.team and user in user.team.get_editors:
            assign_to_list = [user]
        elif assign_to:
            assign_to_list = assign_to.split(',')
        else: assign_to_list = []
        queryset = progress_filter(queryset,value,assign_to_list)
        return queryset
    

    def filter_not_empty(self,queryset, name, value):
        #project type filter
        # if value == "assets":
        #     queryset = queryset.filter(Q(glossary_project__isnull=False))
        if value == "glossary":
            queryset = queryset.filter(Q(glossary_project__isnull=False)).exclude(project_type_id=10)
        elif value == "voice":
            queryset = queryset.filter(Q(voice_proj_detail__isnull=False))
        elif value == "transcription":
            queryset = queryset.filter(Q(voice_proj_detail__isnull=False)&Q(voice_proj_detail__project_type_sub_category_id = 1))
        elif value == "ai_voice":
            queryset = queryset.filter(Q(voice_proj_detail__isnull=False)&Q(voice_proj_detail__project_type_sub_category_id = 2))
        elif value == "translation":
            queryset = queryset.filter(Q(glossary_project__isnull=True)&Q(voice_proj_detail__isnull=True)).exclude(project_type_id__in = [6,7,8])#.exclude(project_file_create_type__file_create_type="From insta text")#.exclude(project_type_id = 5)
        elif value == "designer":
            queryset = queryset.filter(project_type_id=6)
        elif value == "news":
            queryset = queryset.filter(project_type_id=8) 
        elif value == "word_choices":
            queryset = queryset.filter(project_type_id=10)
        #print("QRF-->",queryset)
        return queryset



class QuickProjectSetupView(viewsets.ModelViewSet):
    '''
    This view is to list, create, update and delete the projects
    '''
    permission_classes = [IsAuthenticated]
    paginator = PageNumberPagination()
    serializer_class = ProjectQuickSetupSerializer
    filter_backends = [DjangoFilterBackend,SearchFilter,CaseInsensitiveOrderingFilter]
    ordering_fields = ['project_name','team__name','id']
    filterset_class = ProjectFilter
    search_fields = ['project_name','project_files_set__filename','project_jobs_set__source_language__language',\
                    'project_jobs_set__target_language__language']
    ordering = ('-id')#'-project_jobs_set__job_tasks_set__task_info__task_assign_info__created_at',
    paginator.page_size = 20

    def get_serializer_class(self):
        #if project_type is glossary, then it will return glossarysetupserializer
        project_type = json.loads(self.request.POST.get('project_type','1'))
        if project_type == 3 or project_type == 10:
        # if project_type == 3:
            return GlossarySetupSerializer
        print("project")
        return ProjectQuickSetupSerializer

    def get_object(self):
        pk = self.kwargs.get("pk", 0)
        try:
            obj = get_object_or_404(Project.objects.all(), id=pk)
        except:
            raise Http404
        return obj
    

    def query_filter_project_news(self,queryset):
        
        return self.filter_queryset(self.queryset)

    #@cached(timeout=60 * 15)
    def get_queryset(self):
        from ai_auth.models import InternalMember
        pr_managers = self.request.user.team.get_project_manager if self.request.user.team and self.request.user.team.owner.is_agency else [] 
        user = self.request.user.team.owner if self.request.user.team and self.request.user.team.owner.is_agency and \
            self.request.user in pr_managers else self.request.user

        # Checking for team access and indivual user access
        
        queryset = Project.objects.filter(((Q(project_jobs_set__job_tasks_set__task_info__assign_to = user) & ~Q(ai_user = user))\
                    |Q(project_jobs_set__job_tasks_set__task_info__assign_to = self.request.user))\
                    |Q(ai_user = self.request.user)
                    |Q(team__internal_member_team_info__in = self.request.user.internal_member.filter(role=1))).distinct()
        
        return queryset
 
    def get_user(self):
        # returns main account holder 
        user = self.request.user
        user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user
        return user_1
    

    # Tried project list with limit and offset instead of pagination to minimize time taken

    # def list(self, request, *args, **kwargs):
    #     st_time = time.time()     
    #     queryset = self.filter_queryset(self.get_queryset())
    #     user_1 = self.get_user()
    #     limit = request.query_params.get('limit')
    #     offset = request.query_params.get('offset')
    #     print("Limit Offset----------->",limit,offset)
    #     if limit is not None and offset is not None:
    #         queryset = queryset[int(offset):int(offset) + int(limit)]
    #     din = AddStoriesView.check_user_dinamalar(user_1)
    #     if din:
    #         serializer = ProjectSimpleSerializer(queryset, many=True, context={'request': request,'user_1':user_1})
    #     else:
    #         serializer = ProjectQuickSetupSerializer(queryset, many=True, context={'request': request,'user_1':user_1})
    #     et_time = time.time()
    #     print("Time taken for list------------------>",et_time-st_time)
    #     return Response(serializer.data)
        
    def list(self, request, *args, **kwargs):
        # filter the projects. Now assign_status filter is used only for Dinamalar flow 
        queryset = self.get_queryset()
        user_1 = self.get_user()
        din = AddStoriesView.check_user_dinamalar(user_1)

        queryset = self.filter_queryset(queryset)

        pagin_tc = self.paginator.paginate_queryset(queryset, request, view=self) ###--> 
        # check for dinamalar user. if so, it will return simple serializer with only required fields
        if din:
            self.paginator.page_size = 10  ### pagination changed to 10 for Din
            # serializer = ProjectSimpleSerializer(pagin_tc, many=True,\
                        #  context={'request': request,'user_1':user_1})     
        # else:
        ## the above code is commanded for standard project for din
        serializer = ProjectQuickSetupSerializer(pagin_tc, many=True,context={'request': request,'user_1':user_1})
        response = self.get_paginated_response(serializer.data)
        return  response


    def retrieve(self, request, pk):
        query = Project.objects.get(id=pk)
        user_1 = self.get_user()
        serializer = ProjectQuickSetupSerializer(query, many=False,context={'request': request,'user_1':user_1})
        return Response(serializer.data)

    def create(self, request):

        punctuation='''!"#$%&'``()*+,-./:;<=>?@[\]^`{|}~_'''

        # Getting the direct text data if project is Express / Instant translation project

        text_data = request.POST.get('text_data')
        ser = self.get_serializer_class()
        pdf_obj_id = request.POST.get('pdf_obj_id', None)

        # Getting the audio files if the project type is Voice project

        audio_file = request.FILES.getlist('audio_file',None)

        # Finding the main account holder
        user_1 = self.get_user()

        # project create with text data
        if text_data:
            if urlparse(text_data).scheme:
                return Response({"msg": "Url not Accepted"}, status=406)

            name = text_data.split()[0].strip(punctuation) + ".txt" if len(text_data.split()[0]) <= 15 else text_data[:5].strip(punctuation) + ".txt"
            im_file = DjRestUtils.convert_content_to_inmemoryfile(filecontent=text_data.encode(), file_name=name)
            serlzr = ser(data={**request.data, "files": [im_file], "from_text": ['true']}, context={"request": request})


        # project create from pdf 
        elif pdf_obj_id:
            files_ = request.FILES.getlist('files')
            file_obj = get_file_from_pdf(pdf_obj_id,None)
            files_.append(file_obj)
            serlzr = ser(data={**request.data,"files":files_},context={"request": request,'user_1':user_1})    
        
        # normal create
        else:
            serlzr = ser(data=\
            {**request.data, "files": request.FILES.getlist("files"),"audio_file":audio_file},context={"request": request,'user_1':user_1})
        
            
        if serlzr.is_valid(raise_exception=True):
            serlzr.save()
            pr = Project.objects.get(id=serlzr.data.get('id'))
            #checks for pre-translation option and initiates the celery task
            if pr.pre_translate == True:
                mt_only.apply_async((serlzr.data.get('id'), str(request.auth)),queue='high-priority' )
            return Response(serlzr.data, status=201)
        return Response(serlzr.errors, status=409)

    def update(self, request, pk, format=None):
        instance = self.get_object()
        ser = self.get_serializer_class()
        task_id=request.POST.get('task_id',None)
        pdf_obj_id = request.POST.get('pdf_obj_id',None)
        pdf_task_id = request.POST.get('pdf_task_id',None)
        team = request.POST.get('team',None)
        target_languages = request.POST.get('target_languages',None)
        req_copy = copy.copy( request._request)
        req_copy.method = "DELETE"

        user_1 = self.get_user()

        file_delete_ids = self.request.query_params.get(\
            "file_delete_ids", [])
        job_delete_ids = self.request.query_params.get(\
            "job_delete_ids", [])
        content_delete_ids = self.request.query_params.get(\
            "content_delete_ids", [])
        subject_delete_ids = self.request.query_params.get(\
            "subject_delete_ids", [])
        step_delete_ids = self.request.query_params.get(\
            "step_delete_ids", [])
        
        if file_delete_ids and target_languages:
            return Response({'msg':"dont delete file"},status=400)
        #Deletion of steps,files,jobs,content_type,subject_fields
        if step_delete_ids:
            for task_obj in instance.get_tasks:
                task_obj.task_info.filter(task_assign_info__isnull=True).filter(step_id__in=step_delete_ids).delete()
            instance.proj_steps.filter(steps__in=step_delete_ids).delete()

        if file_delete_ids:
            file_res = FileView.as_view({"delete": "destroy"})(request=req_copy,\
                        pk='0', many="true", ids=file_delete_ids)

        if job_delete_ids:
            # job_ids = job_delete_ids.split(",")
            # for i in job_ids:
            #     job_instance = Job.objects.get(id=i)
            #     for j in job_instance.job_tasks_set.all():
            #         for k in j.task_info.all():
            #             try:
            #                 return Response({'msg':'task is assigned'},status=status.HTTP_400_BAD_REQUEST)
            #             except:
            #                 print("task is not assigned")
            job_res = JobView.as_view({"delete": "destroy"})(request=req_copy,\
                        pk='0', many="true", ids=job_delete_ids)

        if content_delete_ids:
            content_res = ProjectContentTypeView.as_view({"delete": "destroy"})(request=req_copy,\
                        pk='0', many="true", ids=content_delete_ids)

        if subject_delete_ids:
            subject_res = ProjectSubjectView.as_view({"delete": "destroy"})(request=req_copy,\
                        pk='0', many="true", ids=subject_delete_ids)
        
        if not team:
            team = True if instance.team else False
        
        # update from writer
        if task_id:
            file_obj = update_project_from_writer(task_id)
            serlzr = ser(instance, data=\
                {**request.data, "files":[file_obj],"team":[team]},context={"request": request,'user_1':user_1}, partial=True)

        # update from pdf flow    
        elif pdf_obj_id or pdf_task_id:
            if pdf_obj_id:file_obj = get_file_from_pdf(pdf_obj_id,None)
            else:file_obj = get_file_from_pdf(None,pdf_task_id)
            serlzr = ser(instance, data=\
                {**request.data, "files":[file_obj],"team":[team]},context={"request": request,'user_1':user_1}, partial=True)
            
        else:
            serlzr = ser(instance, data=\
                {**request.data, "files": request.FILES.getlist("files"),"team":[team]},
                context={"request": request,'user_1':user_1}, partial=True)

        if serlzr.is_valid(raise_exception=True):
            serlzr.save()
            pr = Project.objects.get(id=serlzr.data.get('id'))
            # checks for project_type and create extra info needed for it
            if pr.project_type_id == 8:
                NewsProjectSetupView.create_news_detail(pr)
            # if instance.pre_translate == True:
            #     mt_only.apply_async((serlzr.data.get('id'), str(request.auth)), )    
            #mt_only.apply_async((serlzr.data.get('id'), str(request.auth)), )
            return Response(serlzr.data)
        return Response(serlzr.errors, status=409)

    def destroy(self, request, *args, **kwargs):
        project = self.get_object()
        if project.assigned == True:
            return Response({'msg':'some tasks are assigned in this project. Unassign and delete'},status=400)
        else:
            project.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)



class VendorDashBoardFilter(django_filters.FilterSet):
    status = django_filters.CharFilter(method='filter_status')
   
    class Meta:
        model = Task
        fields = ('status',)

    def filter_status(self, queryset, name, value):
        assign_filter = self.request.query_params.get('assign_to')
        users = assign_filter.split(',') if assign_filter else None
        queryset_1 = queryset.filter(task_info__task_assign_info__isnull = False)
        
        if value == 'inprogress':
            # it filters the task which is in status inprogress,yet_to_start,return_request
            # and client_response in rework
            if users:
                # it is for editors filter
                tsk_ids = queryset_1.filter(Q(task_info__status__in=[1,2,4])|Q(task_info__client_response = 2),Q(task_info__assign_to__in=users)).\
                            distinct().values_list('id',flat=True)
            else:
                tsk_ids = queryset.filter(Q(task_info__status__in=[1,2,4])|Q(task_info__client_response = 2)).\
                            distinct().values_list('id',flat=True)
        elif value == 'submitted':
            # it filters the task which is in status completed exclude the client_response approved
            if users:
                tsk_ids = queryset_1.filter(task_info__status = 3,task_info__assign_to__in=users).exclude(task_info__client_response=1).\
                            distinct().values_list('id',flat=True)
            else:
                tsk_ids = queryset.filter(task_info__status = 3).exclude(task_info__client_response=1).\
                            distinct().values_list('id',flat=True)
        elif value =='approved':
            # it filters the task in which client_response is approved
            if users:
                tsk_ids = queryset_1.filter(Q(task_info__client_response = 1),Q(task_info__assign_to__in=users)).\
                            distinct().values_list('id',flat=True)
            else:
                tsk_ids = queryset.filter(Q(task_info__client_response = 1)).distinct().values_list('id',flat=True)
        queryset = queryset.filter(id__in=tsk_ids)
        return queryset
    



class VendorDashBoardView(viewsets.ModelViewSet):
    ''' To get task details with project id '''
    permission_classes = [IsAuthenticated]
    paginator = PageNumberPagination()
    paginator.page_size = 20
    filterset_class = VendorDashBoardFilter

    @staticmethod
    def get_tasks_by_projectid(request, pk):
        ''' 
        checking for team access 
        if user is admin or project_owner, it will return all tasks in that project.
        otherwise it will return only assigned tasks.
        '''
    
        project = get_object_or_404(Project.objects.all(),
                    id=pk)
        pr_managers = request.user.team.get_project_manager if request.user.team and request.user.team.owner.is_agency else []
        user_1 = request.user.team.owner if request.user.team and request.user.team.owner.is_agency and request.user in pr_managers else request.user  #####For LSP
        if project.ai_user == request.user:
            return project.get_tasks
        if project.team:
            if ((project.team.owner == request.user)|(request.user in project.team.get_project_manager)):
                return project.get_tasks
            else:
                return project.get_tasks.filter(task_info__assign_to=user_1)
        else:
            return project.get_tasks.filter(task_info__assign_to=user_1)

    # To get the account holder and project_managers
    def get_user(self):
        project_managers = self.request.user.team.get_project_manager if self.request.user.team else []
        user = self.request.user.team.owner if self.request.user.team and self.request.user in project_managers else self.request.user
        project_managers.append(user)
        return user,project_managers

    def get_object(self):
        tasks = Task.objects.order_by("-id").all()
        tasks = get_list_or_404(tasks, file__project__ai_user=self.request.user)
        tasks = authorize_list(tasks,"read",self.request.user)
        return tasks

    def list(self, request, *args, **kwargs):
        tasks = self.get_object()
        user,pr_managers = self.get_user()
        pagin_queryset = self.paginator.paginate_queryset(tasks, request, view=self)
        serlzr = VendorDashBoardSerializer(pagin_queryset, many=True,context={'request':request,'user':user,'pr_managers':pr_managers})
        return self.get_paginated_response(serlzr.data)

    def retrieve(self, request, pk, format=None):
        status = request.query_params.get('status')
        tasks = self.get_tasks_by_projectid(request=request,pk=pk)
        # filter the tasks based on working status. currently it is used in dinamalar flow.
        queryset = self.filter_queryset(tasks)
        tasks = queryset.order_by('-id')
        user,pr_managers = self.get_user()
        serlzr = VendorDashBoardSerializer(tasks, many=True,context={'request':request,'user':user,'pr_managers':pr_managers})
        return Response(serlzr.data, status=200)

# class VendorProjectBasedDashBoardView(viewsets.ModelViewSet):
#     permission_classes = [IsAuthenticated]
#     paginator = PageNumberPagination()
#     paginator.page_size = 20
    

#     def get_user(self):
#         project_managers = self.request.user.team.get_project_manager if self.request.user.team else []
#         user = self.request.user.team.owner if self.request.user.team and self.request.user in project_managers else self.request.user
#         project_managers.append(user)
#         return user,project_managers


#     def get_object(self, project_id):
#         tasks = Task.objects.filter(job__project_id=project_id).all()
#         tasks = get_list_or_404(tasks, file__project__ai_user=self.request.user)
#         return tasks

#     def list(self, request, project_id, *args, **kwargs):
#         tasks = self.get_object(project_id)
#         user,pr_managers = self.get_user()
#         serlzr = VendorDashBoardSerializer(tasks, many=True,context={'request':request,'user':user,'pr_managers':pr_managers})
#         return Response(serlzr.data, status=200)

class TM_FetchConfigsView(viewsets.ViewSet):
    def get_object(self, pk):
        project = get_object_or_404(
            Project.objects.all(), id=pk)
        return project

    def update(self, request, pk, format=None):
        project = self.get_object(pk)
        authorize(request,resource=project,action="update",actor=self.request.user)
        ser = ProjectSerializerV2(project, data=request.data, partial=True)
        if ser.is_valid(raise_exception=True):
            ser.save()
            return Response(ser.data, status=201)


class ReferenceFilesView(viewsets.ModelViewSet):
    '''
    This view is to add, delete reference files to project.
    Not using now
    '''
    serializer_class = ReferenceFileSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = None
    # https://www.django-rest-framework.org/api-guide/filtering/

    def get_object(self):
        return get_object_or_404(ReferenceFiles.objects.all(),
            id=self.kwargs.get("pk"))

    def get_project(self, project_id):
        try:
            project = get_object_or_404(Project.objects.all(),\
                        id=project_id)
        except:
            raise Http404("project_id should be int type!!!")
        return project

    def get_queryset(self):
        project_id = self.request.query_params.get("project", None)
        project = self.get_project(project_id)
        ref_files = ReferenceFiles.objects.none()
        if project:
            ref_files = project.ref_files
        return ref_files

    def create(self, request):
        files = request.FILES.getlist('ref_files')
        project_id = request.data.get("project", None)
        project = self.get_project(project_id)
        data = \
          [{"project": project_id, "ref_files": file} for file in files]
        ser = ReferenceFileSerializer(data=data, many=True)
        if ser.is_valid(raise_exception=True):
            ser.save()
            return Response(ser.data, status=201)

    def destroy(self, request, *args, **kwargs):
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            objs=authorize_list(objs,"delete",request.user)
            for obj in objs:
                obj.delete()
            return Response(status=204)
        obj=self.get_object()
        authorize(request,resource=obj,action="delete",actor=self.request.user)
        return super().destroy(request, *args, **kwargs)

@api_view(["DELETE"])
def test_internal_call(request):
    view = (ReferenceFilesView.as_view({"delete":"destroy"})\
        (request=request._request, pk=0, many="true", ids="6,7")).data
    return Response(view, status=200)


class TbxFileListCreateView(APIView):
    '''
    This view is to list and add tbx files for the project
    returns TbxFileSerializer data
    '''

    def get(self, request, project_id):
        files = TbxFile.objects.filter(project_id=project_id).all()
        serializer = TbxFileSerializer(files, many=True)
        return Response(serializer.data)

    def post(self, request, project_id):
        project=get_object_or_404(Project,id=project_id)
        authorize(request,resource=project,action="create",actor=self.request.user)
        data = {**request.POST.dict(), "tbx_file" : request.FILES.getlist('tbx_file'),'project_id':project_id}
        ser_data = TbxFileSerializer.prepare_data(data)
        serializer = TbxFileSerializer(data=ser_data,many=True)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
        return Response(serializer.data, status=201)

class TbxFileDetail(APIView):

    '''
    This view is to update and delete tbx files for the project
    returns TbxFileSerializer data
    '''

    def get_object(self, id):
        try:
            return TbxFile.objects.get(id=id)
        except TbxFile.DoesNotExist:
            return HttpResponse(status=404)

    def put(self, request, id):
        tbx_asset = self.get_object(id)
        authorize(request,resource=tbx_asset,action="update",actor=self.request.user)
        #tbx_file = request.FILES.get('tbx_file')
        job_id = request.POST.get("job_id", None)
        serializer = TbxFileSerializer(tbx_asset, data={"job" : job_id}, partial=True)
        if serializer.is_valid():
            serializer.save_update()
            return Response(serializer.data, status=200)

    def delete(self, request, id):
        tbx_asset = self.get_object(id)
        authorize(request,resource=tbx_asset,action="delete",actor=self.request.user)
        tbx_asset.delete()
        return Response(data={"Message": "Removed Terminology asset"}, status=204)

class TmxList(APIView):

    def get(self, request, project_id):
        files = TmxFile.objects.filter(project_id=project_id).all()
        serializer = TmxFileSerializer(files, many=True)
        return Response(serializer.data)

#### Glossary template lite #############
@api_view(['GET',])
def tbx_template(request):
    '''
    This function is to download tbx_template.
    '''
    response = HttpResponse(content_type='application/vnd.ms-excel')
    response['Content-Disposition'] = 'attachment; filename=Glossary_Lite.xlsx'
    xlsx_data = WriteToExcel_lite()
    response.write(xlsx_data)
    return response

class TbxTemplateUploadView(APIView):
    '''
    This function is to upload tbx_template file in project.
    '''
    def post(self, request, project_id):
        pro=get_object_or_404(Project,id=project_id)
        authorize(request,resource=pro,action="create",actor=self.request.user)
        data = {**request.POST.dict(), "tbx_template_file" : request.FILES.get('tbx_template_file')}
        data.update({'project_id': project_id})
        prep_data = TbxTemplateSerializer.prepare_data(data)

        serializer = TbxTemplateSerializer(data=prep_data)
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            saved_data = serializer.data
            file_id = saved_data.get("id")
            job_id = prep_data["job"]
            if bool(upload_template_data_to_db(file_id, job_id)):
                tbx_file = user_tbx_write(job_id, project_id)
                fl = open(tbx_file, 'rb')
                file_obj1 = DJFile(fl) #,name=os.path.basename(tbx_file))
                serializer2 = TbxFileSerializer(data={'tbx_file':file_obj1,'project':project_id,'job':job_id})
    
                if serializer2.is_valid():
                    serializer2.save()
                else:
                    return Response(serializer2.errors)
                fl.close()
                TemplateTermsModel.objects.filter(job_id = job_id).delete()
                os.remove(os.path.abspath(tbx_file))
                return Response({'msg':"Template File uploaded and TBX created & uploaded","data":serializer.data})#,"tbx_file":tbx_file})
            else:
                return Response({'msg':"Template file seems empty or partially empty. Fill up terms and try again", "data":{}},
                        status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response(serializer.errors)

@api_view(['GET',])
def tbx_download(request,tbx_file_id):
    obj = TbxFile.objects.get(id=tbx_file_id)
    tbx_asset=obj.tbx_file
    authorize(request,resource=obj,action="download",actor=request.user)
    return download_file(tbx_asset.path)



class UpdateTaskCreditStatus(APIView):
    '''
    The code defines a class UpdateTaskCreditStatus with static methods to update user credits and addon credits based on actual used credits and credit differences.
    The update_addon_credit method deducts credits from addon credit packs based on the actual used credits or credit differences.
    The update_usercredit method updates user subscription credits, considering expiry dates and available credits.
    The update_credits method orchestrates the credit update process, handling user and addon credits within a transaction.
    '''

    permission_classes = [IsAuthenticated]

    @staticmethod
    def update_addon_credit(user, query, actual_used_credits=None, credit_diff=None):
        add_ons = query.filter(Q(user=user) & Q(credit_pack_type="Addon")).\
                    filter(Q(expiry__isnull=True) | Q(expiry__gte=timezone.now())).order_by('expiry')
        if add_ons.exists():
            case = credit_diff if credit_diff != None else actual_used_credits
            for addon in add_ons:
                if addon.credits_left >= case:
                    addon.credits_left -= case
                    addon.save()
                    case = None
                    break
                else:
                    diff = case - addon.credits_left
                    addon.credits_left = 0
                    addon.save()
                    case = diff
            return False if case != None else True
        else:
            return False

    @staticmethod
    def update_usercredit(user, query, actual_used_credits):
        present = datetime.now()
        try:

            user_credit = query.get(Q(user=user) & Q(credit_pack_type__icontains="Subscription") & Q(ended_at=None))

            if present.strftime('%Y-%m-%d %H:%M:%S') <= user_credit.expiry.strftime('%Y-%m-%d %H:%M:%S'):
                if not actual_used_credits > user_credit.credits_left:
                    user_credit.credits_left -= actual_used_credits
                    user_credit.save()
                    return True
                else:
                    credit_diff = actual_used_credits - user_credit.credits_left
                    user_credit.credits_left = 0
                    user_credit.save()
                    from_addon = UpdateTaskCreditStatus.update_addon_credit( user, query, credit_diff)
                    return from_addon
            else:
                raise Exception

        except Exception as e:
            from_addon = UpdateTaskCreditStatus.update_addon_credit(user, query, actual_used_credits)
            return from_addon

    @staticmethod
    def update_credits( user, actual_used_credits):
        with transaction.atomic():
            query = UserCredits.objects.select_for_update().filter(Q(user=user))

            credit_status = UpdateTaskCreditStatus.update_usercredit(user, query, actual_used_credits)

            if credit_status:
                msg = "Successfully debited MT credits"
                status = 200
            else:
                msg = "Insufficient credits to apply MT"
                status = 424

            return {"msg" : msg}, status



########### To get User credit balance #################
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def dashboard_credit_status(request):
    '''
    This function is to return admin account credit_balance
    '''
    pr_managers = request.user.team.get_project_manager if request.user.team else []
    if request.user.is_internal_member == True and request.user in pr_managers:
        user = request.user.team.owner
    else: user = request.user
    return Response({"credits_left": user.credit_balance}, status=200)

######### Tasks Assign to vendor #################
from ai_workspace.serializers import TaskViewSerializer
class TaskViewDetail(APIView):
    def get(self, request, pk):
        try:
            item = Task.objects.get(pk=pk)
        except Task.DoesNotExist:
            return Response(status=status.HTTP_404_NOT_FOUND)
        
        serializer = TaskViewSerializer(item)
        return Response(serializer.data)


class TaskView(APIView):
    permission_classes = [IsAuthenticated]
    serializer_class = TaskSerializer

    def get_queryset(self,):
        tasks = [ task for project in get_list_or_404(Project.objects.all(), ai_user=self.request.user)
                    for task in project.get_tasks
                  ]
        return  tasks

    def get(self, request):
        tasks = self.get_queryset()
        tasks_serlzr = TaskSerializer(tasks, many=True)
        return Response(tasks_serlzr.data, status=200)

    @staticmethod
    def get_object(data):
        obj = Task.objects.filter(**data).first()
        return obj

    def post(self, request):
        print(self.request.POST.dict())
        obj = self.get_object({**request.POST.dict()})
        if obj:
            task_ser = TaskSerializer(obj)
            return Response(task_ser.data, status=200)

        task_serlzr = TaskSerializer(data=request.POST.dict(), context={"request":request,
            "assign_to": self.request.POST.get('assign_to', self.request.user.id)})#,"customer":self.request.user.id})
        if task_serlzr.is_valid(raise_exception=True):
            task_serlzr.save()
            return Response({"msg": task_serlzr.data}, status=200)

        else:
            return Response({"msg": task_serlzr.errors}, status=400)
    
    def delete(self, request, id):
        try:
            task = Task.objects.get(id=id)
        except Task.DoesNotExist:
            return Response(data={"Message": "Task not found"}, status=404)

        print("task--instance", task)
        authorize(request, resource=task, action="delete", actor=self.request.user)

        if task.task_info.filter(task_assign_info__isnull=False).exists():
            return Response(data={"Message": "Task is assigned. Unassign and Delete"}, status=400)
        
        if len(task.job.project.get_tasks) == 1:
            print("The deleting task's project contains only one task instance")
            task.job.project.delete()
        else:
            if task.file:
                if os.path.splitext(task.file.filename)[1] == ".pdf":
                    task.file.delete()
            if task.document:
                task.document.delete()
        
        try:
            task.delete()
        except Exception as e:
            print(f"An error occurred: {e}")
            return Response(data={"Message": "An error occurred while deleting the task"}, status=500)

        return Response(status=status.HTTP_204_NO_CONTENT)

    # def delete(self, request, id):
    #     task = Task.objects.get(id = id)
    #     print("task--instance",task)
    #     authorize(request,resource=task,action="delete",actor=self.request.user)
    #     # it will check for task is assigned to any editor or not. if so, it will return error message
    #     if task.task_info.filter(task_assign_info__isnull=False): ### checking if the task instance is assigned to anyone
    #         return Response(data={"Message":"Task is assigned.Unassign and Delete"},status=400)
    #     else:
    #         # if it is the single task in the project, then it will internally delete the project
    #         if len(task.job.project.get_tasks) == 1:
    #             #check_delete_default_gloss_task(task,is_single_task=True)
    #             print("the deleting task 's project contains only one task instance ")
    #             task.job.project.delete()
    #         elif task.file:
    #             if os.path.splitext(task.file.filename)[1] == ".pdf":
    #                 task.file.delete()
    #             if task.document:
    #                 task.document.delete()
    #             #check_delete_default_gloss_task(task,is_single_task=False)
    #         else:
    #             task.delete()
    #         return Response(status=status.HTTP_204_NO_CONTENT)

# def check_delete_default_gloss_task(task,is_single_task=False):
#     from ai_glex.models import TermsModel
#     from ai_workspace.serializers import ProjectQuickSetupSerializer
#     job = task.job

#     if task.is_default_glossary_task:  
#         if is_single_task:
#             project = job.project
#             ProjectQuickSetupSerializer().create_default_gloss(project=project,jobs=[job],ai_user=project.ai_user)
#             print("project with Task created for individual gloss project --> default gloss")
#         else:        
#             TermsModel.objects.filter(job=job).delete()
#             Task.objects.create_glossary_tasks_of_jobs(jobs=[job], klass=Task)
#             print("Task created for individual gloss project --> default gloss")
#     else:
#         print("Task is not --> default gloss")

################# Create Project from Temp project ################
@api_view(['POST',])
@permission_classes([AllowAny])
def create_project_from_temp_project_new(request):
    '''
    The function extracts user and project details from the request data.
    It retrieves files and jobs associated with the temporary project, prepares data for the new project, and creates a project using a serializer.
    If the serializer is valid, the new project is saved, and the serialized data is returned; otherwise, errors are returned.
    need to pass all values in list to ProjectQuickSetupSerializer.
    '''
    ai_user_id = request.POST.get("user_id")
    ai_user = AiUser.objects.get(id=ai_user_id)
    user_1 = ai_user.team.owner if ai_user.team and ai_user.team.owner.is_agency and (ai_user in ai_user.team.get_project_manager) else ai_user
    temp_proj_id = request.POST.get("temp_project")
    temp_proj =  TempProject.objects.get(temp_proj_id =temp_proj_id)
    files_list = TempFiles.objects.filter(temp_proj_id =temp_proj.id)
    jobs_list = Templangpair.objects.filter(temp_proj_id=temp_proj.id)
    mt_engine = [temp_proj.mt_engine_id]
    source_language = [str(jobs_list[0].source_language_id)]
    target_languages = [str(i.target_language_id) for i in jobs_list]
    files = [DJFile(i.files,name=i.filename) for i in files_list]
    filename,extension = os.path.splitext((files_list[0].filename))
    serializer = ProjectQuickSetupSerializer(data={'project_name':[filename +'-tmp'+ str(temp_proj.id)],\
    'source_language':source_language,'target_languages':target_languages,'files':files,'mt_engine':mt_engine},\
    context={'ai_user':ai_user,'user_1':user_1})
    if serializer.is_valid():
        serializer.save()
        return JsonResponse({"data":serializer.data},safe=False)
    else:
        return JsonResponse({"data":serializer.errors},safe=False)

##############   PROJECT ANALYSIS BY STORING ONLY COUNT DATA   ###########

class ProjectAnalysisProperty(APIView):
    '''
    This view ProjectAnalysisProperty that checks all the tasks in the project is analysed or not. 
    if analysed, it returns the word_count, char_count and seg_count of each task.
    if it is not analysed, then it will initiate celery task and return the result.
    '''

    permission_classes = [IsAuthenticated]

    @staticmethod
    def exact_required_fields_for_okapi_get_document():
        fields = ['source_file_path', 'source_language', 'target_language',
                     'extension', 'processor_name', 'output_file_path']
        return fields

    erfogd = exact_required_fields_for_okapi_get_document

    # erfogd = DocumentViewByTask.exact_required_fields_for_okapi_get_document()

    @staticmethod
    def correct_fields(data):
        check_fields = ProjectAnalysisProperty.erfogd()
        remove_keys = []
        for i in data.keys():
            if i in check_fields:
                check_fields.remove(i)
            else:
                remove_keys.append(i)
        [data.pop(i) for i in remove_keys]
        if check_fields != []:
            raise ValueError("File processing fields not set properly !!!")

    @staticmethod
    def get_data_from_docs(project):
        '''
        some tasks do not have document instance(like voice source projects), for that we will store its counts in TaskDetail table
        if tasks in project has document, then it will aggregate document word_count,char_count,segment_count detail
        '''
        proj_word_count = proj_char_count = proj_seg_count = 0
        task_words = []

        for task in project.get_mtpe_tasks:
            doc = Document.objects.get(id=task.document_id)
            proj_word_count += doc.total_word_count
            proj_char_count += doc.total_char_count
            proj_seg_count += doc.total_segment_count

            task_words.append({task.id:doc.total_word_count})

        return {"proj_word_count": proj_word_count, "proj_char_count":proj_char_count, "proj_seg_count":proj_seg_count,\
                                "task_words" : task_words }

    @staticmethod
    def get_data_from_analysis(project):
        '''
        Filter the TaskDetails by the given project and aggregate the word_count,char_count
        and seg_count of all tasks in that project.
        '''
        out = TaskDetails.objects.filter(project_id=project.id).aggregate(Sum('task_word_count'),Sum('task_char_count'),Sum('task_seg_count'))
        task_words = []
        for task in project.get_mtpe_tasks:
            task_words.append({task.id : task.task_details.first().task_word_count})
        return {"proj_word_count": out.get('task_word_count__sum'), "proj_char_count":out.get('task_char_count__sum'), \
                        "proj_seg_count":out.get('task_seg_count__sum'),
                        "task_words":task_words}

    @staticmethod
    def get_analysed_data(project_id):
        project = Project.objects.get(id=project_id)
        if project.is_all_doc_opened:
            return ProjectAnalysisProperty.get_data_from_docs(project)
        else:
            return ProjectAnalysisProperty.get_data_from_analysis(project)

    @staticmethod
    def analyse_project(project_id):
        '''
        This function is to return project_word_count, char_count and segment_count and task_words
        '''
        project = Project.objects.get(id=project_id)
        project_tasks = Project.objects.get(id=project_id).get_mtpe_tasks
        tasks = []
        # Filtering the tasks which is not having task_count_details.
        for _task in project_tasks:
            if _task.task_details.first() == None:
                tasks.append(_task)
        task_words = []
        file_ids = []

        for task in tasks:
            if task.file_id not in file_ids:
                ser = TaskSerializer(task)
                data = ser.data
                ProjectAnalysisProperty.correct_fields(data)
                params_data = {**data, "output_type": None}
                res_paths = get_res_path(params_data["source_language"])
                #calling springAPI getDocument to process the file and get the json which contains segments, textunits, and total_count_details
                doc = requests.post(url=f"http://{spring_host}:8080/getDocument/", data={
                    "doc_req_params":json.dumps(params_data),
                    "doc_req_res_params": json.dumps(res_paths)
                })

                try:
                    if doc.status_code == 200 :
                        doc_data = doc.json()
                        task_write_data = json.dumps(doc_data, default=str)
                        # Writing json data into models Segment, TextUnit by calling celery task
                        write_doc_json_file.apply_async((task_write_data, task.id),queue='high-priority')
                        
                        #storing count_details from json to TaskDetail table
                        task_detail_serializer = TaskDetailSerializer(data={"task_word_count":doc_data.get('total_word_count', 0),
                                                                "task_char_count":doc_data.get('total_char_count', 0),
                                                                "task_seg_count":doc_data.get('total_segment_count', 0),
                                                                "task": task.id,"project":project_id}
                                                                     )

                        if task_detail_serializer.is_valid(raise_exception=True):
                            task_detail_serializer.save()
                        else:
                            print("error-->", task_detail_serializer.errors)
                    else:
                        # if file is having any issue in  processing in spring
                        logger.debug(msg=f"error raised while process the document, the task id is {task.id}")
                        raise  ValueError("Sorry! Something went wrong with file processing.")
                except:
                    print("No entry")
                # to update processed file_ids
                file_ids.append(task.file_id)

            else:
                # If the file is processed already, then it will just duplicate the existing details for new job.
                print("*************  File taken only once  **************")
                tasks = [i for i in Task.objects.filter(file_id=task.file_id)]
                task_details = TaskDetails.objects.filter(task__in = tasks).first()
                task_details.pk = None
                task_details.task_id = task.id
                task_details.save()
               

        [task_words.append({task.id : task.task_details.first().task_word_count})for task in project.get_mtpe_tasks]
        #Adding all the word_count, char_count and seg_count of tasks in the given project using aggregate
        out = TaskDetails.objects.filter(project_id=project_id).aggregate(Sum('task_word_count'),Sum('task_char_count'),Sum('task_seg_count'))
        
        return {"proj_word_count": out.get('task_word_count__sum'), "proj_char_count":out.get('task_char_count__sum'), \
                        "proj_seg_count":out.get('task_seg_count__sum'),
                        "task_words":task_words}

    @staticmethod
    def get(project_id):

        if bool(Project.objects.get(id=project_id).is_proj_analysed):
            return ProjectAnalysisProperty.get_analysed_data(project_id)
        else:
            return ProjectAnalysisProperty.analyse_project(project_id)

#######################################

class ProjectAnalysis(APIView):
    '''
    To check whether the project is analysed or not. if it is not analysed then it will call 
    analyse_project() and return count_details
    if it is analysed already, then it will return the count by aggregating all the tasks in project from TaskDetail model.
    '''
    permission_classes = [IsAuthenticated]
    def get(self, request, project_id):
        return Response(ProjectAnalysisProperty.get(project_id))

#########################################


########### To send message notifying assign status #####################
def msg_send(sender,receiver,task,step):
    obj = Task.objects.get(id=task)
    work = "Post Editing" if int(step) == 1 else "Reviewing"
    proj = obj.job.project.project_name
    receivers = []
    receivers =  receiver.team.get_project_manager if (receiver.team and receiver.team.owner.is_agency) else []
    receivers.append(receiver)
    
    for i in receivers:
        thread_ser = ThreadSerializer(data={'first_person':sender.id,'second_person':i.id})
        if thread_ser.is_valid():
            thread_ser.save()
            thread_id = thread_ser.data.get('id')
        else:
            thread_id = thread_ser.errors.get('thread_id')
        
        if thread_id:
            message = "You have been assigned a new task in "+proj+ " for "+ work +"."
            msg = ChatMessage.objects.create(message=message,user=sender,thread_id=thread_id)
            notify.send(sender, recipient=i, verb='Message', description=message,thread_id=int(thread_id))


class TaskAssignUpdateView(viewsets.ViewSet):

    '''
    This view is to update Task Assign Information, 
    it takes task, step, and reassigned as the input
    In this we can also add instruction file addition and deletion for particular task.
    '''
    permission_classes = [IsAuthenticated]

    def update(self, request,pk=None):
        task = request.POST.get('task')
        step = request.POST.get('step')
        reassigned = request.POST.get('reassigned',False)
        file = request.FILES.getlist('instruction_file')
        req_copy = copy.copy( request._request)
        req_copy.method = "DELETE"

        inst = Task.objects.get(id=task)
       

        file_delete_ids = self.request.query_params.get(\
            "file_delete_ids", [])

        # To delete InstructionFiles by getting list of ids
        if file_delete_ids:
            file_res = InstructionFilesView.as_view({"delete": "destroy"})(request=req_copy,\
                        pk='0', many="true", ids=file_delete_ids)
        if not task:
            return Response({'msg':'Task Id required'},status=status.HTTP_400_BAD_REQUEST)
        
        task_assign = TaskAssign.objects.get(Q(task_id = task) & Q(step_id = step) & Q(reassigned=reassigned))
       
        if file:
            serializer =TaskAssignUpdateSerializer(task_assign,data={**request.POST.dict(),'files':file},context={'request':request},partial=True)
        else:
            serializer =TaskAssignUpdateSerializer(task_assign,data={**request.POST.dict()},context={'request':request},partial=True)
        if serializer.is_valid():
            serializer.save()
            if request.POST.get('account_raw_count'):
                weighted_count_update.apply_async((None,None,task_assign.task_assign_info.assignment_id,),queue='medium-priority')
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        return Response(task, status=status.HTTP_200_OK)



@api_view(['POST',])
@permission_classes([IsAuthenticated])
def bulk_task_accept(request):
    '''
    This function is to bulk accept the tasks from editor side in PO Accept.
    input: task_accept_detail (task,step,reassigned)
    Update TaskAssign model with the status task_accepted
    output: return TaskAssignUpdateSerializer data
    '''
    task_accept_detail = request.POST.get('task_accept_detail')
    task_accept_detail = json.loads(task_accept_detail)
    for i in task_accept_detail:
        try:
            task_assign = TaskAssign.objects.get(Q(task_id = i.get('task')) & Q(step_id = i.get('step')) & Q(reassigned=i.get('reassigned'))) 
            serializer =  TaskAssignUpdateSerializer(task_assign,data={'task_ven_status':'task_accepted'},context={'request':request},partial=True)
            if serializer.is_valid():
                serializer.save()
            else:
                print("Error--------->",serializer.errors)
        except:
            pass
    return Response({'msg':'Task Accept Succeded'})

class TaskAssignInfoCreateView(viewsets.ViewSet):
    permission_classes = [IsAuthenticated]

    def list(self,request):
        tasks = request.GET.getlist('tasks')
        step = request.GET.get('step')
        reassigned = request.GET.get('reassigned',False)

        tsks = get_list_or_404(Task,id__in=tasks)
        task=authorize_list(tsks,"read",self.request.user)
        task=[i.id for i in task]

        try:
            task_assign_info = TaskAssignInfo.objects.filter(Q(task_assign__task_id__in = task) & Q(task_assign__reassigned = reassigned))
        except TaskAssignInfo.DoesNotExist:
            return HttpResponse(status=404)

        ser = TaskAssignInfoSerializer(task_assign_info,many=True)
        return Response(ser.data)

    def history(self,instance):
        # To record previous assigned user, number of segments confirmed and unassigned_by user
        segment_count=0 if instance.task_assign.task.document == None else instance.task_assign.task.get_progress.get('confirmed_segments')
        task_history = TaskAssignHistory.objects.create(task_assign =instance.task_assign,\
                                                        previous_assign_id=instance.task_assign.assign_to_id,\
                                                        task_segment_confirmed=segment_count,unassigned_by=self.request.user)


    def reassign_check(self,tasks):
        #To check whether the task is reassigned or not.
        user = self.request.user.team.owner if self.request.user.team else self.request.user

        if user.is_agency == True:
            for i in tasks:
                print(i)
                if TaskAssignInfo.objects.filter(task_assign__task = i).filter(task_assign__reassigned=False).exists() == False:
                    return "There is no assign. you can't reassign" 
            return None
        else:
            return "user is not an agency. Reassign is not allowed"
        

    @integrity_error
    def create(self,request):

        ''' 
            Create TaskAssignInfo by getting list of tasks and assign to. 
            Checks for reassigned. it is the boolean field to be set to know whether it is the first time sign or reassign
            It edits assign_to in TaskAssign(change assign_to from self_assign to editor) and then create TaskAssignInfo 
            It also create InstructionFile associated with task. 
        '''

        step = request.POST.get('step')
        task_assign_detail = request.POST.get('task_assign_detail')
        files = request.FILES.getlist('instruction_file')
        sender = self.request.user
        receiver = request.POST.get('assign_to')
        reassign = request.POST.get('reassigned') 
        own_agency_email = os.getenv("AILAYSA_AGENCY_EMAIL")
    
        Receiver = AiUser.objects.get(id = receiver)
        data = request.POST.dict()
        ################################Need to change########################################
        user = request.user.team.owner  if request.user.team  else request.user
        if Receiver.email == own_agency_email:
            HiredEditors.objects.get_or_create(user_id=user.id,hired_editor_id=receiver,defaults = {"role_id":2,"status":2,"added_by_id":request.user.id})
        ##########################################################################################
       
        hired_editors = sender.get_hired_editors if sender.get_hired_editors else []

       
        assignment_id = create_assignment_id()
        extra = {'assignment_id':assignment_id,'files':files}
        final =[]
        task_assign_detail = data.pop('task_assign_detail')
        task_assign_detail = json.loads(task_assign_detail)    #
        tasks = list(itertools.chain(*[d['tasks'] for d in task_assign_detail]))

        # For authorization
        tsks = Task.objects.filter(id__in=tasks)
        for tsk in tsks:
            authorize(request, resource=tsk, actor=request.user, action="read")

        if reassign == 'true':
            msg = self.reassign_check(tasks)
            if msg:
                return Response({'Error':msg},status=400)

        for i in task_assign_detail: 
            i.update(data) 
            i.update(extra)
            final.append(i)
        

        with transaction.atomic():
            serializer = TaskAssignInfoSerializer(data=final,context={'request':request},many=True)
            if serializer.is_valid():
                serializer.save()
            
        task_assgn_objs = TaskAssignInfo.objects.filter(assignment_id = assignment_id)
        if task_assgn_objs.count() >0 :
            # this is to update count based on tm_analysis
            weighted_count_update.apply_async((receiver,sender.id,assignment_id),queue='medium-priority')

            try:msg_send(sender,Receiver,tasks[0],step)
            except:
                print("Inside Exception")
                pass
            # if Receiver in hired_editors:
            #     ws_forms.task_assign_detail_mail(Receiver,assignment_id)
            # notify.send(sender, recipient=Receiver, verb='Task Assign', description='You are assigned to new task.check in your project list')
            return Response({"msg":"Task Assigned"})
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self,request):
        
        ''' 
        Flow 1: Customer--->Editor. if taskassign is deleted, then delete the taskassign info and reassign back to customer(self-assign) and update status and client_response
        Flow 2: Customer---->Agency---->Editor. if taskassign is reassigned and deleted, then delete the taskassign info and assign back to the first user(agency) and update status and client_response
        In Both flows, History is maintained to know details of previous assigned user and their word count at the time of reassign.
        '''

        task = request.GET.getlist('task')
        steps = request.GET.getlist('step')
        reassigned = request.GET.get('reassigned',False)
        task_assign_info_ids = request.GET.getlist('task_assign_info')
        if task and steps:
            assigns = TaskAssignInfo.objects.filter(Q(task_assign__task_id__in=task) & Q(task_assign__step_id__in=steps) & Q(task_assign__reassigned=reassigned))
        if task_assign_info_ids:
            assigns = TaskAssignInfo.objects.filter(id__in = task_assign_info_ids )
        for obj in assigns:
            authorize(request, resource=obj, actor=request.user, action="read")
            try:
                po_modify(obj.id,['unassigned',])
            except BaseException as e:
                logger.error(f"po unassign error id :{obj.id} -ERROR:{str(e)}")
            user = obj.task_assign.task.job.project.ai_user
            with transaction.atomic():
                try:self.history(obj)
                except:
                    pass
                if obj.task_assign.reassigned == True:
                    obj.task_assign.assign_to = self.request.user.team.owner if self.request.user.team else self.request.user #if unassigned..it is assigned back to LSP 
                    obj.task_assign.status = 1
                    obj.task_assign.client_response = None
                    obj.task_assign.save()
                    role = get_assignment_role(obj,obj.task_assign.step,obj.task_assign.reassigned)
                    assigned_user = obj.task_assign.assign_to
                    unassign_task(assigned_user,role,obj.task_obj)   
                    obj.delete()
                else:
                    reassigns = TaskAssign.objects.filter(Q(task=obj.task_assign.task) & Q(step=obj.task_assign.step) & Q(reassigned = True))
                    if reassigns:
                        try:obj_1 = reassigns.first().task_assign_info
                        except:obj_1=None
                        
                        if obj_1:
                            self.history(obj_1)
                            obj_1.task_assign.assign_to = user
                            obj_1.task_assign.status = 1
                            obj_1.task_assign.client_response = None
                            obj_1.task_assign.save()
                            obj_1.delete()
                        else:
                            rr = reassigns.first()
                            rr.assign_to = user
                            rr.save()
                    assigned_user = obj.task_assign.assign_to
                    obj.task_assign.assign_to = user
                    obj.task_assign.status = 1
                    obj.task_assign.client_response = None
                    obj.task_assign.save()
                    role = get_assignment_role(obj,obj.task_assign.step,obj.task_assign.reassigned)
                    unassign_task(assigned_user,role,obj.task_obj)             
                    obj.delete()
        return Response({"msg":"Tasks Unassigned Successfully"},status=200)

# Not using now
# @api_view(['GET',])
# @permission_classes([IsAuthenticated])
# def get_assign_to_list(request):
#     project = request.GET.get('project')
#     job_id = request.GET.get('job',None)
#     proj = Project.objects.get(id = project)
#     jobs = Job.objects.filter(id = job_id) if job_id else proj.get_jobs
#     internalmembers = []
#     hirededitors = []
#     try:
#         internal_team = proj.ai_user.team.internal_member_team_info.filter(role = 2).order_by('id')
#         for i in internal_team:
#             try:profile = i.internal_member.professional_identity_info.avatar_url
#             except:profile = None
#             internalmembers.append({'name':i.internal_member.fullname,'id':i.internal_member_id,\
#                                     'status':i.get_status_display(),'avatar': profile})
#     except:
#         print("No team")
#     external_team = proj.ai_user.team.owner.user_info.filter(role=2) if proj.ai_user.team else proj.ai_user.user_info.filter(role=2)
#     hirededitors = find_vendor(external_team,jobs)
#     return JsonResponse({'internal_members':internalmembers,'Hired_Editors':hirededitors})

# def find_vendor(team,jobs):
#     externalmembers=[]
#     for j in team:
#         for job in jobs:
#             try:profile = j.hired_editor.professional_identity_info.avatar_url
#             except:profile = None
#             vendor = j.hired_editor.vendor_lang_pair.filter(Q(source_lang_id=job.source_language.id)&Q(target_lang_id=job.target_language.id)&Q(deleted_at=None))
#             if vendor:
#                 externalmembers.append({'name':j.hired_editor.fullname,'id':j.hired_editor_id,'status':j.get_status_display(),"avatar":profile,\
#                                         'lang_pair':job.source_language.language+'->'+job.target_language.language,\
#                                         'unique_id':j.hired_editor.uid})
#     return externalmembers


# Not using Now
# class ProjectListView(viewsets.ModelViewSet):
#     permission_classes = [IsAuthenticated]
#     serializer_class = ProjectListSerializer
#     filter_backends = [DjangoFilterBackend,SearchFilter,OrderingFilter]
#     paginator = PageNumberPagination()
#     paginator.page_size = 20
#     search_fields = ['project_name','id']
    

#     def get_queryset(self):
#         print(self.request.user)
#         queryset = Project.objects.prefetch_related('project_jobs_set','project_jobs_set__job_tasks_set__task_info').filter(Q(ai_user = self.request.user)|Q(team__owner = self.request.user)\
#                     |Q(team__internal_member_team_info__in = self.request.user.internal_member.filter(role=1))).distinct().order_by('-id').only('id','project_name')
#         return queryset


#     def list(self,request):
#         queryset = self.filter_queryset(self.get_queryset())
#         pagin_tc = self.paginator.paginate_queryset(queryset, request , view=self)
#         serializer = ProjectListSerializer(pagin_tc, many=True, context={'request': request})
#         data_1 = [i for i in serializer.data if i.get('assignable')==True ]
#         response = self.get_paginated_response(data_1)
#         return response
        


class WriterProjectListView(viewsets.ModelViewSet):
    '''
    This is to list only file based projects(exclude glossary and other types) in ai_writer.
    Existing project list in AIWriter
    Input: Request object
    Output: ProjectListSerializer data
    '''
    permission_classes = [IsAuthenticated]
    serializer_class = ProjectListSerializer
    paginator = PageNumberPagination()
    paginator.page_size = 20

    def list(self,request):
        queryset = Project.objects.filter(Q(ai_user = request.user)|Q(team__owner = request.user)\
                        |Q(team__internal_member_team_info__in = request.user.internal_member.filter(role=1))).filter(project_type_id__in=[1,2]).distinct().order_by('-id')
        pagin_tc = self.paginator.paginate_queryset(queryset, request , view=self)
        serializer = ProjectListSerializer(pagin_tc, many=True, context={'request': request})
        response = self.get_paginated_response(serializer.data)
        return response



@permission_classes([IsAuthenticated])
@api_view(['GET',])
def tasks_list(request):
    '''
    This function is to return the task_list taking input as job or project
    it takes input as job or project
    Takes tasks from project and pass it to VendorDashboardSerializer and returns serializer.data
    '''
    project_managers = request.user.team.get_project_manager if request.user.team else []
    user = request.user.team.owner if request.user.team and request.user in project_managers else request.user
    project_managers.append(user)
    job_id = request.GET.get("job")
    project_id = request.GET.get('project')
    if job_id:
        job = Job.objects.get(id = job_id)
        project_id = job.project.id
    vbd = VendorDashBoardView
    res = vbd.get_tasks_by_projectid(request=request,pk=project_id)
    if job_id:
        res = [i for i in res if i.job == job ]
    try:
        tasks=[]
        for task in res:
            if (task.job.target_language == None):
                if (task.file.get_file_extension == '.mp3'):
                    tasks.append(task)
                else:pass
            else:tasks.append(task)
        ser = VendorDashBoardSerializer(tasks,many=True,context={'request':request,'user':user,'pr_managers':project_managers})
        return Response(ser.data)
    except:
        return JsonResponse({"msg":"something went wrong"})


@api_view(['GET',])
def instruction_file_download(request,instruction_file_id):
    '''
    This function is to download instruction file associated with task_assign.
    '''
    inst = Instructionfiles.objects.get(id=instruction_file_id)
    authorize(request,resource=inst,actor=request.user,action="download")
    instruction_file =inst.instruction_file
    if instruction_file:
        return download_file(instruction_file.path)
    else:
        return JsonResponse({"msg":"no file associated with it"})



class AssignToListView(viewsets.ModelViewSet):
    '''
    This view is to list the editors who works in given language pair.
    In GetAssignToSerializer, it will filter editors with list of lang_pairs taken from job 
    or if project is the input then jobs = project.get_jobs.
    Returns GetAssignToSerializer data
    '''
    permission_classes = [IsAuthenticated]
    def list(self, request, *args, **kwargs):
        project = self.request.GET.get('project')
        job = self.request.GET.getlist('job')
        reassign = self.request.GET.get('reassign',None)
        pro = Project.objects.get(id = project)
        task=Task.objects.filter(job__project__id=project)
        tsk=filter_authorize(request,task,"read",self.request.user)
        if not tsk:
            return JsonResponse({"msg":"You do not have permission to perform this action."})
        try:
            job_obj = Job.objects.filter(id__in = job).first() 
        except Job.DoesNotExist:
            pass
       
        if reassign:
            # sending user as agency user. so that editors(who works in given lang pair) associated with agency list will be return 
            user = self.request.user.team.owner if self.request.user.team else self.request.user
            serializer = GetAssignToSerializer(user,context={'request':request,'pro_user':pro.ai_user})
        else:
            user =pro.ai_user   
            serializer = GetAssignToSerializer(user,context={'request':request})
        return Response(serializer.data, status=201)




class InstructionFilesView(viewsets.ModelViewSet):
    '''
    This class is to delete instruction files which is given in task_assign. 
    Internally called in TaskAssignUpdate
    '''
    serializer_class = InstructionfilesSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = None

    def get_object(self, many=False):
        objs = []
        obj = None
        if not many:
            try:
                obj = get_object_or_404(Instructionfiles.objects.all(),\
                    id=self.kwargs.get("pk"))
            except:
                raise Http404
            return  obj

        objs_ids_list =  self.kwargs.get("ids").split(",")

        for obj_id in objs_ids_list:

            try:
                objs.append(get_object_or_404(Instructionfiles.objects.all(),\
                    id=obj_id))
            except:
                raise Http404
        return objs

    def destroy(self, request, *args, **kwargs):
        if kwargs.get("many")=="true":
            objs = self.get_object(many=True)
            for obj in objs:
                obj.delete()
            return Response(status=204)
        return super().destroy(request, *args, **kwargs)


class StepsView(viewsets.ViewSet):
    permission_classes = [AllowAny,]
    def list(self,request):
        queryset = Steps.objects.all()
        serializer = StepsSerializer(queryset,many=True)
        return Response(serializer.data)

#############Not using now################################
# class CustomWorkflowCreateView(viewsets.ViewSet):
#     permission_classes = [IsAuthenticated]
#     def list(self,request):
#         queryset = Workflows.objects.all()
#         serializer = WorkflowsSerializer(queryset,many=True)
#         return Response(serializer.data)

#     def create(self,request):
#         steps = request.POST.getlist('steps')
#         serializer = WorkflowsStepsSerializer(data={**request.POST.dict(),"user":self.request.user.id,"steps":steps})
#         if serializer.is_valid():
#             serializer.save()
#             return Response({"msg":"workflow created"})
#         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

#     def update(self,request,pk):
#         queryset = Workflows.objects.all()
#         steps = request.POST.getlist('steps')
#         step_delete_ids = request.POST.getlist('step_delete_ids')
#         workflow = get_object_or_404(queryset, pk=pk)
#         if step_delete_ids:
#             [WorkflowSteps.objects.filter(workflow=workflow,steps=i).delete() for i in step_delete_ids]
#         serializer= WorkflowsStepsSerializer(workflow,data={**request.POST.dict(),"steps":steps},partial=True)
#         if serializer.is_valid():
#             serializer.save()
#             return Response(serializer.data)
#         else:
#             return Response(serializer.errors)

#     def delete(self,request,pk):
#         queryset = Workflows.objects.all()
#         obj = get_object_or_404(queryset, pk=pk)
#         obj.delete()
#         return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(["GET"])
def previously_created_steps(request):
    used_steps = []
    pr = Project.objects.filter(Q(created_by = request.user)\
         & Q(proj_steps__isnull=False) & ~Q(project_type=1)).distinct()
    for obj in pr:
        if obj.get_steps_name not in [step for step in used_steps]:
            used_steps.append(obj.get_steps_name)
    return Response({'used_steps':used_steps})


def file_write(pr):
    # used in project_download (Instant project Download)
    for i in pr.get_tasks:
        express_obj = ExpressProjectDetail.objects.filter(task=i).first()
        file_name,ext = os.path.splitext(i.file.filename)
        target_filename = file_name + "_out" +  "(" + i.job.source_language_code + "-" + i.job.target_language_code + ")" + ext
        target_filepath = os.path.join(pr.project_dir_path,'source',target_filename)
        source_filename = file_name + "_source" +  "(" + i.job.source_language_code + "-" + i.job.target_language_code + ")" + ext
        source_filepath = os.path.join(pr.project_dir_path,'source',source_filename)
        if express_obj.source_text:
            with open(source_filepath,'w') as f:
                f.write(express_obj.source_text)
        if express_obj.target_text:
            with open(target_filepath,'w') as f:
                f.write("Source:" + "\n")
                f.write(express_obj.source_text) 
                f.write('\n')
                f.write("---------" + "\n\n")
                f.write("Standard:" + "\n")
                target = express_obj.target_text if express_obj.target_text else ''
                f.write(target)
                f.write('\n')
                f.write("---------" + "\n\n")
                rewrite_obj = express_obj.express_src_text.filter(customize__customize='Rewrite')
                if rewrite_obj.exists():
                    f.write("Rewrite:" + "\n")
                    f.write(rewrite_obj.last().final_result)
                    f.write("\n")
                    f.write("---------" + "\n\n")
                simplified_obj = express_obj.express_src_text.filter(customize__customize='Simplify')
                if simplified_obj.exists():
                    f.write("Simplify:" + "\n")
                    f.write(simplified_obj.last().final_result)
                    f.write("\n")
                    f.write("---------" + "\n\n")
                shorten_obj =express_obj.express_src_text.filter(customize__customize='Shorten')
                if shorten_obj.exists():
                    f.write("Shorten:" + "\n")
                    f.write(shorten_obj.last().final_result)
                    f.write("\n")
                    f.write("---------" + "\n\n")
                


@api_view(["GET"])
#@permission_classes([AllowAny])
def project_download(request,project_id):
    '''
    This function is for project zip download. 
    giving the path of project folder which contains all source and target files and zip it and returns file path
    '''
    pr = Project.objects.get(id=project_id)
    authorize(request,resource=pr,action="download",actor=request.user)
    if pr.project_type_id == 5:
        file_write(pr)

    elif pr.project_type_id not in [3,5]:
        for i in pr.get_mtpe_tasks:
            if i.document:
                from ai_workspace_okapi.api_views import DocumentToFile
                res_1 = DocumentToFile.document_data_to_file(request,i.document.id)

    if os.path.exists(os.path.join(pr.project_dir_path,'source')):
        shutil.make_archive(pr.project_name, 'zip', pr.project_dir_path + '/source')
        if os.path.exists(os.path.join(pr.project_dir_path,'Audio')):
            shutil.make_archive(pr.project_name, 'zip', pr.project_dir_path + '/Audio')
        res = download_file(pr.project_name+'.zip')
        os.remove(pr.project_name+'.zip')
        return res
    else:
        return Response({'msg':'something went wrong'},status=400)



class ShowMTChoices(APIView):
    '''
    This class is to return all MT results for given text.
    It will call all the MT APIs through get_translation() and return translations 
    '''
    permission_classes = [AllowAny]

    @staticmethod
    def get_lang_code(lang_id):
        return LanguagesLocale.objects.filter(language_id = lang_id).first().locale_code

    @staticmethod
    def reduce_text(text,lang_code):
        '''
        split the text into sentences and checks for total word count of 100. if it exceeds then strip it and return the text
        else return the input text.
        '''
        punctuation='''!"#$%&'``()*+,-./:;<=>?@[\]^`{|}~_'''
        info =''
        lang_list = ['hi','bn','or','ne','pa']
        if lang_code in lang_list:
            sents = sentence_split(text, lang_code, delim_pat='auto')
            # sents = regex.split(u"([.।?!])?[\n]+|[.।?!] ", text)
        else:
            sents = nltk.sent_tokenize(text)
        for i in sents:
            info =info +' '+ i
            nltk_tokens = nltk.word_tokenize(info)
            count = len([word for word in nltk_tokens if word not in punctuation])
            if count in range(90,100) or count>100:
                return info.lstrip()
            else:
                continue
        return info.lstrip()


    def post(self, request):
        user = None if request.user.is_anonymous == True else request.user
        data = request.POST.dict()
        text = data.get("text", "")
        target_languages = json.loads(data["target_language"])
        sl_code = json.loads(data["source_language"])
        text_1 = self.reduce_text(text,self.get_lang_code(sl_code))
        res = {}
        for tl in target_languages:
            mt_responses = {}
            for mt_engine in AilaysaSupportedMtpeEngines.objects.all():
                if user:
                    initial_credit = user.credit_balance.get("total_left")
                    consumable_credits =  get_consumable_credits_for_text(text_1,source_lang=self.get_lang_code(sl_code),target_lang=self.get_lang_code(tl))
                    if initial_credit > consumable_credits:
                        try:
                            mt_responses[mt_engine.name] = get_translation(mt_engine.id, text_1, ShowMTChoices.get_lang_code(sl_code), ShowMTChoices.get_lang_code(tl),user_id=user.id)
                        except:
                            mt_responses[mt_engine.name] = None
                    else:
                        mt_responses[mt_engine.name] = 'Insufficient Credits'
                    res[tl] = mt_responses
                else:
                    try:
                        mt_responses[mt_engine.name] = get_translation(mt_engine.id, text_1, ShowMTChoices.get_lang_code(sl_code), ShowMTChoices.get_lang_code(tl))
                    except:
                        mt_responses[mt_engine.name] = None
                    res[tl] = mt_responses
                    
        return Response(res, status=status.HTTP_200_OK)


###########################Transcribe Short File#####################################

def transcribe_short_file(speech_file,source_code,obj,length,user,hertz):
    '''
    This is to transcribe short files with google API 
    and update transcripted text in TaskTranscriptDetail model using TaskTranscriptDetailSerializer
    and returns serializer.data
    '''

    news_transcribe = False
    if obj:
        ai_user = obj.job.project.ai_user
    else:
        ai_user = user
        news_transcribe = True
    client = speech.SpeechClient()

    with io.open(speech_file, "rb") as audio_file:
        content = audio_file.read()

    audio = speech.RecognitionAudio(content=content)

    sample_hertz = hertz if hertz >= 48000 else 8000

    config = speech.RecognitionConfig(encoding=speech.RecognitionConfig.AudioEncoding.MP3,sample_rate_hertz=sample_hertz,language_code=source_code, enable_automatic_punctuation=True,)
    try:
        response = client.recognize(config=config, audio=audio)
        transcript=''
        for result in response.results:
            transcript += result.alternatives[0].transcript

        file_length = int(response.total_billed_time.seconds)
        
        record_api_usage.apply_async(("GCP","Transcription",ai_user.uid,ai_user.email,file_length), queue='low-priority')

        if news_transcribe:
            return {"transcripted_text":transcript,"audio_file_length":file_length}
        ser = TaskTranscriptDetailSerializer(data={"transcripted_text":transcript,"task":obj.id,"audio_file_length":file_length,"user":user.id})
        if ser.is_valid():
            ser.save()
            return (ser.data)
        return (ser.errors)
    except:
        return ({'msg':'Something  went wrong in Google Cloud Api'})




###########################Transcribe Long File##############################

def upload_blob(bucket_name, source_file_name, destination_blob_name):
    """Uploads a file to the bucket."""
    storage_client = storage.Client()
    bucket = storage_client.get_bucket(bucket_name)
    blob = bucket.blob(destination_blob_name)

    blob.upload_from_filename(source_file_name)


def delete_blob(bucket_name, blob_name):
    """Deletes a blob from the bucket."""
    storage_client = storage.Client()
    bucket = storage_client.get_bucket(bucket_name)
    blob = bucket.blob(blob_name)

    blob.delete()



def transcribe_long_file(speech_file,source_code,filename,obj,length,user,hertz):
    ai_user = obj.job.project.ai_user
    bucket_name = os.getenv("BUCKET")
    source_file_name = speech_file
    destination_blob_name = filename
    
    #Upload long audio file in bucket
    upload_blob(bucket_name, source_file_name, destination_blob_name)

    gcs_uri = os.getenv("BUCKET_URL") + filename
    transcript = ''
    sample_hertz = hertz if hertz >= 48000 else 8000
    client = speech.SpeechClient()
    audio = speech.RecognitionAudio(uri=gcs_uri)

    config =  speech.RecognitionConfig(encoding=speech.RecognitionConfig.AudioEncoding.MP3,sample_rate_hertz=sample_hertz,language_code=source_code, enable_automatic_punctuation=True,)


    # Detects speech in the audio file and call the API
    operation = client.long_running_recognize(config=config, audio=audio)
    response = operation.result(timeout=10000)
    for result in response.results:
        #concatenate the results
        transcript += result.alternatives[0].transcript
    
    # Take the file_length from response
    file_length = int(response.total_billed_time.seconds)

    #Delete long audio file in bucket
    delete_blob(bucket_name, destination_blob_name)

    #Record API Usage
    record_api_usage.apply_async(("GCP","Transcription",ai_user.uid,ai_user.email,file_length), queue='low-priority')

    #Save the transcripted_data in TaskTranscriptDetail model using TaskTranscriptDetailSerializer and return serializer.data
    ser = TaskTranscriptDetailSerializer(data={"transcripted_text":transcript,"task":obj.id,"audio_file_length":length,"user":user.id})
    if ser.is_valid():
        ser.save()
        return (ser.data)
    return (ser.errors)


def audio_read(speech_file):
    # used in transcribe_file() to return length and hertz
    audio = AudioSegment.from_file(speech_file)
    length = int(audio.duration_seconds)###seconds####
    hertz = audio.frame_rate
    return length,hertz




################################speech-to-text##################################
@api_view(["POST"])
@permission_classes([IsAuthenticated])
def transcribe_file(request):
    '''
    This is to convert audio file to text using google's API.
    In this, we first calculate the audio file length. 
    If the length is greater than 60 seconds then it will intiate transcribe_long_file celery task and return task status.
    else it calls transcribe_short_file and return TaskTranscriptDetailSerializer data.
    '''
    from ai_workspace.models import MTonlytaskCeleryStatus
    task_id = request.POST.get('task')
    user = request.user
    target_language = request.POST.getlist('target_languages')
    queryset = TaskTranscriptDetails.objects.filter(task_id = task_id)
    if queryset:
        ser = TaskTranscriptDetailSerializer(queryset,many=True)
        return Response(ser.data)
    ins = MTonlytaskCeleryStatus.objects.filter(Q(task_id=task_id) & Q(task_name = 'transcribe_long_file_cel')).last()
    state = transcribe_long_file_cel.AsyncResult(ins.celery_task_id).state if ins else None

    if state == 'PENDING' or state == 'STARTED':
        return Response({'msg':'Transcription is ongoing. Pls Wait','celery_id':ins.celery_task_id},status=400)
    else:
        obj = Task.objects.get(id = task_id)
        project = obj.job.project
        account_debit_user = project.team.owner if project.team else project.ai_user
        source = [obj.job.source_language.id]
        source_code = obj.job.source_language_code
        filename = obj.file.filename
        speech_file = obj.file.file.path
        try:
            length,hertz = audio_read(speech_file)
            # audio = AudioSegment.from_file(speech_file)
            # length = int(audio.duration_seconds)###seconds####
            # hertz = audio.frame_rate
        except:
            length=None
        #print("Length in main----->",length)
        if length==None:
            return Response({'msg':'something wrong in input file'},status=400)
        initial_credit = account_debit_user.credit_balance.get("total_left")
        consumable_credits = get_consumable_credits_for_speech_to_text(length)

        if initial_credit > consumable_credits:
            if length and length<60:
                res = transcribe_short_file(speech_file,source_code,obj,length,user,hertz)
                if res.get('msg') == None:
                    consumable_credits = get_consumable_credits_for_speech_to_text(res.get('audio_file_length'))
                    debit_status, status_code = UpdateTaskCreditStatus.update_credits(account_debit_user, consumable_credits)
            else:
                ins = MTonlytaskCeleryStatus.objects.filter(Q(task_id=obj.id) & Q(task_name = 'transcribe_long_file_cel')).last()
                state = transcribe_long_file_cel.AsyncResult(ins.celery_task_id).state if ins else None
            
                if state == 'PENDING' or state == 'STARTED':
                    return Response({'msg':'Transcription is ongoing. Pls Wait','celery_id':ins.celery_task_id},status=400)
                elif (not ins) or state == 'FAILURE' or state == 'REVOKED':#need to revert credits
                    res = transcribe_long_file_cel.apply_async((speech_file,source_code,filename,obj.id,length,user.id,hertz),queue='high-priority')
                    debit_status, status_code = UpdateTaskCreditStatus.update_credits(account_debit_user, consumable_credits)
                    return Response({'msg':'Transcription is ongoing. Pls Wait','celery_id':res.id},status=400)
                elif state == 'SUCCESS':
                    ser = TaskTranscriptDetailSerializer(queryset,many=True)
                    return Response(ser.data)
           
            return JsonResponse(res,safe=False,json_dumps_params={'ensure_ascii':False})
        else:
            return Response({'msg':'Insufficient Credits'},status=400)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def transcribe_file_get(request):
    task_id = request.GET.get('task')
    task=Task.objects.get(id=task_id)
    authorize(request,resource=task,action="read",actor=request.user)
    queryset = TaskTranscriptDetails.objects.filter(task_id = task_id)
    ser = TaskTranscriptDetailSerializer(queryset,many=True)
    return Response(ser.data)

def google_long_text_file_process(file,obj,language,gender,voice_name):
    '''
    This is to convert long text to file of translated task.
    '''
    final_name,ext =  os.path.splitext(file)
    size_limit = 4500 
    final_audio = final_name  + "_" + obj.job.source_language_code + "-" + obj.job.target_language_code  + ".mp3"
    dir_1 = os.path.join('/ai_home/',"output_"+str(obj.id))
    if not os.path.exists(dir_1):
        os.mkdir(dir_1)
    lang=language if language else obj.job.target_language_code 
    split_file_by_size(file, dir_1, lang, size_limit)
    for file in os.listdir(dir_1):
        filepath = os.path.join(dir_1, file)
        if file.endswith('.txt'):
            name,ext = os.path.splitext(file)
            dir = os.path.join('/ai_home/',"outputAudio_"+str(obj.id))
            if not os.path.exists(dir):
                os.mkdir(dir)
            audio_ = name + '.mp3'
            audiofile = os.path.join(dir,audio_)
            text_to_speech_long(filepath,language if language else obj.job.target_language_code ,audiofile,gender if gender else 'FEMALE',voice_name)
    
    list_of_audio_files = [AudioSegment.from_mp3(mp3_file) for mp3_file in sorted(glob(os.path.join(dir, '*.mp3')),key=lambda x: int(os.path.splitext(os.path.basename(x))[0].split('_')[-1])) if len(mp3_file)!=0]
    combined = AudioSegment.empty()
    for aud in list_of_audio_files:
        combined += aud
    combined.export(final_audio, format="mp3")
    f2 = open(final_audio, 'rb')
    file_obj = DJFile(f2,name=os.path.basename(final_audio))
    shutil.rmtree(dir)
    shutil.rmtree(dir_1)
    os.remove(final_audio)
    #os.remove(out_filename)
    return file_obj,f2


def long_text_source_process(consumable_credits,user,file_path,task,language,voice_gender,voice_name):

    res1,f2 = google_long_text_source_file_process(file_path,task,language,voice_gender,voice_name)
    ser = TaskTranscriptDetailSerializer(data={"source_audio_file":res1,"task":task.id,"user":user.id})
    if ser.is_valid():
        ser.save()
    
    f2.close()


def google_long_text_source_file_process(file,obj,language,gender,voice_name):

    '''
    For source only files.
    This function is to split the long data into small chunks
    and then send it to google speech to task and returns the audio_file 
    by merging multiple audio files with numbers in series.
    '''
    
    project_id  = obj.job.project.id
    final_name,ext =  os.path.splitext(file)
    lang_list = ['hi','bn','or','ne','pa']
    final_audio = final_name + "_" + obj.job.source_language_code  + ".mp3"#+ "_" + obj.ai_taskid
    dir_1 = os.path.join('/ai_home/',"Output_"+str(obj.id))
    if not os.path.exists(dir_1):
        os.mkdir(dir_1)
    count=0
    out_filename = final_name + '_out.txt'
    size_limit = 4000 #if obj.job.source_language_code in ['ta','ja'] else 3500
    lang = language if language else obj.job.source_language_code

    ## split_file_by_size is the python library used to split by size
    split_file_by_size(file, dir_1, lang, size_limit)
    for file in os.listdir(dir_1):
        filepath = os.path.join(dir_1, file)
        if file.endswith('.txt') :
            name,ext = os.path.splitext(file)
            dir = os.path.join('/ai_home/',"OutputAudio_"+str(obj.id))
            if not os.path.exists(dir):
                os.mkdir(dir)
            audio_ = name + '.mp3'
            audiofile = os.path.join(dir,audio_)
            rr = text_to_speech_long(filepath,language if language else obj.job.source_language_code ,audiofile,gender if gender else 'FEMALE',voice_name)

    # Merging audio files by sorting with AudioSegment and returns the file
    list_of_audio_files = [AudioSegment.from_mp3(mp3_file) for mp3_file in sorted(glob(os.path.join(dir, '*.mp3')),key=lambda x: int(os.path.splitext(os.path.basename(x))[0].split('_')[-1])) if len(mp3_file)!=0]
    combined = AudioSegment.empty()
    for aud in list_of_audio_files:
        combined += aud
    combined.export(final_audio, format="mp3")
    f2 = open(final_audio, 'rb')
    file_obj = DJFile(f2,name=os.path.basename(final_audio))

    # Remove temporary directories created for splitting and merging back
    shutil.rmtree(dir)
    shutil.rmtree(dir_1)
    os.remove(final_audio)
    return file_obj,f2


@api_view(["GET"])
#@permission_classes([IsAuthenticated])
def convert_and_download_text_to_speech_source(request):
    '''
    It is to convert text_to_speech by internally calling text_to_speech_task and called download_file()
    '''    
    tasks =[]
    user = request.user
    project = request.GET.get('project',None)
    language = request.GET.get('language_locale',None)
    gender = request.GET.get('gender',None)
    pr = Project.objects.get(id=project)
    authorize(request,resource=pr,action="read",actor=request.user)
    for _task in pr.get_source_only_tasks:
        if _task.task_transcript_details.first() == None:
            tasks.append(_task)
    for obj in tasks:
        text_to_speech_task(obj,language,gender,user,None)
    shutil.make_archive(pr.project_name, 'zip', pr.project_dir_path + '/source/Audio')
    res = download_file(pr.project_name+'.zip')
    os.remove(pr.project_name+'.zip')
    return res


def text_to_speech_task(obj,language,gender,user,voice_name):
    '''
    it will call text_to_speech_celery_task and returns TaskTranscriptDetailSerializer data
    '''
    
    from ai_workspace.models import MTonlytaskCeleryStatus
    project = obj.job.project
    ai_user = project.ai_user
    
    # Take the filepath and get the extention.
    
    file,ext = os.path.splitext(obj.file.file.path)
    dir,name_ = os.path.split(os.path.abspath(file))

    # If ext is docx, it will convert it to txt and read the content and store it in data variable

    if ext == '.docx':
        name = file + '.txt'
        data = docx2txt.process(obj.file.file.path)
        with open(name, "w") as out:
            out.write(data)
    else:
        name = obj.file.file.path
        text_file = open(name, "r")
        data = text_file.read()
        text_file.close()

    # calculate the word_count of data by calling spring API, it returns word count.
    # update word_count and char_count in TaskDetails database
    seg_data = {"segment_source":data, "source_language":obj.job.source_language_code, "target_language":obj.job.source_language_code,\
                 "processor_name":"plain-text-processor", "extension":".txt"}
    res1 = requests.post(url=f"http://{spring_host}:8080/segment/word_count", data={"segmentWordCountdata":json.dumps(seg_data)})
    wc = res1.json() if res1.status_code == 200 else None
    TaskDetails.objects.get_or_create(task = obj,project = obj.job.project,defaults = {"task_word_count": wc,"task_char_count":len(data)})

    # checking for available credits, if credits available initiate the celery task else return insufficient_credits
    account_debit_user = project.team.owner if project.team else project.ai_user
    consumable_credits = get_consumable_credits_for_text_to_speech(len(data))
    initial_credit = account_debit_user.credit_balance.get("total_left")
    if initial_credit > consumable_credits:
        record_api_usage.apply_async(("GCP","Text to Speech",ai_user.uid,ai_user.email,consumable_credits), queue='low-priority')
        
        #checking for data len and decide either it is a text_to_speech_long or text_to_speech_short
        if len(data.encode("utf8"))>4500:

            # if length exceeds 4500(google's limit is 5000) then initiate text_to_speech_long celery task and returns celery status
            ins = MTonlytaskCeleryStatus.objects.filter(Q(task_id=obj.id) & Q(task_name='text_to_speech_long_celery')).last()
            state = text_to_speech_long_celery.AsyncResult(ins.celery_task_id).state if ins else None
            if state == 'PENDING' or state == 'STARTED':
                return Response({'msg':'Text to Speech conversion ongoing. Please wait','celery_id':ins.celery_task_id},status=400)
            elif (obj.task_transcript_details.exists()==False) or (not ins) or state == "FAILURE" or state == 'REVOKED':
                if state == "FAILURE":
                    user_credit = UserCredits.objects.get(Q(user=user) & Q(credit_pack_type__icontains="Subscription") & Q(ended_at=None))
                    user_credit.credits_left = user_credit.credits_left + consumable_credits
                    user_credit.save()
                celery_task = text_to_speech_long_celery.apply_async((consumable_credits,account_debit_user.id,name,obj.id,language,gender,voice_name),queue='high-priority' )
                debit_status, status_code = UpdateTaskCreditStatus.update_credits(user, consumable_credits)
                return Response({'msg':'Text to Speech conversion ongoing. Please wait','celery_id':celery_task.id},status=400)
        else:
            audio_file = name_ + "_source" + "_" + obj.job.source_language_code + ".mp3"#+ "_" + obj.ai_taskid
            
            # calling text_to_speech of google with required args and returns audio file object
            res2,f2 = text_to_speech(name,language if language else obj.job.source_language_code ,audio_file,gender if gender else 'FEMALE',voice_name)
            
            # Debit the used credits from main account
            debit_status, status_code = UpdateTaskCreditStatus.update_credits(account_debit_user, consumable_credits)
            os.remove(audio_file)

            # Returns serializer data
            ser = TaskTranscriptDetailSerializer(data={"source_audio_file":res2,"task":obj.id,"user":user.id})
            if ext == '.docx':
                os.remove(name)
            if ser.is_valid():
                ser.save()
                f2.close()
                return Response(ser.data)
            f2.close()
            #os.remove(name)
            return Response(ser.errors)
    else:
        return Response({'msg':'Insufficient Credits'},status=400)


###########Not used for now#################
# @api_view(["GET"])
# @permission_classes([IsAuthenticated])
# def get_media_link(request,task_id):
#     obj = Task.objects.get(id = task_id)
#     try:
#         task_transcript_obj = TaskTranscriptDetails.objects.filter(task = obj).first()
#         return Response({'url':task_transcript_obj.source_audio_file.url})
#     except:
#         return Response({'msg':'something went wrong'})


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def convert_text_to_speech_source(request):
    '''
    This function is to call text_to_speech_task of source_only_tasks. 
    it will get input either as task or project
    if task 
        it will call the celery task and return its status, if it is in PENDING or STARTED.
        if status is in SUCCESS, then it returns TaskTranscriptDetailSerializer data
    if project
        it will get source_only_tasks within project and repeat above and returns 
        TaskTranscriptDetailSerializer data
    '''
    task = request.GET.get('task')
    project  = request.GET.get('project')
    language = request.GET.get('language_locale',None)
    gender = request.GET.get('gender')
    voice_name = request.GET.get('voice_name')
    user = request.user
    if task:
        obj = Task.objects.get(id = task)
        # authorize(request,resource=obj,action="read",actor=request.user)
        tt = text_to_speech_task(obj,language,gender,user,voice_name)
        if tt!=None and tt.status_code == 400:
            return tt
        else:
            ser = TaskTranscriptDetailSerializer(obj.task_transcript_details.first())
            return Response(ser.data)
    if project:
        tasks =[]
        task_list = []
        pr = Project.objects.get(id=project)
        for _task in pr.get_source_only_tasks:
            if _task.task_transcript_details.first() == None:
                tasks.append(_task)
        # tasks=filter_authorize(request,tasks,"read",request.user)
        if tasks:
            for obj in tasks:
                conversion = text_to_speech_task(obj,language,gender,user,voice_name)
                if conversion.status_code == 200:
                    task_list.append(obj.id)
                elif conversion.status_code == 400:
                    return conversion
        queryset = TaskTranscriptDetails.objects.filter(task__in = pr.get_source_only_tasks)
        ser = TaskTranscriptDetailSerializer(queryset,many=True)
        return Response(ser.data)
    else:
        return Response({'msg':'task_id or project_id must'})


@api_view(["GET"])
#@permission_classes([IsAuthenticated])
def download_text_to_speech_source(request):
    '''
    This function is to download Text_to_speech(with only source_lang) task.
    '''
    task = request.GET.get('task')
    language = request.GET.get('language_locale',None)
    gender = request.GET.get('gender')
    user = request.user
    obj = Task.objects.get(id = task)
    authorize(request,resource=obj,action="download",actor=request.user)
    file = obj.task_transcript_details.first().source_audio_file
    return download_file(file.path)



#######################Not using now ###############################################
# @api_view(["GET"])
# #@permission_classes([IsAuthenticated])
# def download_speech_to_text_source(request):
#     task = request.GET.get('task')
#     obj = Task.objects.get(id = task)
#     authorize(request,resource=obj,action="download",actor=request.user)
#     try:
#         output_from_writer =  obj.task_transcript_details.first().transcripted_file_writer
#         return download_file(output_from_writer.path)
#     except BaseException as e:
#         print(f"Error : {str(e)}")
#         return Response({'msg':'something went wrong'})



@api_view(["GET"])
#@permission_classes([IsAuthenticated])
def download_task_target_file(request):
    '''
    This function is to download "google file translate" flow task
    '''
    task = request.GET.get('task')
    obj = Task.objects.get(id = task)
    authorize(request,resource=obj,action="download",actor=request.user)
    try:
        output_file =  obj.task_file_detail.first().target_file
        return download_file(output_file.path)
    except BaseException as e:
        print(f"Error : {str(e)}")
        return Response({'msg':'something went wrong'})


def zipit(folders, zip_filename):
    # used in project_download
    zip_file = zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED)

    for folder in folders:
        for dirpath, dirnames, filenames in os.walk(folder):
            for filename in filenames:
                zip_file.write(
                    os.path.join(dirpath, filename),
                    os.path.relpath(os.path.join(dirpath, filename), os.path.join(folders[0], '../..')))

    zip_file.close()

@api_view(['GET',])
@permission_classes([IsAuthenticated])
def project_list_download(request):
    '''
    This function is to download list of projects as one zip file.
    '''
    projects = request.GET.getlist('project')

    pro = get_list_or_404(Project,id__in=projects)
    proj=authorize_list(pro,"read",request.user)
    projects=[str(i.id) for i in proj]
    
    if projects:
        dest = []
        for obj in projects:
            pr = Project.objects.get(id=obj)
            path = pr.project_dir_path + '/source'
            dest.append(path)
        zip_file_name = "Project-"+projects[0]+"-"+projects[-1]+'.zip' if len(projects) > 1 else "Project-"+projects[0]+'.zip'
        zipit(dest,zip_file_name)
        tt = download_file(zip_file_name)
        os.remove(zip_file_name)
        return tt
    return JsonResponse({"msg":"you are not authorized to perform this action"})


############################Deprecated###########################################
@permission_classes([IsAuthenticated])
def task_unassign(request):
    task = request.GET.getlist('task')
    assigns = TaskAssignInfo.objects.filter(Q(task_id__in=task))
    for obj in assigns:
        user = obj.task.job.project.ai_user
        team_members = [i.internal_member for i in user.team.internal_member_team_info.filter(role=1)] if user.team else []
        if request.user == user or request.user in team_members:
            segment_count=0 if obj.task.document == None else obj.task.get_progress.get('confirmed_segments')
            task_history = TaskAssignHistory.objects.create(task =obj.task,previous_assign_id=obj.task.assign_to_id,task_segment_confirmed=segment_count,unassigned_by=request.user)
            obj.task.assign_to = user
            obj.task.save()
            obj.delete()
        else:
            return Response({'msg':'Permission Denied'})
    return Response({"msg":"Tasks Unassigned Successfully"},status=200)

###########Not using now ###############
# def docx_save(name,data):
#     from docx import Document
#     document = Document()
#     new_parser = HtmlToDocx()
#     quill_data = data.get('ops')
#     docx = html.render(quill_data)
#     new_parser.add_html_to_document(docx, document)
#     document.save(name)
#     f2 = open(name, 'rb')
#     file_obj = DJFile(f2)
#     return file_obj,name,f2


# def target_exists(project):
#     for i in project.project_jobs_set.all():
#         if i.target_language != None:
#             return True
#     return False



def update_project_from_writer(task_id):
    # used in QuickProjectSetupView update to update transcripted project
    obj = TaskTranscriptDetails.objects.filter(task_id = task_id).first()
    writer_project_updated_count = 1 if obj.writer_project_updated_count==None else obj.writer_project_updated_count+1
    obj.writer_project_updated_count = writer_project_updated_count
    obj.save()
    writer_filename = obj.writer_filename + '_edited_'+ str(obj.writer_project_updated_count)+'.docx'
    file_obj = ContentFile(obj.transcripted_file_writer.file.read(),name=writer_filename)
    return file_obj



####################### Not using now ############################
@permission_classes([IsAuthenticated])
def get_quill_data(request):
    task_id = request.GET.get('task_id')
    obj = TaskTranscriptDetails.objects.filter(task_id = task_id).first()
    try:
        data = json.loads(obj.quill_data)
        res = data.get('ops')
    except:res = None
    return Response({'data':res})


def update_task_assign(task_obj,user):
    try:
        obj = TaskAssignInfo.objects.filter(task_assign__task = task_obj).filter(task_assign__assign_to = user).first().task_assign
        if obj.status != 2:
            obj.status = 2
            obj.save()
    except:pass



@api_view(['POST',])
@permission_classes([IsAuthenticated])
def writer_save(request):
    # This is to update task_transcripted text from writer
    task_id = request.POST.get('task_id')
    transcripted_file_writer = request.FILES.get('file',None)
    task_obj = Task.objects.get(id=task_id)
    edited_text = request.POST.get('edited_text')
    obj = TaskTranscriptDetails.objects.filter(task_id = task_id).first()
    filename,ext = os.path.splitext(task_obj.file.filename)
    data1 = {"writer_filename":filename,"task":task_id,"html_data":edited_text,'user':request.user.id}
    if transcripted_file_writer:
        data1.update({"transcripted_file_writer":transcripted_file_writer})
    if obj:
        ser1 = TaskTranscriptDetailSerializer(obj,data=data1,partial=True)#"transcripted_file_writer":file_obj,
    else:
        ser1 = TaskTranscriptDetailSerializer(data=data1,partial=True)#"transcripted_file_writer":file_obj,
    if ser1.is_valid():
        ser1.save()
        # this is to update editor status to inprogress
        update_task_assign(task_obj,request.user)
        return Response(ser1.data)
    return Response(ser1.errors)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_voice_task_status(request):
    '''
    it returns pre_translation status of all tasks in voice project. 
    '''
    from ai_workspace.models import MTonlytaskCeleryStatus
    from ai_auth.tasks import transcribe_long_file_cel,google_long_text_file_process_cel
    project_id = request.GET.get('project')
    sub_category = request.GET.get('sub_category')
    pr = Project.objects.get(id=project_id)
    res= []
    if pr.project_type_id == 4:
        tasks = pr.get_source_only_tasks
        for i in tasks:
            obj = MTonlytaskCeleryStatus.objects.filter(task=i).filter(Q(task_name = 'transcribe_long_file_cel') | Q(task_name = 'google_long_text_file_process_cel')).last()
            if obj:
                if obj.task_name == 'transcribe_long_file_cel':
                    state = transcribe_long_file_cel.AsyncResult(obj.celery_task_id).state if obj and obj.celery_task_id else None
                elif obj.task_name == 'google_long_text_file_process_cel':
                    state = google_long_text_file_process_cel.AsyncResult(obj.celery_task_id).state if obj and obj.celery_task_id else None
                if state == 'STARTED':
                    status = 'False'
                else:
                    status = 'True'
            else:
                status = 'True'
            res.append({'task':i.id,'open':status})
        return Response({'res':res})
    else:
        return Response({'msg':'No Detail'})



def celery_check(obj):
    # used in get_task_status
    from ai_auth.tasks import pre_translate_update
    state = None
    if obj.task_name == 'mt_only':
        state = mt_only.AsyncResult(obj.celery_task_id).state if obj and obj.celery_task_id else None
    elif obj.task_name == 'pre_translate_update':
        state = pre_translate_update.AsyncResult(obj.celery_task_id).state if obj and obj.celery_task_id else None
    if state == 'STARTED':
        status = 'False'
    else:
        status = 'True'
    return status




@api_view(['GET',])
@permission_classes([IsAuthenticated])
def get_task_status(request):
    '''
    it returns pre_translation status of single task or all tasks in project. 
    '''
    from ai_workspace_okapi.api_views import DocumentViewByTask
    from ai_workspace.models import MTonlytaskCeleryStatus
    from ai_tm.api_views import get_json_file_path
    project_id = request.GET.get('project')
    task_id = request.GET.get('task')
    if project_id:
        pr = Project.objects.get(id=project_id)
        pre_t = pr.pre_translate
        tasks = pr.get_mtpe_tasks
    elif task_id:
        pre_t = Task.objects.get(id=task_id).job.project.pre_translate
        tasks = Task.objects.filter(id=task_id)
    if pre_t == True:
        res = []
        for i in tasks:
            msg,progress = None,None
            document = i.document                    
            obj = MTonlytaskCeleryStatus.objects.filter(task=i).filter(Q(task_name = 'mt_only') | Q(task_name = 'pre_translate_update')).last()
            if document:
                if not obj or obj.status == 2:
                    status = 'True'
                elif obj.status == 1 and obj.error_type == "Insufficient Credits":
                    status = 'True'
                else:
                    status = celery_check(obj)
            else:
                if obj:
                    status = celery_check(obj)
                else:
                    status = 'True' 
            if status == 'True':
                progress = i.get_progress
            res.append({'task':i.id,'document':i.document_id,'open':status,'progress':progress,'msg':msg})
        return Response({'res':res})
    else:
        return Response({'msg':'No Detail'})


class ExpressProjectSetupView(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]
    '''
    This view is to create instant project.
    '''
    def create(self, request):
        punctuation='''!"#$%&'``()*+,-./:;<=>?@[\]^`{|}~_'''
        text_data=request.POST.get('text_data')
        text_data = text_data.replace('\r','')
        user = self.request.user
        user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user
        name =  text_data.split()[0].strip(punctuation)+ ".txt" if len(text_data.split()[0])<=15 else text_data[:5].strip(punctuation)+ ".txt"
        im_file= DjRestUtils.convert_content_to_inmemoryfile(filecontent = text_data.encode(),file_name=name)
        serializer =ProjectQuickSetupSerializer(data={**request.data,"files":[im_file],"project_type":['5']},context={"request": request,'user_1':user_1})
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            pr = Project.objects.get(id=serializer.data.get('id'))
            authorize(request,resource=pr,action="create",actor=self.request.user)
            source_lang = pr.project_jobs_set.first().source_language_id
            res=[{'task_id':i.id,'target_lang_name':i.job.target_language.language,"target_lang_id":i.job.target_language.id} for i in pr.get_mtpe_tasks]
            return Response({'project_id':pr.id,'source_lang_id':source_lang,'Res':res})
        return Response(serializer.errors)


def get_consumable_credits_for_text(source,target_lang,source_lang):
    '''
    This is to calculate the credits for text by calling API in spring_boot which intern returns word_count.
    '''
    seg_data = { "segment_source" : source,
                 "source_language" : source_lang,
                 "target_language" : target_lang,
                 "processor_name" : "plain-text-processor",
                 "extension":".txt"
                 }
    res = requests.post(url=f"http://{spring_host}:8080/segment/word_count", \
        data={"segmentWordCountdata":json.dumps(seg_data)})#,timeout=3)

    if res.status_code == 200:
        return res.json()
    else:
        logger.info(">>>>>>>> Error in segment word count calculation <<<<<<<<<")
        raise  ValueError("Sorry! Something went wrong with word count calculation.")

def exp_proj_save(task_id,mt_change):
    '''
    This function is to get all the segments of latest version from task_id 
    and group it with text_unit_id for maintaining paragraphs and returns it.
    '''
    vers = ExpressProjectSrcSegment.objects.filter(task_id = task_id).last().version
    exp_obj = ExpressProjectSrcSegment.objects.filter(task_id = task_id,version=vers)
    obj = Task.objects.get(id=task_id)
    express_obj = ExpressProjectDetail.objects.get(task_id = task_id)
    tar = ''
    for i in exp_obj.distinct('src_text_unit'):
        rr = exp_obj.filter(src_text_unit=i.src_text_unit)
        for i in rr:
            try:tar_1 = i.express_src_mt.filter(mt_engine_id=express_obj.mt_engine_id).first().mt_raw 
            except:tar_1 = None
            tar = tar +' '+tar_1 if tar_1 else ''
        tar = tar + '\n\n'
    express_obj.mt_raw = tar.strip().strip('\n')
    express_obj.target_text = tar.strip('\n')
    express_obj.save()
    try:wc = get_consumable_credits_for_text(express_obj.source_text,None,obj.job.source_language_code)
    except:wc = 0
    td = TaskDetails.objects.update_or_create(task = obj,project = obj.job.project,defaults = {"task_word_count": wc,"task_char_count":len(express_obj.source_text)})
    if mt_change == None:
        ExpressProjectSrcSegment.objects.filter(task_id = task_id).exclude(version = vers).delete()
    return express_obj


def seg_create(task_id,content,from_mt_edit=None):
    '''
    This is to tokenize sentences from input text and get its translation and store it in
    ExpressProjectSrcSegment.
    '''
    from ai_workspace.models import ExpressProjectSrcSegment,ExpressProjectSrcMTRaw
    obj = Task.objects.get(id=task_id)
    lang_code = obj.job.source_language_code
    user = obj.job.project.ai_user
    express_obj = ExpressProjectDetail.objects.get(task_id = task_id)
    if from_mt_edit == None:
        express_obj.source_text = content
        express_obj.mt_engine = obj.job.project.mt_engine
        express_obj.save()

    NEWLINES_RE = re.compile(r"\n{1,}")
    no_newlines = content.strip("\n")  # remove leading and trailing "\n"
    split_text = NEWLINES_RE.split(no_newlines)
    lang_list = ['hi','bn','or','ne','pa']
    lang_list_2 = ['zh-Hans','zh-Hant','ja']
    for i,j  in enumerate(split_text):
        if lang_code in lang_list_2:
            sents = cust_split(j)

        elif lang_code in lang_list:
            sents = sentence_split(j, lang_code, delim_pat='auto')
        else:
            sents = nltk.sent_tokenize(j)

        for l,k in enumerate(sents):
            ExpressProjectSrcSegment.objects.create(task_id=task_id,src_text_unit=i,src_segment=k.strip(),seq_id=l,version=1)

    for i in ExpressProjectSrcSegment.objects.filter(task_id=task_id,version=1):
        tar = get_translation(mt_engine_id=express_obj.mt_engine_id,source_string = i.src_segment ,source_lang_code=i.task.job.source_language_code , target_lang_code=i.task.job.target_language_code,user_id=user.id)
        ExpressProjectSrcMTRaw.objects.create(src_seg = i,mt_raw = tar,mt_engine_id=express_obj.mt_engine_id)
    res = exp_proj_save(task_id,None)
    return res


def get_total_consumable_credits(source_lang,prompt_string_list):
    credit = 0
    for i in prompt_string_list:
        if i != None:
            consumable_credit = get_consumable_credits_for_text(i,None,source_lang)
            credit+=consumable_credit
    return credit

def cust_split(text):
    import re
    tt = []
    for sent in re.findall(u'[^!?。\.\!\?]+[!?。\.\!\?]?', text, flags=re.U):
        tt.append(sent)
    final = [i for i in tt if i.strip()]
    return final


def seg_edit(express_obj,task_id,src_text,from_mt_edit=None):
    # called in task_segments_save
    obj = Task.objects.get(id=task_id)
    user = obj.job.project.ai_user
    NEWLINES_RE = re.compile(r"\n{1,}")
    no_newlines = src_text.strip("\n")  # remove leading and trailing "\n"
    split_text = NEWLINES_RE.split(no_newlines)
    lang_code = obj.job.source_language_code
    lang_list = ['hi','bn','or','ne','pa']
    lang_list_2 = ['zh-Hans','zh-Hant','ja']
    exp_src_obj = ExpressProjectSrcSegment.objects.filter(task_id=task_id)
    if not exp_src_obj:
        initial_credit = user.credit_balance.get("total_left")
        consumable_credits = get_consumable_credits_for_text(src_text,source_lang=obj.job.source_language_code,target_lang=obj.job.target_language_code)
        res = seg_create(task_id,src_text,from_mt_edit)
        return None
    vers = exp_src_obj.last().version
    for i,j  in enumerate(split_text):
        if lang_code in lang_list_2:
            sents = cust_split(j)
        elif lang_code in lang_list:
            sents = sentence_split(j, lang_code, delim_pat='auto')
        else:
            sents = nltk.sent_tokenize(j)
        for l,k in enumerate(sents):
            ExpressProjectSrcSegment.objects.create(task_id=task_id,src_text_unit=i,src_segment=k.strip(),seq_id=l,version=vers+1)
    latest =  ExpressProjectSrcSegment.objects.filter(task_id=task_id).last().version
    for i in ExpressProjectSrcSegment.objects.filter(task=task_id,version=latest):
        print(i.src_segment)
        tt = ExpressProjectSrcSegment.objects.filter(task=task_id,version=latest-1).filter(src_segment__iexact = i.src_segment)
        if tt:
            mt_obj = tt.first().express_src_mt.filter(mt_engine_id=express_obj.mt_engine_id).first()
            if mt_obj: 
                ExpressProjectSrcMTRaw.objects.create(src_seg = i,mt_raw = mt_obj.mt_raw,mt_engine_id=express_obj.mt_engine_id)

            else:
                consumed = get_consumable_credits_for_text(i.src_segment,None,obj.job.source_language_code)
                tar = get_translation(mt_engine_id=express_obj.mt_engine_id,source_string = i.src_segment ,source_lang_code=i.task.job.source_language_code , target_lang_code=i.task.job.target_language_code,user_id=user.id)
                ExpressProjectSrcMTRaw.objects.create(src_seg = i,mt_raw = tar,mt_engine_id=express_obj.mt_engine_id)
        else:
            consumable = get_consumable_credits_for_text(i.src_segment,None,obj.job.source_language_code)
            tar = get_translation(mt_engine_id=express_obj.mt_engine_id,source_string = i.src_segment ,source_lang_code=i.task.job.source_language_code , target_lang_code=i.task.job.target_language_code,user_id=user.id)
            tt = ExpressProjectSrcMTRaw.objects.create(src_seg = i,mt_raw = tar,mt_engine_id=express_obj.mt_engine_id)
    res = exp_proj_save(task_id,None)
    return None


@api_view(['GET',])
@permission_classes([IsAuthenticated])
def task_get_segments(request):
    '''
    This function is to get Instant project task's translation (sentence tokenization-->translation-->merging back translated as like original)
    and return
    '''
    from ai_workspace.models import ExpressProjectDetail
    user = request.user.team.owner  if request.user.team  else request.user
    task_id = request.GET.get('task_id')
    task=get_object_or_404(Task,id=task_id)
    # authorize(request,resource=task,actor=request.user,action="read")
    express_obj = ExpressProjectDetail.objects.filter(task_id=task_id).first()
    obj = Task.objects.get(id=task_id)
    if express_obj.source_text == None:
        with open(obj.file.file.path, "r") as file:
            content = file.read()
    else:content = express_obj.source_text

    if express_obj.mt_raw == None and express_obj.target_text == None:
        initial_credit = user.credit_balance.get("total_left")
        consumable_credits = get_consumable_credits_for_text(content,source_lang=obj.job.source_language_code,target_lang=obj.job.target_language_code)

        if initial_credit > consumable_credits:
            res = seg_create(task_id,content)
            express_obj = ExpressProjectDetail.objects.filter(task_id=task_id).first()
            ser = ExpressProjectDetailSerializer(express_obj)
            return Response({'Res':ser.data})

        else:
            express_obj.source_text = content
            express_obj.mt_engine = obj.job.project.mt_engine
            express_obj.save()
            ser = ExpressProjectDetailSerializer(express_obj)
            out = ser.data
            return Response({'msg':'Insufficient Credits','Res':out})

    else:
        ser = ExpressProjectDetailSerializer(express_obj)
        return Response({'Res':ser.data})



def seg_get_new_mt(task,mt_engine_id,user,express_obj):
    # called in task_segments_save
    exp_src_obj = ExpressProjectSrcSegment.objects.filter(task_id=task.id).last()
    if not exp_src_obj:
        seg_create(task.id,express_obj.source_text,True)
        return None
    else:
        latest =  ExpressProjectSrcSegment.objects.filter(task=task).last().version
        for i in ExpressProjectSrcSegment.objects.filter(task=task,version=latest):
            mt_obj = i.express_src_mt.filter(mt_engine_id=express_obj.mt_engine_id).first() 
            if not mt_obj:
                consumable_credit = get_consumable_credits_for_text(i.src_segment,None,task.job.source_language_code)
                tar = get_translation(express_obj.mt_engine.id,i.src_segment ,i.task.job.source_language_code,i.task.job.target_language_code,i.task.job.project.ai_user.id)
                ExpressProjectSrcMTRaw.objects.create(src_seg = i,mt_raw = tar,mt_engine_id=mt_engine_id)
        exp_proj_save(task.id,True)

def inst_create(obj,option):
    # called in task_segments_save
    customize = AiCustomize.objects.get(customize = option)
    created_obj = ExpressProjectAIMT.objects.create(express_id=obj.id,source=obj.source_text,customize_id=customize.id,mt_engine_id=obj.mt_engine_id)
    return created_obj


def sent_tokenize(text,lang_code):
    # called in task_segments_save
    '''
    This is to sent tokenize the text in instant translation
    '''
    lang_list = ['hi','bn','or','ne','pa']
    lang_list_2 = ['zh-Hans','zh-Hant','ja']
    NEWLINES_RE = re.compile(r"\n{1,}")
    no_newlines = text.strip("\n")  # remove leading and trailing "\n"
    split_text = NEWLINES_RE.split(no_newlines)
    out = []
    for i,j  in enumerate(split_text):
        if lang_code in lang_list_2:
            sents = cust_split(j)
        elif lang_code in lang_list:
            sents = sentence_split(j, lang_code, delim_pat='auto')
        else:
            sents = nltk.sent_tokenize(j)
        out.extend(sents)
    return out


import difflib
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def task_segments_save(request):
    '''
    This function is to update instant project's task.
    '''
    task_id = request.POST.get('task_id')
    task=get_object_or_404(Task,id=task_id)
    if not task_id:
        return Response({'msg':'task_id required'},status=400)
    from_history = request.POST.get('from_history',None)
    target_text = request.POST.get('target_text',None)
    simplified_text = request.POST.get('simplified_text')
    shortened_text = request.POST.get('shortened_text')
    rewrite_text = request.POST.get('rewrite_text')
    mt_engine_id = request.POST.get('mt_engine',None)
    source_text = request.POST.get('source_text')
    apply_all = request.POST.get('apply_all',None)
    obj = Task.objects.get(id=task_id)
    user = obj.job.project.ai_user
    express_obj = ExpressProjectDetail.objects.filter(task_id=task_id).first()
    authorize(request,resource=express_obj,actor=request.user,action="update")

    # This is to replace src_text and tar_txt from history obj
    if from_history:
        task_hist_obj = ExpressTaskHistory.objects.get(id = from_history)
        express_obj.source_text = task_hist_obj.source_text
        express_obj.target_text = task_hist_obj.target_text
        express_obj.mt_raw = ''
        express_obj.save()
        ExpressProjectSrcSegment.objects.filter(task_id = task_id).delete()

    # This is to update target_text
    elif target_text:# or target_text!=None:
        express_obj.target_text = target_text.replace('\r','')
        express_obj.save()

    # This is to update simplified text 
    elif simplified_text:
        inst_cust_obj = express_obj.express_src_text.filter(customize__customize='Simplify').last()
        if not inst_cust_obj:
            inst_cust_obj = inst_create(express_obj,'Simplify')
        inst_cust_obj.final_result = simplified_text
        inst_cust_obj.save()

    # This is to update shortened text 
    elif shortened_text:
        inst_cust_obj = express_obj.express_src_text.filter(customize__customize='Shorten').last()
        if not inst_cust_obj:
            inst_cust_obj = inst_create(express_obj,'Shorten')
        inst_cust_obj.final_result = shortened_text
        inst_cust_obj.save()

    # This is to update rewrite text 
    elif rewrite_text:
        inst_cust_obj = express_obj.express_src_text.filter(customize__customize='Rewrite').last()
        if not inst_cust_obj:
            inst_cust_obj = inst_create(express_obj,'Rewrite')
        inst_cust_obj.final_result = rewrite_text
        inst_cust_obj.save()
      
    
    #In instant translation, we seperate text into sentences by using segmentation rules(adapting different langs) 
    #   and then translate and then stored in the model ExpressProjectSrcSegment with version detail for history.
    #This is to update source_text and mt_engine_id, 
    #   which intern checks for new addition of sentences from previous stored text and translate that alone.
    
    elif ((source_text) or (source_text and mt_engine_id)):
        source_text = source_text.replace('\r','')
        if mt_engine_id:
            express_obj.mt_engine_id = mt_engine_id
            express_obj.save()
        if apply_all == 'True':tasks = obj.job.project.get_tasks
        else: tasks = Task.objects.filter(id=task_id)
        for i in tasks:
            express_obj = ExpressProjectDetail.objects.filter(task_id=i.id).first()
            exp_src_obj = ExpressProjectSrcSegment.objects.filter(task_id=task_id)
            previous_stored_source = express_obj.source_text.strip() if express_obj.source_text else ''
            text1 = sent_tokenize(previous_stored_source,i.job.source_language_code)
            text2 = sent_tokenize(source_text.strip(),i.job.source_language_code)
            output_list = [li for li in difflib.ndiff(text1,text2) if li[0] == '+']
            initial_credit = user.credit_balance.get("total_left")
            if exp_src_obj:
                consumable_credits = get_total_consumable_credits(i.job.source_language_code,output_list)
            else:
                consumable_credits = get_consumable_credits_for_text(source_text,None,i.job.source_language_code)

            if initial_credit < consumable_credits:
                return  Response({'msg':'Insufficient Credits'},status=400) 
            else:
                express_obj = i.express_task_detail.first()
                express_obj.source_text = source_text
                express_obj.save()
                if i.id == json.loads(task_id) and mt_engine_id:
                    seg_edit(express_obj,i.id,source_text,True)
                else:
                    seg_edit(express_obj,i.id,source_text)

    # This is to update mt_engine, which intern checks for already existing or else called get_translation and then returns the result.
    elif mt_engine_id:
        initial_credit = user.credit_balance.get("total_left")
        consumable_credits = get_consumable_credits_for_text(express_obj.source_text,target_lang=None,source_lang=obj.job.source_language_code)
        if initial_credit < consumable_credits:
            return  Response({'msg':'Insufficient Credits'},status=400) 
        else:
            express_obj.mt_engine_id = mt_engine_id
            express_obj.save()
            seg_get_new_mt(obj,mt_engine_id,user,express_obj)
    express_obj = ExpressProjectDetail.objects.filter(task_id=task_id).first()
    ser = ExpressProjectDetailSerializer(express_obj)
    return Response(ser.data)

@api_view(['GET'])
#@permission_classes([IsAuthenticated])
def express_task_download(request,task_id):###############permission need to be added and checked##########################
    '''
    This is for instant project download
    '''
    obj = Task.objects.get(id = task_id)
    proj_name = obj.job.project.project_name
    express_obj = ExpressProjectDetail.objects.filter(task_id=task_id).first()
    authorize(request,resource=express_obj,actor=request.user,action="download")
    file_name,ext = os.path.splitext(obj.file.filename)
    target_filename = proj_name + "_out" +  "(" + obj.job.source_language_code + "-" + obj.job.target_language_code + ")" + ext
    with open(target_filename,'w') as f:
        f.write("Source:" + "\n")
        f.write(express_obj.source_text) 
        f.write('\n')
        f.write("---------" + "\n\n")
        f.write("Standard:" + "\n")
        target = express_obj.target_text if express_obj.target_text else ''
        f.write(target)
        f.write('\n')
        f.write("---------" + "\n\n")
        rewrite_obj = express_obj.express_src_text.filter(customize__customize='Rewrite')
        if rewrite_obj.exists():
            f.write("Rewrite:" + "\n")
            f.write(rewrite_obj.last().final_result)
            f.write("\n")
            f.write("---------" + "\n\n")
        simplified_obj = express_obj.express_src_text.filter(customize__customize='Simplify')
        if simplified_obj.exists():
            f.write("Simplify:" + "\n")
            f.write(simplified_obj.last().final_result)
            f.write("\n")
            f.write("---------" + "\n\n")
        shorten_obj =express_obj.express_src_text.filter(customize__customize='Shorten')
        if shorten_obj.exists():
            f.write("Shorten:" + "\n")
            f.write(shorten_obj.last().final_result)
            f.write("\n")
            f.write("---------" + "\n\n")
    res = download_file(target_filename)
    os.remove(target_filename)
    return res

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def express_project_detail(request,project_id):
    '''
    This function is to give instant project details. 
    this is called when instant project edit action is initiated. 
    '''
    obj = Project.objects.get(id=project_id)
    jobs = obj.project_jobs_set.all()
    authorize(request,resource=obj,actor=request.user,action="read")
    jobs_data = JobSerializer(jobs, many=True)
    source_lang_id =  obj.project_jobs_set.first().source_language_id
    mt_engine_id = obj.mt_engine_id
    project_name = obj.project_name
    team = obj.get_team
    project_file = obj.project_files_set.all().first()
    with open(project_file.file.path, "r") as file:
        content = file.read()
    return JsonResponse({'jobs':jobs_data.data,'source_lang':source_lang_id,\
                        'mt_engine':mt_engine_id,'project_name':project_name,'team':team,\
                        'source_text':content})


def voice_project_progress(pr,tasks):
    '''
    This function is to calculate the progress of Voice projects.
    it checks for source_only tasks and mtpe_tasks in that project.
    Do calculation seperately for both and return the combined result.
    '''
    from ai_workspace_okapi.models import Document, Segment
    count=0
    progress = 0
    source_tasks = tasks.filter(job__target_language=None)
    if source_tasks:
        if pr.voice_proj_detail.project_type_sub_category_id==1:
            for i in source_tasks:
                if TaskTranscriptDetails.objects.filter(task_id = i).exists():
                    if TaskTranscriptDetails.objects.filter(task_id = i).last().transcripted_text !=None:
                        count+=1
        elif pr.voice_proj_detail.project_type_sub_category_id==2:
            for i in source_tasks:
                if TaskTranscriptDetails.objects.filter(task_id = i).exists():
                    if TaskTranscriptDetails.objects.filter(task_id = i).last().source_audio_file !=None:
                        count+=1
    mtpe_tasks = tasks.filter(~Q(job__target_language=None))
    if mtpe_tasks:
        assigned_jobs = [i.job.id for i in mtpe_tasks]
        docs = Document.objects.filter(job__in=assigned_jobs).all()

        if not docs:
            count+=0
        if docs.count() == mtpe_tasks.count():
            total_seg_count = 0
            confirm_count  = 0
            confirm_list = [102, 104, 106, 110, 107]

            segs = Segment.objects.filter(text_unit__document__job__project_id=pr.id)
            for seg in segs:
                if seg.is_merged == True and seg.is_merge_start is None:
                    continue
                else:
                     total_seg_count += 1

                     seg_new = seg.get_active_object()
                     if seg_new.status_id in confirm_list:
                        confirm_count += 1

            if total_seg_count == confirm_count:
                count+=mtpe_tasks.count()
            else:
                progress+=1

    if count == 0 and progress == 0:
        return "Yet to Start"
    elif count == tasks.count():
        return "Completed"
    elif count != tasks.count() or progress != 0:
        return "In Progress"

### Not using now ############
@api_view(['PUT'])
@permission_classes([IsAuthenticated])
def translate_from_pdf(request,task_id):
    from ai_exportpdf.models import Ai_PdfUpload
    from ai_exportpdf.views import get_docx_file_path
    task_obj = Task.objects.get(id = task_id)
    pdf_obj = Ai_PdfUpload.objects.filter(task_id = task_id).last()
    user = request.user
    user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user

    if pdf_obj.pdf_api_use == "convertio":
        docx_file_path = get_docx_file_path(pdf_obj.id)
        file = open(docx_file_path,'rb')
        file_obj = ContentFile(file.read(),name= os.path.basename(docx_file_path))#name=docx_file_name
    else:
        file_obj = ContentFile(pdf_obj.docx_file_from_writer.file.read(),name= os.path.basename(pdf_obj.docx_file_from_writer.path))
    ins = task_obj.job.project
    team = True if ins.team else False
    serlzr = ProjectQuickSetupSerializer(ins, data={"files":[file_obj],'team':[team]},context={"request": request,'user_1':user_1}, partial=True)
    if serlzr.is_valid():
        serlzr.save()
        return Response(serlzr.data)
    return Response(serlzr.errors)



class MyDocFilter(django_filters.FilterSet):
    doc_name = django_filters.CharFilter(field_name='doc_name',lookup_expr='icontains')#related_docs__doc_name
    class Meta:
        model = MyDocuments
        fields = ['doc_name']

from django.db.models import Value, CharField, IntegerField
from ai_openai.models import BlogCreation,BookCreation
from functools import reduce

class MyDocumentsView(viewsets.ModelViewSet):
    '''
    This view is to list, create, update and delete Documents in AIWriter.
    '''

    serializer_class = MyDocumentSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend,SearchFilter,CaseInsensitiveOrderingFilter]
    ordering_fields = ['doc_name','id']#'proj_name',
    search_fields = ['doc_name']
    filterset_class = MyDocFilter
    paginator = PageNumberPagination()
    ordering = ('-created_at')
    paginator.page_size = 20
    # https://www.django-rest-framework.org/api-guide/filtering/


    # it combines multiple models MyDocument,BlogCreation,BookCreation querysets into one queryset and return it.
    def get_queryset_new(self):
        query = self.request.GET.get('doc_name')
        ordering = self.request.GET.get('ordering')
        user = self.request.user
        project_managers = user.team.get_project_manager if user.team else []
        owner = user.team.owner if (user.team and user in project_managers) else user
        queryset = MyDocuments.objects.filter(Q(ai_user=user)|Q(ai_user__in=project_managers)|Q(ai_user=owner)).distinct()
        q1 = queryset.annotate(open_as=Value('Document', output_field=CharField())).values('id','created_at','doc_name','word_count','open_as','document_type__type')
        q1 = q1.filter(doc_name__icontains =query) if query else q1
        q2 = BlogCreation.objects.filter(Q(user = user)|Q(created_by__in = project_managers)|Q(user=owner)).distinct().filter(blog_article_create__document=None).distinct().annotate(word_count=Value(0,output_field=IntegerField()),document_type__type=Value(None,output_field=CharField()),open_as=Value('BlogWizard', output_field=CharField())).values('id','created_at','user_title','word_count','open_as','document_type__type')
        q2 = q2.filter(user_title__icontains = query) if query else q2
        q4 = BookCreation.objects.filter(Q(user = user)|Q(user=owner)).distinct().annotate(word_count=Value(0,output_field=IntegerField()),document_type__type=Value('Book',output_field=CharField()),open_as=Value('Book', output_field=CharField())).values('id','created_at','project__project_name','word_count','open_as','document_type__type')
        q4 = q4.filter(project__project_name__icontains=query) if query else q4
        q3 = q1.union(q2,q4)
        final_queryset = q3.order_by('-created_at')
        if ordering:
            field_name = ordering.lstrip('-')
            if ordering.startswith('-'):
                queryset = final_queryset.order_by(F(field_name).desc(nulls_last=True))
            else:
                queryset = final_queryset.order_by(F(field_name).asc(nulls_last=True))

            return queryset
        
        return final_queryset

    # It is to return combined list of blogs and documents    
    def get_queryset_for_combined(self):
        user = self.request.user
        project_managers = user.team.get_project_manager if user.team else []
        owner = user.team.owner if (user.team and user in project_managers) else user
        queryset = MyDocuments.objects.filter(Q(ai_user=user)|Q(ai_user__in=project_managers)|Q(ai_user=owner)).distinct()
        q1 = queryset.annotate(open_as=Value('Document', output_field=CharField())).values('id','created_at','doc_name','word_count','open_as','document_type__type')
        q2 = BlogCreation.objects.filter(Q(user = user)|Q(created_by__in = project_managers)|Q(user=owner))\
            .distinct().filter(blog_article_create__document=None).distinct().annotate(word_count=Value(0,output_field=IntegerField()),\
            document_type__type=Value(None,output_field=CharField()),open_as=Value('BlogWizard', output_field=CharField()),doc_name=F('user_title'))\
                .values('id','created_at','doc_name','word_count','open_as','document_type__type')
        q3 = list(chain(q1, q2))
        return q3

    # It is to return only documents
    def get_queryset(self):
        user = self.request.user
        project_managers = self.request.user.team.get_project_manager if self.request.user.team else []
        owner = self.request.user.team.owner if self.request.user.team  else self.request.user
        queryset = MyDocuments.objects.filter(Q(ai_user=user)|Q(ai_user__in=project_managers)|Q(ai_user=owner)).order_by('-id')
        return queryset

    def list(self, request, *args, **kwargs):
        paginate = request.GET.get('pagination',True)
        queryset = self.get_queryset_new()
        # Pagination=false is used for the purpose of list the documents in open_document inside Aiwriter.
        if paginate == 'False':
            serializer = MyDocumentSerializer(self.filter_queryset(self.get_queryset()), many=True)
            return Response(serializer.data)
        pagin_tc = self.paginator.paginate_queryset(queryset, request , view=self)
        serializer = MyDocumentSerializerNew(pagin_tc, many=True)
        response = self.get_paginated_response(serializer.data)
        return  response


    def retrieve(self, request, pk):
        queryset = self.get_queryset()
        ins = get_object_or_404(queryset, pk=pk)
        serializer = MyDocumentSerializer(ins)
        return Response(serializer.data)

    def create(self, request):
        file = request.FILES.get('file',None)
        ai_user = request.user.team.owner if request.user.team else request.user
        writer_proj = request.POST.get('project',None)
        if not writer_proj:
            writer_obj = WriterProject.objects.create(ai_user_id = ai_user.id)
            writer_proj = writer_obj.id
        ser = MyDocumentSerializer(data={**request.POST.dict(),'project':writer_proj,'file':file,'ai_user':ai_user.id,'created_by':request.user.id})
        if ser.is_valid(raise_exception=True):
            ser.save()
            return Response(ser.data, status=201)
        return Response(ser.errors)
        
    def update(self, request, pk, format=None):
        ins = MyDocuments.objects.get(id=pk)
        file = request.FILES.get('file')
        if file:
            ser = MyDocumentSerializer(ins,data={**request.POST.dict(),'file':file},partial=True)
        else:
            ser = MyDocumentSerializer(ins,data={**request.POST.dict()},partial=True)
        if ser.is_valid(raise_exception=True):
            ser.save()
            return Response(ser.data, status=200)
        return Response(ser.errors)

    def destroy(self, request, pk):
        # it is to delete the document instance and its related blog_creation and blog_articles and its files if any
        ins = MyDocuments.objects.get(id=pk)
        ins.blog_doc.all().delete()
        ins.ai_doc_blog.all().delete()
        if ins.file:
            os.remove(ins.file.path)
        ins.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
        
from django.db.models import Subquery
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def default_proj_detail(request):
    '''
    This is to get user's recent source and target languages and mt_engine
    '''
    last_pr = Project.objects.filter(Q(ai_user=request.user)|Q(created_by=request.user)).last()
    if last_pr:
        query =  Project.objects.filter(Q(ai_user=request.user)|Q(created_by=request.user)).exclude(project_jobs_set__target_language=None).exclude(project_type_id=3).order_by('-id').annotate(target_count = Count('project_jobs_set__target_language')).filter(target_count__gte = 1)[:20]
        out = []
        for i in query:
            res={'src':i.project_jobs_set.first().source_language.id}
            # res['tar']=[j.target_language.id for j in i.project_jobs_set.all()]
            res['tar']=[j.target_language.id for j in i.project_jobs_set.all() if j.target_language.id != i.project_jobs_set.first().source_language.id ]
            
            if res not in out:
                if res['tar']:
                    out.append(res)


        mt_engine =last_pr.mt_engine_id
        out_1 = [a[0] for a in itertools.groupby(out)][:5]
        return JsonResponse({'recent_pairs':out_1,'mt_engine_id':mt_engine})
    else:
        return JsonResponse({'recent_pairs':[],'mt_engine_id':None})


def express_custom(request,exp_obj,option):
    '''
    This function is to call customization function (instant_customize_response) in ai_openai
    and then stores the result in ExpressProjectAIMT. 
    it will update, if the instance already exists in ExpressProjectAIMT else it will create new_instance.
    returns ExpressProjectAIMTSerializer data
    '''
    from ai_openai.serializers import AiPromptSerializer
    from ai_openai.api_views import instant_customize_response
    user = exp_obj.task.job.project.ai_user
    instant_text = exp_obj.source_text
    tone=1
    txt_generated = None
    if not instant_text:
        with open(exp_obj.task.file.file.path, "r") as file:
            instant_text = file.read()
    target_lang_code = exp_obj.task.job.target_language_code
    source_lang_code = exp_obj.task.job.source_language_code
    customize = AiCustomize.objects.get(customize = option)
    total_tokens = 0
    if source_lang_code != 'en':
        initial_credit = user.credit_balance.get("total_left")
        consumable_credits_user_text =  get_consumable_credits_for_text(instant_text,source_lang=source_lang_code,target_lang='en')
        if initial_credit > consumable_credits_user_text:
            if target_lang_code!='en':
                user_insta_text_mt_en = get_translation(mt_engine_id=exp_obj.mt_engine_id , source_string = instant_text,
                                source_lang_code=source_lang_code , target_lang_code='en',user_id=user.id,from_open_ai=True)
                
                total_tokens += get_consumable_credits_for_text(user_insta_text_mt_en,source_lang=target_lang_code,target_lang='en')
            else:
                user_insta_text_mt_en = exp_obj.target_text
            result_txt,total_tokens = instant_customize_response(customize,user_insta_text_mt_en.replace('\r',''),total_tokens)
           
            if target_lang_code != 'en':
                txt_generated = get_translation(mt_engine_id=exp_obj.mt_engine_id , source_string = result_txt.strip(),
                            source_lang_code='en' , target_lang_code=target_lang_code,user_id=user.id,from_open_ai=True)
                total_tokens += get_consumable_credits_for_text(result_txt,source_lang='en',target_lang=target_lang_code)

        else:
            return ({'msg':'Insufficient Credits'})
    
    else:##english
        initial_credit = user.credit_balance.get("total_left")
        consumable_credits_user_text =  get_consumable_credits_for_text(instant_text,source_lang=source_lang_code,target_lang=target_lang_code)
        if initial_credit < consumable_credits_user_text:
            return ({'msg':'Insufficient Credits'})
        result_txt,total_tokens = instant_customize_response(customize,instant_text.replace('\r',''),total_tokens)
       
        if target_lang_code != 'en':
            txt_generated = get_translation(mt_engine_id=exp_obj.mt_engine_id , source_string = result_txt.strip(),
                        source_lang_code='en' , target_lang_code=target_lang_code,user_id=user.id,from_open_ai=True)
            total_tokens += get_consumable_credits_for_text(result_txt,source_lang='en',target_lang=target_lang_code)
    AiPromptSerializer().customize_token_deduction(instance = None,total_tokens= total_tokens,user=user)
   
    inst_data = {'express':exp_obj.id,'source':instant_text, 'customize':customize.id,
                'api_result':result_txt.strip() if result_txt else None,'mt_engine':exp_obj.mt_engine_id,'final_result':txt_generated if txt_generated else result_txt.strip()}
    
    ins = ExpressProjectAIMT.objects.filter(express=exp_obj,customize=customize)
    if ins:
        queryset = ins.last()
        serializer = ExpressProjectAIMTSerializer(queryset,data=inst_data,partial=True)
    else:
        serializer = ExpressProjectAIMTSerializer(data=inst_data)
    if serializer.is_valid():
        serializer.save()
        return (serializer.data)
    return (serializer.errors)



@api_view(['POST',])
@permission_classes([IsAuthenticated])
def instant_translation_custom(request):
    '''
    This function is to use customization options in instant translation.
    It will check for changes in incoming source_text with already existing source_text with difflib library.
    if change detected, then it calls for the customize option again in the function express_custom(), 
    else it will return the already stored one from the model ExpressProjectAIMT.
    '''
    from ai_openai.serializers import AiPromptSerializer
    from ai_openai.api_views import customize_response
    task = request.POST.get('task')
    output_list = []
    option = request.POST.get('option')#Shorten#Simplify#Rewrite
    customize = AiCustomize.objects.get(customize = option)
    exp_obj = ExpressProjectDetail.objects.get(task_id = task)
    user = exp_obj.task.job.project.ai_user
    ins = ExpressProjectAIMT.objects.filter(express=exp_obj,customize=customize)
    if ins:
        queryset = ins.last()
        text1 = exp_obj.source_text.strip()
        text2 = queryset.source.strip()
        output_list = [li for li in difflib.ndiff(text1.replace('\n',''), text2.replace('\n','')) if li[0]=='+' or li[0]=='-' if li[-1].strip()]
        if output_list == []:
            if exp_obj.mt_engine_id == queryset.mt_engine_id:
                serializer = ExpressProjectAIMTSerializer(queryset)
                return Response(serializer.data)
            elif queryset.api_result:
                input_src = queryset.api_result
                initial_credit = user.credit_balance.get("total_left")
                consumable_credit = get_consumable_credits_for_text(input_src,exp_obj.task.job.target_language_code,exp_obj.task.job.source_language_code)
                if initial_credit > consumable_credit:
                    txt_generated = get_translation(mt_engine_id=exp_obj.mt_engine_id , source_string = input_src,
                                    source_lang_code=exp_obj.task.job.source_language_code , target_lang_code=exp_obj.task.job.target_language_code,user_id=exp_obj.task.job.project.ai_user_id)
                    serializer = ExpressProjectAIMTSerializer(queryset,data={'final_result':txt_generated,'mt_engine':exp_obj.mt_engine.id},partial=True)
                    if serializer.is_valid():
                        serializer.save()
                        return Response(serializer.data)
                    return Response(serializer.errors)
                else:
                    return Response({'msg':'Insufficient Credits'},status=400)
            else:
                res = express_custom(request,exp_obj,option)
                if res.get('msg'):return Response(res,status=400)
                else:return Response(res)
        else:
            res = express_custom(request,exp_obj,option)
            if res.get('msg'):return Response(res,status=400)
            else:return Response(res)
            
    elif not ins:
        res = express_custom(request,exp_obj,option)
        if res.get('msg'):return Response(res,status=400)
        else:return Response(res)
        


class DocumentImageView(viewsets.ViewSet):
    permission_classes = [IsAuthenticated]

    '''
    This view is to list, create and delete document related images in AIWriter.
    '''
    def list(self,request):
        user = request.user
        image = DocumentImages.objects.filter(ai_user_id=user.id).all()
        serializer = DocumentImagesSerializer(image, many=True)
        return Response(serializer.data)

    def create(self, request):
        doc = request.POST.get('document')
        image = request.FILES.get('image')
        serializer = DocumentImagesSerializer(data={**request.POST.dict(),'image':image,'ai_user':request.user.id})
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update(self,request,pk):
        pass
        # doc = MyDocuments.objects.get(id=pk)
        # image = request.FILES.get('image')
        # if image:
        #     serializer = DocumentImagesSerializer(doc,data={**request.POST.dict(),'image':image},partial=True)
        # else:
        #     serializer= DocumentImagesSerializer(doc,data={**request.POST.dict()},partial=True)
        # if serializer.is_valid():
        #     serializer.save()
        #     return Response(serializer.data)
        # else:
        #     return Response(serializer.errors)


    def delete(self,request):
        image_url = request.GET.get('image_url')
        doc = request.GET.get('document')
        pdf = request.GET.get('pdf')
        task = request.GET.get('task')
        book = request.GET.get('book')
        if doc:queryset = DocumentImages.objects.filter(document_id=doc).all()
        if pdf:queryset = DocumentImages.objects.filter(pdf_id=pdf).all()
        if task:queryset = DocumentImages.objects.filter(task_id=task).all()
        if book:queryset = DocumentImages.objects.filter(book_id=book).all()
        for i in queryset:
            if i.image.url == image_url:
                i.delete()
            else:
                print("No match")
        return Response(status=status.HTTP_204_NO_CONTENT)





class ExpressTaskHistoryView(viewsets.ViewSet):

    '''
    This View is for instant task translation history list, create and delete.
    '''

    permission_classes = [IsAuthenticated]

    # def get_queryset(self):
    #     queryset=ExpressTaskHistory.objects.filter(id=self.id).all()
    #     return queryset

    def list(self,request):
        task_id = request.GET.get('task')
        task=get_object_or_404(Task,id=task_id)
        authorize(request,resource=task,action="read",actor=self.request.user)
        queryset = ExpressTaskHistory.objects.filter(task_id=task_id).exclude(target_text=None).all().order_by('-id')
        serializer = ExpressTaskHistorySerializer(queryset,many=True)
        return Response(serializer.data)

    def create(self,request):
        source = request.POST.get('source_text')
        target = request.POST.get('target_text')
        task = request.POST.get('task')
        action = request.POST.get('action')
        obj=get_object_or_404(Task,id=task)
        authorize(request,resource=obj,action="create",actor=self.request.user)
        serializer = ExpressTaskHistorySerializer(data={'source_text':source.replace('\r',''),'target_text':target.replace('\r',''),'action':action,'task':task})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update(self,request,pk):
        pass

    def delete(self,request,pk):
        obj = ExpressTaskHistory.objects.get(id=pk)
        task=obj.task
        authorize(request,resource=task,action="delete",actor=self.request.user)
        obj.delete()
        return Response(status=204)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def docx_convertor(request):
    from docx import Document
    from ai_workspace.html2docx_custom import HtmlToDocx
    import re
    html = request.POST.get('html')
    name = request.POST.get('name')
    document = Document()
    new_parser = HtmlToDocx()
    new_parser.table_style = 'TableGrid'
    target_filename = name + '.docx'

    def replace_hex_color(match):
        hex_color = match.group(1)
        red, green, blue = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        rgb_color = f"rgb({red}, {green}, {blue})"
        return rgb_color

    def replace_hex_colors_with_rgb(html):
        hex_color_regex = re.compile("'#([A-Fa-f0-9]{6}|[A-Fa-f0-9]{3})'")
        html = hex_color_regex.sub(replace_hex_color, html)
        return html

    updatedHtml = replace_hex_colors_with_rgb(html)  
    htmlupdates = updatedHtml.replace('<br />', '')
    #new_parser.add_html_to_document(htmlupdates, document)
    try:new_parser.add_html_to_document(htmlupdates, document)
    except:return Response({'msg':"Unsupported formatting"}, status=400)
    document.save(target_filename)
    res = download_file(target_filename)
    os.remove(target_filename)
    return res

##########test doc to docx convert

import os
import mimetypes
import urllib.parse
from django.http import HttpResponse

def download_file_doc(file_path):
    filename = os.path.basename(file_path)
    fl = open(file_path, 'rb')
    
    # Guess MIME type
    mime_type, _ = mimetypes.guess_type(file_path)
    if not mime_type:
        # Default to DOCX MIME type if guessing fails
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    response = HttpResponse(fl, content_type=mime_type)
    
    # Encode filename
    encoded_filename = urllib.parse.quote(filename, encoding='utf-8')
    
    # Set Content-Disposition header
    response['Content-Disposition'] = 'attachment;filename*=UTF-8\'\'{}'.format(encoded_filename)
    
    # Set custom header
    response['X-Suggested-Filename'] = encoded_filename
    
    # Expose Content-Disposition header
    response['Access-Control-Expose-Headers'] = 'Content-Disposition'
    
    return response




import subprocess
from django.core.files.storage import default_storage

 


@api_view(["POST"])
def doc2docx(request):
    file = request.FILES.get('file',None)
    if not file:
        return Response({'msg':'Need file to upload'})
    temp_doc_path = default_storage.save('temp/' + file.name, file)
    temp_docx_path_full = os.path.join(settings.MEDIA_ROOT, temp_doc_path)
    path_data = subprocess.run(['libreoffice', 
                        '--headless',
                        '--convert-to', 'docx',
                        "output.docx"], check=True)

    print("path_data", path_data)
      
 
    res = download_file_doc(temp_docx_path_full) #download_file_doc
    # os.remove(temp_docx_path_full)
    return res


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def project_word_char_count(request):
    '''
    It will get the list of project_ids.
    This function is to check project analysis property of each project. 
    if it is analysed, it will return char_count, word_count and seg_count.
    else it will initiate celery task of project_analysis_property
    '''
    from .api_views import ProjectAnalysisProperty
    from .models import MTonlytaskCeleryStatus
    prs = request.GET.getlist('project_id')
    final =[]
    for pr in prs:
        pr_obj = Project.objects.get(id=pr)
        obj = MTonlytaskCeleryStatus.objects.filter(project_id = pr).filter(task_name = 'project_analysis_property').last()
        state = project_analysis_property.AsyncResult(obj.celery_task_id).state if obj else None
        if state == 'STARTED' or state == 'PENDING':
            res = {"proj":pr_obj.id,'msg':'project analysis ongoing. Please wait','celery_id':obj.celery_task_id}
        elif state =='None' or state == 'FAILURE' or state == 'REVOKED':
            celery_task = project_analysis_property.apply_async((pr_obj.id,), queue='high-priority')
            res = {"proj":pr_obj.id,'msg':'project analysis ongoing. Please wait','celery_id':celery_task.id}
        elif state == "SUCCESS" or pr_obj.is_proj_analysed == True:
            task_words = []
            tasks = pr_obj.get_analysis_tasks
            if pr_obj.is_all_doc_opened:

                [task_words.append({i.id:i.document.total_word_count}) for i in tasks]
                out=Document.objects.filter(id__in=[j.document_id for j in tasks]).aggregate(Sum('total_word_count'),\
                    Sum('total_char_count'),Sum('total_segment_count'))

                res = ({"proj": pr_obj.id, "proj_word_count": out.get('total_word_count__sum'), "proj_char_count":out.get('total_char_count__sum'), \
                    "proj_seg_count":out.get('total_segment_count__sum'),\
                                    "task_words" : task_words })
            else:
                out = TaskDetails.objects.filter(task_id__in=[j.id for j in tasks]).aggregate(Sum('task_word_count'),Sum('task_char_count'),Sum('task_seg_count'))
                task_words = []
                [task_words.append({i.id:i.task_details.first().task_word_count if i.task_details.first() else 0}) for i in tasks]

                res = ({"proj":pr_obj.id, "proj_word_count": out.get('task_word_count__sum'), "proj_char_count":out.get('task_char_count__sum'), \
                    "proj_seg_count":out.get('task_seg_count__sum'),
                                "task_words":task_words})
        else:
            try:
                celery_task = project_analysis_property.apply_async((pr_obj_id,), queue='high-priority')
                res = {"proj":pr_obj.id,'msg':'project analysis ongoing. Please wait','celery_id':celery_task.id}

            except:
                res = ({"proj_word_count": 0, "proj_char_count": 0, \
                    "proj_seg_count": 0, "task_words":[]})
        final.append(res)
    return Response({'out':final})


from celery.result import AsyncResult
from django.http import HttpResponse
from celery import Celery

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def stop_task(request):
    '''
    This is to stop the celery task after it gets intiated. but it's not working in main code.
    Need to debug.
    '''
    app = Celery('ai_tms')
    task_id = request.GET.get('task_id')
    task = AsyncResult(task_id)
    if task.state == 'STARTED':
        app.control.revoke(task_id, terminate=True, signal='SIGKILL')
        return HttpResponse('Task has been stopped.') 
    elif task.state == 'PENDING':
        app.control.revoke(task_id)
        return HttpResponse('Task has been revoked.')
    else:
        return HttpResponse('Task is already running or has completed.')



from django.template.loader import render_to_string
from django.core.mail import send_mail
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def msg_to_extend_deadline(request):
    '''
    This function is to extend the deadline. 
    if deadline exceeded, then user can't accept the PO. 
    so he will send message and email to extend deadline to project_manager and admin.
    '''
    from ai_marketplace.serializers import ThreadSerializer
    from ai_marketplace.models import ChatMessage
    task = request.POST.get('task')
    step = request.POST.get('step')
    reassigned = request.POST.get('reassigned')
    task_assign = TaskAssign.objects.get(task=task,step=step,reassigned=reassigned)
    
    sender = task_assign.assign_to
    receivers = []
    receiver =  task_assign.task_assign_info.assigned_by
    receivers =  receiver.team.get_project_manager if (receiver.team and receiver.team.owner.is_agency) or receiver.is_agency else []

    receivers.append(task_assign.task_assign_info.assigned_by)
    if receiver.team:
        receivers.append(task_assign.task_assign_info.assigned_by.team.owner)
    receivers = [*set(receivers)]

    for i in receivers:
        thread_ser = ThreadSerializer(data={'first_person':sender.id,'second_person':i.id})
        if thread_ser.is_valid():
            thread_ser.save()
            thread_id = thread_ser.data.get('id')
        else:
            thread_id = thread_ser.errors.get('thread_id')
      
        message = "Task with task_id "+task_assign.task.ai_taskid+" assigned to "+ task_assign.assign_to.fullname +' for '+task_assign.step.name +" in "+task_assign.task.job.project.project_name+" has requested you to extend deadline."
        
    if thread_id:
        msg = ChatMessage.objects.create(message=message,user=sender,thread_id=thread_id)
        notify.send(sender, recipient=i, verb='Message', description=message,thread_id=int(thread_id))

    context = {'message':message}	
    Receiver_emails = [i.email for i in [*set(receivers)]]	
	
    msg_html = render_to_string("assign_notify_mail.html", context)
    send_mail(
        "Regarding Assigned Task Deadline Extension",None,
        settings.DEFAULT_FROM_EMAIL,
        Receiver_emails,
        #['thenmozhivijay20@gmail.com',],
        html_message=msg_html,
    )

    task_assign.task_assign_info.deadline_extend_msg_sent = True
    task_assign.task_assign_info.save()
    return Response({"msg":"Notification sent"})   


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def translate_file(request):
    '''
    This function is to initiate the google file_translate process.
    In this function, it takes input as list of task_ids and project_id. 
    For each task, First we are checking credits(calculating credits with number of pages)
    then, we are intiating celery task which internally calls google file translate 
    and returns the boolean true or false.
    '''
    tasks = request.GET.getlist('task')
    project  = request.GET.get('project')
    user = request.user
    task_list = []
    if project or tasks:
        if project:
            pr = Project.objects.get(id=project)
            task_objs = pr.get_tasks
        if tasks:
            task_objs = Task.objects.filter(id__in=tasks)
        for obj in task_objs:
            if obj.task_file_detail.first() == None: #need to change for different mt_engines
                conversion = translate_file_task(obj.id)
                if conversion.get('status') == 200:
                    task_list.append({'task':obj.id,'msg':True,'status':200})
                elif conversion.get('status') == 400 or conversion.get('status') == 402 or conversion.get('status') == 404:
                    task_list.append({'task':obj.id,'msg':conversion.get('msg'),'status':conversion.get('status'),'celery':conversion.get('celery_id')})
            else:
                task_list.append({'task':obj.id,'msg':True,'status':200})
        return JsonResponse({"results":task_list}, safe=False)   
    else:
        return Response({'msg':'task_id or project_id must'})    

from ai_exportpdf.utils import pdf_char_check
def translate_file_process(task_id):
    '''
    This function is to call google file translate and store the translated file in the
    model TaskTranslatedFile.
    '''
    tsk = Task.objects.get(id=task_id)
    file,name = file_translate(tsk,tsk.file.get_source_file_path,tsk.job.target_language_code)
    ser = TaskTranslatedFileSerializer(data={"target_file":file,"task":tsk.id})
    if ser.is_valid():
        ser.save()
    print(ser.errors)



def translate_file_task(task_id):
    '''
    For task, First we are checking credits(calculating credits with number of pages)
    then, we are intiating celery task which internally calls google file translate and return
    celery task status.
    '''
    from .models import MTonlytaskCeleryStatus
    from ai_auth.tasks import translate_file_task_cel
    from ai_workspace_okapi.utils import get_consumption_of_file_translate

    tsk = Task.objects.get(id=task_id)
    user = tsk.job.project.ai_user
    consumable_credits = get_consumption_of_file_translate(tsk)
    if consumable_credits == None:
        return {'msg':'something went wrong in calculating page count','status':404}
    if consumable_credits == "exceeded":
        return {'msg':'PDF file page limit should be less then 300','status':404}
    if consumable_credits == "ocr":
        return {'msg':'PDF file is in image pdf should be a text or scanned file','status':404}
    initial_credit = user.credit_balance.get("total_left")
    
    if initial_credit>consumable_credits:
        ins = MTonlytaskCeleryStatus.objects.filter(Q(task_id=tsk.id) & Q(task_name='translate_file_task_cel')).last()
        
        state = translate_file_task_cel.AsyncResult(ins.celery_task_id).state if ins else None
        
        if state == 'SUCCESS':
            return {'status':200}
        elif state == 'PENDING' or state == 'STARTED':
            return ({'msg':'Translation ongoing. Please wait','celery_id':ins.celery_task_id,'task_id':tsk.id,'status':400})
        elif (tsk.task_file_detail.exists()==False) or (not ins) or state == "FAILURE" or state == 'REVOKED':
            if state == "FAILURE":
                user_credit = UserCredits.objects.get(Q(user=user) & Q(credit_pack_type__icontains="Subscription") & Q(ended_at=None))
                user_credit.credits_left = user_credit.credits_left + consumable_credits
                user_credit.save()
            celery_task = translate_file_task_cel.apply_async((tsk.id,),queue='high-priority' )
            debit_status, status_code = UpdateTaskCreditStatus.update_credits(user, consumable_credits)
            return {'msg':'Translation ongoing. Please wait','celery_id':celery_task.id,'task_id':tsk.id,'status':400}
    else:
        return {'msg':'Insufficient Credits','status':402}


from itertools import chain
from .serializers import CombinedSerializer



class CombinedProjectListView(viewsets.ModelViewSet):
    
    permission_classes = [IsAuthenticated]
    serializer_class = CombinedSerializer
    paginator = PageNumberPagination()
    paginator.page_size = 20

    '''
    This is to combine all projects, documents(general+blogs+books) and PDF
    with general name search and ordering(created_at).
    '''

    def list(self,request):
        view_instance_1 = QuickProjectSetupView()

        view_instance_1.request = request
        view_instance_1.request.GET = request.GET

        queryset1 = view_instance_1.get_queryset().exclude(project_type_id__in=[8,10])


        view_instance_2 = MyDocumentsView()

        view_instance_2.request = request
        view_instance_2.request.GET = request.GET

        queryset2 = view_instance_2.get_queryset_for_combined()
       
        user = self.request.user
        user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user
        project_managers = request.user.team.get_project_manager if request.user.team else []
        owner = request.user.team.owner if (request.user.team and request.user in project_managers) else request.user
        queryset3 = Ai_PdfUpload.objects.filter(Q(user = request.user) |Q(created_by=request.user)|Q(created_by__in=project_managers)|Q(user=owner))\
                            .filter(task_id=None).order_by('-id')         
        search_query = request.GET.get('search')

        if search_query:
            queryset1 = queryset1.filter(project_name__icontains=search_query)
            queryset2 = [item for item in queryset2 if search_query.lower() in item.get('doc_name', '').lower()]
            queryset3 = queryset3.filter(pdf_file_name__icontains=search_query)

        merged_queryset = list(chain(queryset1,queryset2,queryset3))
        ordering_param = request.GET.get('ordering', '-created_at')  

        if ordering_param.startswith('-'):
            field_name = ordering_param[1:]  
            reverse_order = True
        else:
            field_name = ordering_param
            reverse_order = False

        ordered_queryset = sorted(merged_queryset, key=lambda obj: obj[field_name] if isinstance(obj, dict) else getattr(obj, field_name), reverse=reverse_order)

        pagin_tc = self.paginator.paginate_queryset(ordered_queryset, request , view=self)
        ser = CombinedSerializer(pagin_tc, many=True, context={'request': request,'user_1':user_1})
        response = self.get_paginated_response(ser.data)
        return response

from django.db.models import F, Value, Case, When
from django.db.models.functions import Coalesce

# Annotate the queryset to get the task_word_count for each task, defaulting to 0 if task_details is None
# Convert the queryset to a dictionary


def analysed_true(pr,tasks):
    '''
    it returns total word_count, char_count and segment_count of all tasks in the project.
    '''
    task_words = []
    if pr.is_all_doc_opened:
        [task_words.append({i.id:i.document.total_word_count}) for i in tasks]
        out=Document.objects.filter(id__in=[j.document_id for j in tasks]).aggregate(Sum('total_word_count'),\
            Sum('total_char_count'),Sum('total_segment_count'))

        return {"proj_word_count": out.get('total_word_count__sum'), "proj_char_count":out.get('total_char_count__sum'), \
            "proj_seg_count":out.get('total_segment_count__sum'),\
                            "task_words" : task_words }
    else:
        out = TaskDetails.objects.filter(task_id__in=[j.id for j in tasks]).aggregate(Sum('task_word_count'),Sum('task_char_count'),Sum('task_seg_count'))
        
        ##### old 
        # task_words = {i.id: i.task_details.first().task_word_count if i.task_details.first() else 0 for i in tasks}
        # print("task_words",task_words)

        tasks_with_word_count = tasks.annotate(task_word_count=Coalesce(F('task_details__task_word_count'), Value(0))
                                                ).values('id', 'task_word_count')
        
        for task in tasks_with_word_count:
            task_words.append({task['id']: task['task_word_count']})

        # task_words = [] #{task['id']: task['task_word_count'] for task in tasks_with_word_count}
            

        ###### old 
        # for i in tasks:
            # task_words.append({i.id:i.task_details.first().task_word_count if i.task_details.first() else 0})

        ########## old
        # [task_words.append({i.id:i.task_details.first().task_word_count if i.task_details.first() else 0}) for i in tasks]

        return {"proj_word_count": out.get('task_word_count__sum'), "proj_char_count":out.get('task_char_count__sum'), \
            "proj_seg_count":out.get('task_seg_count__sum'),
                        "task_words":task_words}


from .serializers import AssertSerializer
from ai_workspace_okapi.models import ChoiceLists
class AssertList(viewsets.ModelViewSet):
    '''
    combined list view for assets (glossary and choicelist)
    '''
    serializer_class = AssertSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend,SearchFilter,CaseInsensitiveOrderingFilter]
    paginator = PageNumberPagination()
    paginator.page_size = 20

    # https://www.django-rest-framework.org/api-guide/filtering/


    def list(self,request):
        project_managers = request.user.team.get_project_manager if request.user.team else []
        user = request.user.team.owner if request.user.team and request.user in project_managers else request.user
        query = request.GET.get('type')
        ordering = request.GET.get('ordering')

        view_instance_1 = QuickProjectSetupView()

        view_instance_1.request = request
        view_instance_1.request.GET = request.GET

        queryset = view_instance_1.get_queryset()

        queryset1 = queryset.filter(glossary_project__isnull=False).exclude(project_type_id=10)
        queryset2 = ChoiceLists.objects.none()  

        user = self.request.user
        user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user

        search_query = request.GET.get('search')

        if query:
            if query == 'glossary':
                queryset2 = ChoiceLists.objects.none()
            elif query == 'assert':
                queryset1 = Project.objects.none()

        if search_query:
            queryset1 = queryset1.filter(Q(project_name__icontains=search_query)|Q(project_jobs_set__source_language__language__icontains=search_query)|Q(project_jobs_set__target_language__language__icontains=search_query))
            queryset2 = queryset2.filter(Q(name__icontains=search_query)|Q(language__language__icontains=search_query))

        merged_queryset = list(chain(queryset1,queryset2))
        ordering_param = request.GET.get('ordering', '-created_at')  

        if ordering_param.startswith('-'):
            field_name = ordering_param[1:]  
            reverse_order = True
        else:
            field_name = ordering_param
            reverse_order = False
        if field_name == 'created_at':
            ordered_queryset = sorted(merged_queryset, key=lambda obj:getattr(obj, field_name), reverse=reverse_order)
        else:
            ordered_queryset = sorted(merged_queryset,key=lambda obj: (getattr(obj, 'project_name', None) or getattr(obj,'name',None)),reverse=reverse_order)

        pagin_tc = self.paginator.paginate_queryset(ordered_queryset, request , view=self)
        ser = AssertSerializer(pagin_tc,many=True,context={'request': request,'user_1':user_1})
        response = self.get_paginated_response(ser.data)
        return response



def get_news_federal_key_and_url(lang):
    ''' 
    This function is to get the federal CMS KEY and URL based on language. 
    '''
    if lang == "Kannada":
        key = settings.KARNATAKA_FEDERAL_KEY
        integration_api_url = settings.KARNATAKA_FEDERAL_URL+"news"
    elif lang == "Telugu":
        key = settings.TELANGANA_FEDERAL_KEY
        integration_api_url = settings.TELUGANA_FEDERAL_URL+"news"
    
    elif lang == "Hindi":
        key = settings.HINDI_FEDERAL_KEY
        integration_api_url = settings.HINDI_FEDERAL_URL+"news"

    else:
        key = settings.FEDERAL_KEY
        integration_api_url = settings.FEDERAL_URL+"news"
    return key,integration_api_url

class GetNewsFederalView(generics.ListAPIView):
    pagination.PageNumberPagination.page_size = 20
    permission_classes = [IsAuthenticated,IsEnterpriseUser]
    '''
    This is to list all the news from federal CMS. 
    pagination, common search and category filter is connected
    '''

    @staticmethod
    def check_user_federal(request_user):
        user = request_user.team.owner if request_user.team else request_user
        try:
            if user.user_enterprise.subscription_name == 'Enterprise - TFN':
                return True
        except:
            return False


    def get_stories(self):
        page = self.request.query_params.get('page', 1)
        count = self.request.query_params.get('count', 20)
        news_id = self.request.query_params.get('news_id', None)
        search = self.request.query_params.get('search', None)
        lang = self.request.query_params.get('lang', None)
        categoryIds = self.request.query_params.getlist('categoryId')
        user = self.request.user.team.owner if self.request.user.team else self.request.user
        key, integration_api_url = get_news_federal_key_and_url(lang)

        headers = {
            's-id': key,
        }

        startIndex  = (int(page) - 1) * int(count)

        params = {
            'startIndex': startIndex,
            'count': count,
        }
        if news_id:
            params.update({'newsId': news_id})
        if search:
            params.update({'search': search})
        if categoryIds:
            params.update({'categoryIds': categoryIds})

        response = requests.request("GET", integration_api_url, headers=headers, params=params)
        if response.status_code == 200:
            news_jsons = response.json().get('news')
            for news_json in news_jsons:
                tasks = TaskNewsDetails.objects.filter(news_id=news_json['newsId'])
                if tasks:
                    tar_code = []
                    news_json['claimed'] = True
                    news_json['src_code'] = tasks[0].task.job.source_language_id
                    news_json['tar_code'] = list(set([task.task.job.target_language_id for task in tasks]))
            response._content = json.dumps(news_jsons).encode('utf-8')
        return response


    def list(self, request, *args, **kwargs):
        #### Need to check request from federal team ####
        allow = GetNewsFederalView.check_user_federal(request.user)
        if allow:
            response = self.get_stories()
            if response.status_code == 500:
                return Response({'msg':"something wrong with API"},status=400)
            return Response(response.json())
        else:
            return Response({"detail": "You do not have permission to perform this action."},status=403)
      

from django.core.files.base import ContentFile
from ai_workspace.utils import split_dict

class NewsProjectSetupView(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated,IsEnterpriseUser]

    '''
    This is to create project for federal by getting news_id as input.

    '''

    def get_files(self,news,lang):
        files =[]
        key,federal_api_url = get_news_federal_key_and_url(lang)
        headers = { 's-id': key,}
        for i in news:
            response = requests.request("GET", federal_api_url, headers=headers, params={'newsId':i})
            if response.status_code == 200:
                name = f"{i}.json"
                im_file = DJFile(ContentFile(json.dumps(response.json())),name=name)
                files.append(im_file)
        return files

    @staticmethod
    def create_news_detail(pr):
        tasks = pr.get_tasks
        for i in tasks:
            file_path = i.file.file.path
            with open(file_path, 'r') as fp:
                json_data = json.load(fp)
            newsID = json_data.get('news')[0].get('newsId')
            obj,created = TaskNewsDetails.objects.get_or_create(task=i,news_id=newsID,defaults = {'source_json':json_data})


        
    def create(self, request):
        '''
        It will get list of news_id and source language to fetch news from respective CMS.
        get_files() is to get the news from news_id and create json file and returns it.
        create_news_detail() is to create the record that the task is created for this news_id in lang pair to avoid duplication.
        ProjectFilesCreateType is updated to refer that file is created from CMS.

        '''
        from ai_workspace.models import ProjectFilesCreateType
        from ai_staff.models import Languages
        allow = GetNewsFederalView.check_user_federal(request.user)

        if allow:
            news = request.POST.getlist('news_id')
            lang = request.POST.get('source_language')
            source_lang = Languages.objects.get(id=lang).language
            files = self.get_files(news,source_lang)
            user = self.request.user
            user_1 = user.team.owner if user.team and user.team.owner.is_agency and (user in user.team.get_project_manager) else user
            serializer =ProjectQuickSetupSerializer(data={**request.data,"files":files,"project_type":['8']},context={"request": request,'user_1':user_1})
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                pr = Project.objects.get(id=serializer.data.get('id'))
                NewsProjectSetupView.create_news_detail(pr)
                ProjectFilesCreateType.objects.filter(project=pr).update(file_create_type=ProjectFilesCreateType.FileType.from_cms)
                authorize(request,resource=pr,action="create",actor=self.request.user)
                return Response(serializer.data)
            return Response(serializer.errors)
        else:
            return Response({"detail": "You do not have permission to perform this action."},status=403) 



from ai_workspace.serializers import TaskNewsDetailsSerializer
from ai_workspace.models import TaskNewsDetails ,TaskNewsMT

class TaskNewsDetailsViewSet(viewsets.ViewSet):
    '''
    This view is to list,create,update and delete TaskNewsDetail in Federal flow. 
    '''
    permission_classes = [IsAuthenticated,IsEnterpriseUser]

    def list(self,request):
        user = request.user
        task_news = TaskNewsDetails.objects.filter(task__file__project__ai_user=user)
        serializer = TaskNewsDetailsSerializer(task_news, many=True,context={'request':request})
        return Response(serializer.data)

    def create(self,request):
        serializer = TaskNewsDetailsSerializer(data=request.data,context={'request':request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors)

    def update(self,request,pk):
        obj = TaskNewsDetails.objects.get(id=pk )  
        serializer = TaskNewsDetailsSerializer(obj,data=request.data,context={'request':request},partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors,status=400)
    
    def retrieve(self, request, pk=None):
        obj = TaskNewsDetails.objects.get(task_id=pk )
        serializer = TaskNewsDetailsSerializer(obj,context={'request':request})
        return Response(serializer.data)

    
    def delete(self,request,pk):
        queryset = TaskNewsDetails.objects.all()
        obj = get_object_or_404(queryset, pk=pk)
        obj.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_federal_categories(request):
    '''
    This function is to get federal categories from their CMS. 
    This is for filter the news with category. to get category id and show the category for the user.
    '''
    page = request.query_params.get('page', 1)
    count = request.query_params.get('count', 20)

    headers = {
        's-id': settings.FEDERAL_KEY,
        }
    
    startIndex = (int(page) - 1) * int(count)
   
    params={ 
        'startIndex':startIndex,
        'count':count,
        }

    CMS_url = settings.FEDERAL_URL+"category"
    payload={
        'sessionId':settings.CMS_SESSION_ID,
    }
    response = requests.request("GET", CMS_url, headers=headers, params=params)
    if response.status_code == 200:
        return Response(response.json())
    else:
        return JsonResponse({'msg':'something went wrong'})
            




@api_view(["POST"])
@permission_classes([IsAuthenticated])
def push_translated_story(request):
    '''
    This function is to push back translated story to the federal CMS.
    with task_id, we identify target language and push back to that CMS.
    '''
    from ai_workspace_okapi.api_views import DocumentToFile
    task_id = request.GET.get('task_id')
    feed_id = request.GET.get('feed_id')
    task = Task.objects.get(id=task_id)
    target_lang = task.job.target_language.language

    if target_lang == "Telugu":
        federal_key = settings.TELANGANA_FEDERAL_KEY
        base_url = settings.TELUGANA_FEDERAL_URL
    elif target_lang == "Kannada":
        federal_key = settings.KARNATAKA_FEDERAL_KEY
        base_url = settings.KARNATAKA_FEDERAL_URL
    elif target_lang == "Hindi":
        federal_key = settings.HINDI_FEDERAL_KEY
        base_url = settings.HINDI_FEDERAL_URL
    else:
        federal_key = settings.FEDERAL_KEY
        base_url = settings.FEDERAL_URL

    src_json, tar_json = {}, {}
    headers = {'s-id': federal_key, 'Content-Type': 'application/json'}
    feed_url = base_url + 'createFeedV2'
    payload = {'sessionId': settings.CMS_SESSION_ID}

    tar_json = task.news_task.first().target_json

    if not tar_json:
        doc = task.document
        if doc:
            src_json = task.news_task.first().source_json.get('news')[0]
            doc_to_file = DocumentToFile()
            res = doc_to_file.document_data_to_file(request, doc.id)
            if res.status_code in [200, 201]:
                with open(res.text, "r") as fp:
                    trans_json = json.load(fp)
                tar_json = merge_dict(trans_json, src_json)
                tt = task.news_task.first()
                tt.target_json = tar_json
                tt.save()

    payload.update({
        'heading': tar_json.get('heading'),
        'description': tar_json.get('description'),
        'story': tar_json.get('story'),
        'location': tar_json.get('location'),
        'locationId': tar_json.get('locationId'),
        'categoryId': tar_json.get('maincategory'),
        'mediaIds': tar_json.get('mediaId'),
        'tags': tar_json.get('tags'),
        'keywords': tar_json.get('keywords'),
        'author': tar_json.get('authorName'),
        'custom_params': [
            {'name': 'image_caption', 'value': tar_json.get('image_caption')},
            {'name': 'story_summary', 'value': tar_json.get('story_summary')}
        ]
    })

    if feed_id:
        payload.update({'feedId': feed_id})

    response = requests.request("POST", feed_url, headers=headers, data=json.dumps(payload))
    
    if response.status_code == 200:
        feed = response.json().get('feedId')
        if feed:
            task.news_task.update(feed_id=feed, pushed=True)
            return Response({'msg': 'pushed successfully'}, status=200)
    
    return Response({'msg': "something went wrong with CMS"}, status=400)



class AddStoriesView(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated,IsEnterpriseUser]
    ''' 
    This view is to add stories for dinamalar.
    For Dinamalar, we are creating one project for one date and 
    update all the stories created in that date as the task in that project.
    In this, we get input as text (internally convert it into txt file) or files and checks for the project already exists. 
    if exists then we will add this as task or else create the project and then 
    create the task by preparing the data and send it to ProjectQuickSetupSerializer. 
    '''

    def pr_check(self,src_lang,tar_langs,user):
        today_date = date.today()
        project_name = today_date.strftime("%b %d, %Y")
        query = Project.objects.filter(ai_user=user).filter(project_type_id=8).filter(project_name__icontains=project_name)\
                .filter(project_jobs_set__source_language_id = src_lang)\
                .filter(project_jobs_set__target_language_id__in = tar_langs)
        if query:
            return query.last()
        return None

    @staticmethod
    def create_news_detail(pr):
        tasks = pr.get_tasks
        for i in tasks:
            tt = TaskNewsDetails.objects.get_or_create(task=i)

    # To check the user is dinamalar or not
    @staticmethod
    def check_user_dinamalar(request_user):
        user = request_user.team.owner if request_user.team else request_user
        try:
            if user.user_enterprise.subscription_name == 'Enterprise - DIN':
                return True
        except:
            return False 

    def get_file(self,text_data,name):
        name = f"{name}.txt"
        im_file = DjRestUtils.convert_content_to_inmemoryfile(filecontent = text_data.encode(),file_name=name)
        return im_file

    def create(self, request):
        from ai_workspace.models import ProjectFilesCreateType
        din = AddStoriesView.check_user_dinamalar(request.user)
        if din:
            text_data = request.POST.get('news_data')
            files = request.FILES.getlist('files')
            today_date = date.today()
            project_name = today_date.strftime("%b %d, %Y")
            src_lang = request.POST.get('source_language')
            tar_langs = request.POST.getlist('target_languages')
            user = self.request.user
            user_1 = user.team.owner if user.team and (user in user.team.get_project_manager) else user
            pr = self.pr_check(src_lang,tar_langs,user_1)
            count = pr.get_tasks.count() if pr else 1
            name = pr.project_name + ' - ' + str(count).zfill(3) if pr else project_name + ' - ' + str(count).zfill(3)
            if text_data:
                file = self.get_file(text_data,name)
                if files:files.append(file)
                else: files=[file]
            if pr:
                data = request.POST.dict()
                data.pop('target_languages')
                serializer =ProjectQuickSetupSerializer(pr,data={**data,"files":files,"project_type":['8']},context={"request": request,'user_1':user_1})
            else:
                serializer =ProjectQuickSetupSerializer(data={**request.data,"files":files,"project_type":['8'],"project_name":[project_name]},context={"request": request,'user_1':user_1})
            if serializer.is_valid(raise_exception=True):
                serializer.save()
                pr = Project.objects.get(id=serializer.data.get('id'))
                self.create_news_detail(pr)
                ProjectFilesCreateType.objects.filter(project=pr).update(file_create_type=ProjectFilesCreateType.FileType.from_stories)
                authorize(request,resource=pr,action="create",actor=self.request.user)
                return Response(serializer.data)
            return Response(serializer.errors)
        else:
            return Response({"detail": "You do not have permission to perform this action."},status=403) 

##########################################Dinamalar Report########################################################################################

def task_count_report(user,owner,start_date,today):
    '''
    This function is to get general report for dinamalar.
    it will get all editors from team and calculate number of tasks in each status during particular data_range
    and return the values
    '''
    managers = user.team.get_project_manager if user.team and user.team.get_project_manager else []
    team_members = user.team.get_team_members if user.team else []
    team_members.append(owner)
    res =[]
    if user in managers  or user == owner:
        tot_queryset = TaskAssign.objects.filter(task__job__project__project_type_id=8).filter(task__job__project__created_at__date__range=(start_date,today)).\
        filter(assign_to__in = team_members).distinct()
        total = tot_queryset.count()
        queryset = tot_queryset.filter(task_assign_info__isnull=False)
        editors = user.team.get_editors_only if user.team else []
        sorted_list = sorted(editors, key=lambda x: x.fullname.lower())
        for i in sorted_list:
            state = "active" if i.is_active == True else "deleted"
            additional_details = {}
            query = queryset.filter(assign_to=i)
            additional_details['user'] = i.fullname
            additional_details['state'] = state
            additional_details['TotalAssigned'] = query.count()
            additional_details['YetToStart']=query.filter(status=1).count()
            additional_details['Inprogress']=query.filter(status=2).count() #filter(task_assign_info__isnull=False).
            additional_details['Completed']=query.filter(status=3).count()
            additional_details['total_completed_words'] = query.filter(status=3).aggregate(total=Sum('task__task_details__task_word_count'))['total']
            additional_details['total_approved_words'] = query.filter(client_response=1).aggregate(total=Sum('task__task_details__task_word_count'))['total']
            res.append(additional_details)
    else:
        queryset = TaskAssign.objects.filter(task__job__project__project_type_id=8).\
                    filter(task__job__project__created_at__date__range=(start_date,today)).\
                    filter(assign_to = user).distinct()
        total = queryset.count()

    total = total
    total_assigned = queryset.count()
    progress = queryset.filter(status=2).count()
    yts = queryset.filter(status=1).count()
    completed = queryset.filter(status=3)
    total_completed_words = completed.aggregate(total=Sum('task__task_details__task_word_count'))['total']
    total_approved_words = queryset.filter(client_response=1).aggregate(total=Sum('task__task_details__task_word_count'))['total']
    data = {'Total':total,'TotalAssigned':total_assigned,'Inprogress':progress,'YetToStart':yts,'Completed':completed.count(),'TotalCompletedWords':total_completed_words,"TotalApprovedWords":total_approved_words,"Additional_info":res}
    return data,res

def billing_report(user,owner,start_date,today):
    '''
    This function is to get billing report for dinamalar.
    it will get all project_managers from team and calculate approved_words in particular data_range
    and return the values
    '''
    managers = user.team.get_project_manager if user.team and user.team.get_project_manager else []
    team_members = user.team.get_team_members if user.team else []
    team_members.append(owner)
    res =[]
    if user in managers  or user == owner:
        tot_queryset = TaskAssign.objects.filter(task__job__project__project_type_id=8).filter(Q(task_assign_status_history__field_name='client_response')&\
                    Q(task_assign_status_history__timestamp__date__range=(start_date,today))).\
                    filter(assign_to__in = team_members).distinct()
        total = tot_queryset.count()
        queryset = tot_queryset.filter(task_assign_info__isnull=False)

        editors = user.team.get_editors_only if user.team else []
        sorted_list = sorted(editors, key=lambda x: x.fullname.lower())
        for i in sorted_list:
            state = "active" if i.is_active == True else "deleted"
            additional_details = {}
            query = queryset.filter(assign_to=i)
            additional_details['user'] = i.fullname
            additional_details['total_approved_words'] = query.filter(client_response=1).aggregate(total=Sum('task__task_details__task_word_count'))['total']
            additional_details['state'] = state
            res.append(additional_details)
    else:
        queryset = TaskAssign.objects.filter(task__job__project__project_type_id=8).filter(Q(task_assign_status_history__field_name='client_response')&\
                    Q(task_assign_status_history__timestamp__date__range=(start_date,today))).\
                    filter(assign_to = user).distinct()
    total_approved_words = queryset.filter(client_response=1).aggregate(total=Sum('task__task_details__task_word_count'))['total']
    data = {"TotalApprovedWords":total_approved_words,"Additional_info":res}
    return data,res

def glossary_report(user,owner,start_date,today):
    '''
    This function is to get glossary report for dinamalar.
    it will get all terminologists from team and calculate terms added in particular data_range
    and return the value
    '''
    from ai_glex.models import MyGlossary
    managers = user.team.get_project_manager if user.team and user.team.get_project_manager else []
    team_members = user.team.get_team_members if user.team else []
    team_members.append(owner)
    res =[]
    if user in managers  or user == owner:
        queryset = MyGlossary.objects.filter(user=owner).filter(Q(created_at__date__range=(start_date,today))).distinct()
        editors = user.team.get_terminologist if user.team else []
        sorted_list = sorted(editors, key=lambda x: x.fullname.lower())
        for i in sorted_list:
            # it is to know the user is deleted or not.
            state = "active" if i.is_active == True else "deleted"
            additional_details = {}
            query = queryset.filter(created_by=i)
            additional_details['user'] = i.fullname
            additional_details['total_terms_added'] = query.count()
            additional_details['state'] = state
            res.append(additional_details)
    else:
        queryset = MyGlossary.objects.filter(created_by=user).filter(Q(created_at__date__range=(start_date,today))).distinct()
    total_terms = queryset.count()
    data = {"Totalterms":total_terms,"Additional_info":res}
    return data,res


from datetime import datetime, timedelta
@api_view(["GET"])
@permission_classes([IsAuthenticated,IsEnterpriseUser])
def get_task_count_report(request):
    '''
    This function will take input as date range and what type of report needs to be calculated.
    if it is billing then it will calculate the number of approved words of each project_owner in the team within the time range. 
    if it is glossary then it will calculate the number of terms added by terminologist within the time range.
    else it will generate the general report with number of tasks assigned, with the status yet-to-start, inprogress, completed, approved 
    '''
    user = request.user
    time_range = request.GET.get('time_range', None)
    from_date = request.GET.get('from_date',None)
    to_date = request.GET.get('to_date',None) 
    download_report = request.GET.get('download_report',False) 
    billing = request.GET.get('billing',False) 
    glossary = request.GET.get('glossary',False)
    owner = user.team.owner if user.team else user
    if owner.user_enterprise.subscription_name == 'Enterprise - DIN':
        today = datetime.now().date()
        if time_range == 'today':
            start_date = today
        elif time_range == '30days':
            start_date = today - timedelta(days=30)
        elif from_date and to_date:
            start_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            today = datetime.strptime(to_date, '%Y-%m-%d').date()
        else:
            start_date = today

        if glossary == 'True':
            data,res = glossary_report(user,owner,start_date,today)

        elif billing == 'True':
            data,res = billing_report(user,owner,start_date,today)

        else:
            data,res = task_count_report(user,owner,start_date,today)

        if download_report:
            if res:
                response = download_editors_report(res,start_date,today) #need date details. today or last month or (from_date, to_date)
                return response

        return JsonResponse(data)
    else:
        return JsonResponse({'msg':'you are not allowed to access this details'},status=400)
    
import io
import pandas as pd
from ai_workspace_okapi.api_views import DocumentToFile

# Use pandas to write the data from db to excel
def download_editors_report(res, from_date, to_date):
    '''
    This function is to download editors report.
    It uses pandas to write the data from db to excel.
    Called internally in get_task_count_report()

    :param res: Data from the database
    :param from_date: Start date for the report
    :param to_date: End date for the report
    :return: Response object for downloading the report file
    '''
    output = io.BytesIO()

    # Creating DataFrame for date details
    date_details = pd.DataFrame([{'From': from_date, 'To': to_date}])

    # Creating DataFrame from the database result
    data = pd.DataFrame(res)
    # Renaming columns for better readability
    data = data.rename(columns={'user': 'Name', 'TotalAssigned': 'No.of stories assigned',
                                'YetToStart': 'Yet to start', 'Inprogress': 'In progress',
                                'Completed': 'Completed', 'total_completed_words': 'Total words completed',
                                'total_approved_words': 'Total words approved'})

    # Filling NaN values with 0
    data.fillna(0, inplace=True)

    with pd.ExcelWriter(output, engine='xlsxwriter', date_format='YYYY-MM-DD') as writer:
        # Write the date details DataFrame to the Excel file at cell A1
        date_details.to_excel(writer, sheet_name='Report', startrow=0, index=False)

        # Filter data based on state and concatenate active and deleted entries
        data_active = data[data['state'] == "active"]
        data_deleted = data[data['state'] == "deleted"]

        df_active_sorted = data_active.sort_values(by='Name', key=lambda x: x.str.lower())
        df_deleted_sorted = data_deleted.sort_values(by='Name', key=lambda x: x.str.lower())

        empty_row = pd.DataFrame([[None] * len(df_active_sorted.columns)], columns=df_active_sorted.columns)
        df_active_sorted = pd.concat([df_active_sorted, empty_row])
        
        # df_active_sorted.loc[len(df_active_sorted)] = [None] * len(df_active_sorted.columns)
        # df_active_sorted.loc[len(df_active_sorted) + 2] = pd.Series()  # Adding empty row

        data = pd.concat([df_active_sorted, df_deleted_sorted])
        # Write the data DataFrame to the same Excel file below the date details
        data.to_excel(writer, sheet_name='Report', startrow=date_details.shape[0] + 2, index=False)

    # Closing the writer
    writer.close()

    # Seeking to the beginning of the output buffer
    output.seek(0)

    # Generating filename for the report
    filename = f"editors_report({from_date},{to_date}).xlsx"

    # Getting response object for downloading the report file
    response = DocumentToFile().get_file_response(output, pandas_dataframe=True, filename=filename)

    return response


def get_file_url(path):
    media_url = settings.MEDIA_URL.rstrip('/')
    url = path.replace(settings.MEDIA_ROOT,media_url)
    return url


@api_view(["GET"])
@permission_classes([IsAuthenticated,IsEnterpriseUser])
def get_news_detail(request):

    ''' 
    This fuction takes task_id as the input, 
    and returns source_json, target_json, source_file_path
    and target_file_path. This is for news preview. 
    '''

    from ai_workspace_okapi.api_views import DocumentToFile
    task_id = request.GET.get('task_id')
    obj = Task.objects.get(id=task_id)
    main_user = obj.job.project.ai_user
    target_json,source_json= {},{}
    target_file_path,source_file_path = None,None
    if obj.job.project.project_type_id == 8:
        if obj.news_task.exists():
            try: source_json = obj.news_task.first().source_json.get('news')[0]
            except: 
                source_json = obj.news_task.first().source_json
                source_file_path = obj.news_task.first().task.file.file.url
        if source_json == None: source_json = {}
        if obj.document:
            doc_to_file = DocumentToFile()
            res = doc_to_file.document_data_to_file(request,obj.document.id)
            
            try:
                with open(res.text,"r") as fp:
                    json_data = json.load(fp)
            except:
                json_data = {}
            trans_json = json_data	
            if obj.job.project.ai_user.user_enterprise.subscription_name == 'Enterprise - TFN':
                target_json = merge_dict(trans_json,source_json)
            else: 
                target_json = trans_json
                target_file_path = get_file_url(res.text)
        else:
           target_json = obj.news_task.first().target_json 
           if target_json == None: target_json = {}

        
    return Response({'source_json':source_json,'target_json':target_json,\
                      'source_file_path':source_file_path,'target_file_path':target_file_path})


@api_view(['GET'])
@permission_classes([IsAuthenticated,IsEnterpriseUser])
def federal_segment_translate(request):
    '''
    This is to get the raw_mt of text. With task_id, we will get source lang and target_lang and checks for
    user_credit and call get_translation to translate the text and returns the result if the credits exists or 
    It will return insufficient credits.
    It is now used for federal flow.
    '''
    task_id = request.query_params.get('task_id',None)
    text =  request.query_params.get('text',None)
    task_instance =  Task.objects.get(id=task_id)
    if text:
        mt_engine = task_instance.job.project.mt_engine
        src_lang = task_instance.job.source_language.locale.first().locale_code
        tar_lang = task_instance.job.target_language.locale.first().locale_code

        consumable_credit = get_consumable_credits_for_text(text,target_lang=None,source_lang=src_lang)
        initial_credit = request.user.credit_balance.get("total_left")
        if initial_credit > consumable_credit:
            trans_text = get_translation(mt_engine.id , source_string=text,source_lang_code=src_lang,
                                        target_lang_code=tar_lang,user_id=request.user.id)
            return Response({'text':trans_text},status=200)
        else:
            return Response({'msg':'Insufficient Credits'},status=400)
    else:
        return Response({'msg':'Text field empty'},status=400)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def get_ner(request):
    '''
    This function is to get the text and find the NER with nlp and returns the list.
    '''
    text = request.POST.get('text')
    doc = nlp(text)
    exclude_labels = ['DATE', 'TIME', 'PERCENT', 'MONEY', 'QUANTITY', 'ORDINAL', 'CARDINAL', 'LANGUAGE']
    ner = [ent.text for ent in doc.ents if ent.label_ not in exclude_labels]
    ner_new = list(set(ner))
    return JsonResponse({"ner": ner_new}, safe=False)



############# function for segment choice #######################
import re

def contains_valid_words(sentence):
    # Define a regular expression pattern to match valid words
    word_pattern = re.compile(r'\b\w+\b')
    
    # Find all words in the sentence
    words = word_pattern.findall(sentence)
    
    # Check if there are any valid words
    if not words:
        return False
    
    # Check if there is only one word and it is alphabetic
    if len(words) == 1 and words[0].isalpha():
        return False
    
    # Check if there are multiple valid words and at least one is alphabetic
    return any(word.isalpha() for word in words)



@api_view(['GET'])
#@permission_classes([IsAuthenticated])
def segment_choice_mt_and_glossary(request):

    from ai_workspace_okapi.api_views import check_source_words
    from ai_staff.models import SegmentChoices
    from rest_framework.response import Response
    from ai_openai.utils import get_consumable_credits_for_openai_text_generator,get_prompt_chatgpt_turbo
    from ai_workspace.models import Segment
    from ai_workspace_okapi.api_views import get_src_tags

    user = request.user
    segment_id = request.GET.get('segment_id',None)
    seg_choice_id = request.GET.get('seg_choice_id',None)
    seg_choice_ins = SegmentChoices.objects.get(id=seg_choice_id)
    segment_instance = Segment.objects.get(id=segment_id)
    initial_credit = user.credit_balance.get("total_left")
    if initial_credit == 0:
        return  Response({'msg':'Insufficient Credits'},status=400)
    src_seg = segment_instance.source
    tar_seg = segment_instance.seg_mt_raw.mt_raw ### taking the segment from the mt_raw table (1 to 1 relation) which is a base target
    seg_task = segment_instance.task_obj
    src_lang = seg_task.job.source_language.language
    tar_lang = seg_task.job.target_language.language
    tags = get_src_tags(src_seg) 

    ### for cleaning sentence and extract tags

    if contains_valid_words(src_seg):

        
        tar_seg = re.sub('<[^<]+?>', '', tar_seg)
        src_seg = re.sub('<[^<]+?>', '', src_seg)


        words,gloss = check_source_words(src_seg,seg_task) ### this function checks the gloss words and select relevent words from the source segment
        #boolean,gloss= target_source_words(tar_seg,seg_task)
        ## output
        # [{'sl_term': 'Neuralink', 'tl_term': 'நியூரோ பாலம்'},
        # {'sl_term': 'neurotechnology', 'tl_term': 'பதட்டமாக தொழில்நுட்பம்'}]


        prompt = ''
        if seg_choice_ins.choice_name == "mt_llm": ## rewrite
            
            prompt = seg_choice_ins.prompt.format(tar_lang,tar_seg)
            print("rewrite----->")
            print("prompt---->",prompt)
            print("tar_lang--->",tar_lang)
            print("tar_seg---->",tar_seg)
            print("opt---->",seg_choice_ins.option)

        elif seg_choice_ins.choice_name in ["mt_glossary","mt_llm_glossary"]: 
            print("gloss")
            print("src_seg------------->",src_seg)
            print("tar_seg------------->",tar_seg)
            print("gloss----->",gloss)
            print("prompt")
            
            prompt = seg_choice_ins.prompt.format(src_seg,tar_seg,gloss)
            print(prompt)

        if prompt:
            ### check the credit 
            consumable_credits_user_text =  get_consumable_credits_for_text(prompt,source_lang='en',target_lang=None)
            if initial_credit >= consumable_credits_user_text:
                
                result_prompt = get_prompt_chatgpt_turbo(prompt,n=1)
                print("result_prompt--->",result_prompt)
                para_sentence = result_prompt["choices"][0]["message"]["content"]
                prompt_usage = result_prompt['usage']
                total_token = prompt_usage['total_tokens']
                consumed_credits = get_consumable_credits_for_openai_text_generator(total_token)
                debit_status, status_code = UpdateTaskCreditStatus.update_credits(user, consumed_credits)

                return Response({'result':para_sentence,'tag':tags},status=200)
            else:
                return  Response({'msg':'Insufficient Credits'},status=400)
        else:
            print("no prompt") ## return the mt_raw because no prompt
            return Response({'result':tar_seg,'tag':tags},status=200) 
    else:
        return Response({'result':tar_seg,'tag':tags},status=200)
    

#### to get the number of insert and delete in a list of segments using task id

@api_view(['GET',])
def get_task_segment_diff(request):
    task = request.GET.get('task',None)
    if not task:
        return Response({'msg':'Need Task'})
    task_ins = Task.objects.get(id=task)
    from ai_workspace.utils import number_of_words_delete,number_of_words_insert
    from tqdm import tqdm
    no_of_insert = []
    no_of_delete = []
    for segment in tqdm(task_ins.document.get_segments()):
        for segment_history in segment.segment_history.all():
            for segment_diff in segment_history.segment_difference.all():
                if segment_diff:
                    insertion_result = number_of_words_insert(segment_diff.sentense_diff_result)
                    deletion_result =  number_of_words_delete(segment_diff.sentense_diff_result)
                    if insertion_result[-1]:
                        no_of_insert.append(insertion_result)
                    if deletion_result[-1]:
                        no_of_delete.append(deletion_result)

    return Response({'insertion_done':len(no_of_insert),
                     'deletion_done':len(no_of_delete)})
