

import os,sys,random
from django.core.files.uploadedfile import InMemoryUploadedFile
from io import StringIO, BytesIO
from indicnlp.tokenize.sentence_tokenize import sentence_split
import nltk, re ,logging ,json,os,string
from django.core.cache import cache
 
 
 
logger = logging.getLogger('django')


 

async def detect_lang(text):
    from googletrans import Translator
    detector = Translator()
    return await detector.detect(text)

class DjRestUtils:

    def get_a_inmemoryuploaded_file():
        io = BytesIO()
        with open("/home/langscape/Documents/translate_status_api.txt", "rb") as f:
            io.write(f.read())
        io.seek(0)
        im = InMemoryUploadedFile(io, None, "text.txt", "text/plain",
                sys.getsizeof(io), None)
        return im

    def convert_content_to_inmemoryfile(filecontent, file_name):
        # text/plain hardcoded may be needs to be change as generic...
        io = BytesIO()
        io.write(filecontent)
        io.seek(0)
        im = InMemoryUploadedFile(io, None, file_name,
                                  "text/plain", sys.getsizeof(io), None)
        return im

def create_dirs_if_not_exists(path):
	if not os.path.isdir(path):
		os.makedirs(path)
		return  path
	return create_dirs_if_not_exists(path+random.choice(["-", "_","@", "!"])+str(random.randint(1,100)))

# def print_key_value(keys, values):
# 	for i, j in zip(keys, values):
# 		print(f'{i}--->{j}')


def create_assignment_id():
	from ai_workspace.models import TaskAssignInfo
	chars=string.ascii_uppercase + string.digits
	size=6
	rand_id = "AS-"+''.join(random.choice(chars) for _ in range(size))
	pr = TaskAssignInfo.objects.filter(assignment_id = rand_id)
	if not pr:
		return  rand_id
	return create_assignment_id()



def create_task_id():
	from ai_workspace.models import Task
	chars=string.ascii_uppercase + string.digits
	size=6
	rand_id = "TK-"+''.join(random.choice(chars) for _ in range(size))
	pr = Task.objects.filter(ai_taskid = rand_id)
	if not pr:
		return  rand_id
	return create_task_id()


def create_ai_project_id_if_not_exists(user):
	from ai_workspace.models import Project
	rand_id = user.uid+"p"+str(random.randint(1,10000))
	pr = Project.objects.filter(ai_project_id = rand_id)
	if not pr:
		return  rand_id
	return create_ai_project_id_if_not_exists(user)
# //////////////// References \\\\\\\\\\\\\\\\\\\\

# random.choice([1,2,3])  ---> 2
# random.choice([1,2,3])  ---> 2
# random.choices([1,2,3])  ---> [2]
import math

def roundup(x):
    return int(math.ceil(x / 15.0)) * 15



def get_consumable_credits_for_text_to_speech(total_chars):
    return round(total_chars/20)

def get_consumable_credits_for_speech_to_text(total_seconds):#######Minimum billable 15 seconds##########
    return round(roundup(total_seconds)/3)

def task_assing_role_ls(task_assign_info_ls):
	from ai_auth.signals import assign_object
	from ai_auth.utils import get_assignment_role
	from ai_workspace.models import TaskAssignInfo
	from ai_workspace.models import AiRoleandStep
	objs = TaskAssignInfo.objects.filter(id__in=task_assign_info_ls)
	for instance in objs:
		role = get_assignment_role(instance,instance.task_assign.step,instance.task_assign.reassigned)
		# role= AiRoleandStep.objects.get(step=instance.task_assign.step).role.name
		assign_object.send(
			sender=TaskAssignInfo,
			instance = instance,
			user=instance.task_assign.assign_to,
			role = role
		)


import copy,json
from django.conf import settings
TRANSLATABLE_KEYS_FEDERAL = settings.TRANSLATABLE_KEYS_FEDERAL.split(" ")
HTML_MIME_FEDARAL = settings.HTML_MIME_FEDARAL.split(" ")


MIME_TYPE_FEDARAL = {'html': 'html', 'text': 'text'}
LIST_KEYS_FEDARAL={'media':['caption']}# , 'news_tags':['name']}


####Need to add credit check and function to get MT for each key####################
def federal_json_translate(json_file,tar_code,src_code,user,translate=True):
	from ai_workspace_okapi.utils import get_translation
	
	try:json_data = json_file['news'][0]
	except: json_data = json_file
	json_file_copy = copy.deepcopy(json_data)
	for key,value in json_file_copy.items():
		if key in TRANSLATABLE_KEYS_FEDERAL:
			format_ = MIME_TYPE_FEDARAL['html'] if key in HTML_MIME_FEDARAL else MIME_TYPE_FEDARAL['text']
			if type(value) == list:
				if key in  LIST_KEYS_FEDARAL.keys(): #news_tags media
					for lists in LIST_KEYS_FEDARAL[key]:
						for list_names in json_file_copy[key]:
							if lists in list_names.keys():
								if translate:
									list_names[lists] = get_translation(mt_engine_id=1,source_string=list_names[lists],target_lang_code=tar_code,
																		source_lang_code=src_code,format_=format_,user_id=user.id)			
			else:
				if translate:
					json_file_copy[key] = get_translation(mt_engine_id=1,source_string=json_file_copy[key],target_lang_code=tar_code,
														source_lang_code=src_code,format_=format_,user_id=user.id)
	return  json_file_copy


# For voice projects and tasks
def split_file_by_size(input_file, output_directory, lang_code, max_size):
    from .api_views import cust_split
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()
    lang_code = lang_code.split('-')[0]
    if lang_code in ['zh','ja']:
        sentences = cust_split(content)
    elif lang_code in ['hi', 'bn', 'or', 'ne', 'pa']:
        sentences = sentence_split(content, lang_code, delim_pat='auto')
    else:
        sentences = nltk.sent_tokenize(content)

    part_number = 1
    current_size = 0
    current_content = []

    for sentence in sentences:
        sentence_size = len(sentence.encode('utf-8'))

        if current_size + sentence_size > max_size and current_content:
            output_file = f"{output_directory}/output_{part_number}.txt"
            with open(output_file, 'w', encoding='utf-8') as output:
                output.write("\n".join(current_content))
            part_number += 1
            current_content = []
            current_size = 0

        current_content.append(sentence)
        current_size += sentence_size

    if current_content:
        output_file = f"{output_directory}/output_{part_number}.txt"
        with open(output_file, 'w', encoding='utf-8') as output:
            output.write("\n".join(current_content))







################################## For Federal Flow #####################################
def split_dict(single_data):
    trans_keys = ["keywords","description","image_caption","heading","newsId","authorName","location","story"]
    trans_key_get_list = {"media":"caption"}#, "news_tags":"name"}
    trans_keys_dict = {}
    json_data = single_data.get('news')[0] if single_data.get('news') else None
    if not json_data:
        json_data = single_data
    for key,value in  json_data.items():
        if key in trans_keys:
            trans_keys_dict[key] = value
        if key in list(trans_key_get_list.keys()):
            trans_list=[]
            for i in value:
                if trans_key_get_list[key] in i.keys():
                    trans_list.append({trans_key_get_list[key] :i[trans_key_get_list[key]]})
            trans_keys_dict[key] = trans_list
    return trans_keys_dict

def split_dict_pib(single_data):
    print(single_data, "single data", type(single_data))
    trans_keys = ["heading","story"]
    trans_keys_dict = {}
    json_data = single_data
    for key,value in  json_data.items():
        if key in trans_keys:
            trans_keys_dict[key] = value
    
    print(trans_keys_dict, "Trans key dict", type(trans_keys_dict))
    return trans_keys_dict

def merge_dict(translated_json,raw_json):
	raw_json_trans = copy.deepcopy(raw_json)
	translated_json_copy = copy.deepcopy(translated_json)
	for key,values in list(LIST_KEYS_FEDARAL.items()):
		if key in list(raw_json_trans.keys()):
			for count,j in enumerate(raw_json_trans[key]):
				if values[count] in j.keys():
					j.update(translated_json_copy[key][count])
		if key in translated_json_copy:
			translated_json_copy.pop(key)
	raw_json_trans.update(translated_json_copy)
	return raw_json_trans

########################################################################################## 


import json
import pypandoc
from docx import Document

 
def html_to_docx(html_content, docx_filename):
	if html_content == None:
		html_content = "<p>"

	html_content = html_content.replace('\n', '<br>')
    # Convert HTML to DOCX using pypandoc
	pypandoc.convert_text(html_content, 'docx', format='html',outputfile=docx_filename)
   
    

def add_additional_content_to_docx(docx_filename, additional_content):
    doc = Document(docx_filename)
    for key, value in additional_content.items():
        if key!='story':
            doc.add_paragraph(f'{key.capitalize()}:')
            doc.add_paragraph(f'{value}')
    doc.save(docx_filename)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import tempfile
import os

# def generate_pib_docx(heading: str, story: str, base_filename: str):
#     """
#     Generates PIB DOCX with formatting and returns the file path.
#     """

#     doc = Document()

#     # Heading centered (H2)
#     if heading:
#         h = doc.add_heading(heading, level=2)
#         h.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         spacer = doc.add_paragraph()
#         spacer.add_run("") 

#     # Split story into paragraphs
#     if story:
#         for block in story.split("\n\n"):
#             para = doc.add_paragraph()
#             run = para.add_run(block.strip())
#             para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#             run.font.name = "Times New Roman"
#             run.font.size = Pt(12)

#     # Save to temp location
#     filename = f"{base_filename}.docx"
#     tmp_path = os.path.join(tempfile.gettempdir(), filename)
#     doc.save(tmp_path)

#     return tmp_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from htmldocx import HtmlToDocx
import tempfile, os

def generate_pib_docx(heading: str, story: str, base_filename: str):
    doc = Document()

    # Add heading
    if heading:
        h = doc.add_heading(heading, level=2)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
    print(heading, type(heading))
    # Add story text (paragraph-wise)
    if story:
        print(story, type(story))
        paragraphs = [p.strip() for p in story.split("\n\n") if p.strip()]

        for block in paragraphs:
            p = doc.add_paragraph(block)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph("")   # spacing

    # Save file
    filename = f"{base_filename}.docx"
    tmp_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(tmp_path)

    return tmp_path

############################# Project Filter #################################################

from django.db.models import Q, Prefetch, Count, F
import time
def progress_filter(queryset,value,users):
	from ai_workspace.models import TaskAssign

	queryset = queryset.filter(project_type_id=8) #News project

	queryset = queryset.prefetch_related(
        Prefetch('project_jobs_set__job_tasks_set__task_info', queryset=TaskAssign.objects.all())
    )
	
	if value == 'inprogress':
		if users:
			pr_ids = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__status__in = [1,2,4])\
			|Q(project_jobs_set__job_tasks_set__task_info__client_response = 2),\
			project_jobs_set__job_tasks_set__task_info__task_assign_info__isnull=False,\
			project_jobs_set__job_tasks_set__task_info__assign_to__in = users).distinct().values_list('id',flat=True)
		else:
			pr_ids = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__status__in = [1,2,4])|\
			Q(project_jobs_set__job_tasks_set__task_info__client_response = 2)).distinct().values_list('id', flat=True)
	elif value == 'submitted':
		if users:
			qs = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__status = 3),\
			project_jobs_set__job_tasks_set__task_info__task_assign_info__isnull=False,\
			project_jobs_set__job_tasks_set__task_info__assign_to__in = users).distinct()
			# Need to change this with annotate
			filtered_qs = [i.id for i in qs if i.get_tasks.filter(task_info__status=3,task_info__assign_to__in=users).count() == i.get_tasks.filter(task_info__client_response=1,task_info__assign_to__in=users).count()]
		else:
			qs = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__status = 3)).distinct()
			# Need to change this with annotate
			filtered_qs = [i.id for i in qs if i.get_tasks.filter(task_info__status=3).count() == i.get_tasks.filter(task_info__client_response=1).count()]
		pr_ids = qs.exclude(id__in=filtered_qs).values_list('id',flat=True)
		
	elif value == 'approved':
		if users:
			pr_ids = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__client_response = 1),\
			project_jobs_set__job_tasks_set__task_info__task_assign_info__isnull=False,\
			project_jobs_set__job_tasks_set__task_info__assign_to__in = users).distinct().values_list('id',flat=True)
		else:
			pr_ids = queryset.filter(Q(project_jobs_set__job_tasks_set__task_info__client_response = 1)).distinct().values_list('id',flat=True)

	queryset = queryset.filter(id__in = pr_ids)

	return queryset

def number_of_words_insert(segment):
	words_inserts = re.findall(r'<ins class="changed-word">(.+?)</ins>', segment)
	words_insert_with_classes =  re.findall(r'<ins>(.+?)</ins>', segment)  
	len_words_inserts = 0
	for words_insert in words_inserts:
		len_words_inserts= len_words_inserts+len(words_insert.split(" "))
	for words_insert_with_class in words_insert_with_classes:
		len_words_inserts=len_words_inserts+ len(words_insert_with_class.split(" "))
	return len_words_inserts

def number_of_words_delete(segment):
    words_deletes = re.findall(r'<del>(.+?)</del>', segment)  
    words_delete_class =  re.findall(r'<del class="removed-word">(.+?)</del>', segment) 
    len_words_deletes = 0
    for i in words_deletes:
        len_words_deletes=len_words_deletes+ len(i.split(" "))
    for j in words_delete_class:
        len_words_deletes=len_words_deletes+ len(j.split(" "))		
    return len_words_deletes

 


def word_count_find(task):
    import requests
    from .serializers import TaskSerializer
    from ai_workspace_okapi.api_views import DocumentViewByTask  
    from ai_workspace_okapi.utils import get_res_path
    from os.path import exists


    spring_host = os.environ.get("SPRING_HOST")
    data = TaskSerializer(task).data
    DocumentViewByTask.correct_fields(data)
    params_data = {**data, "output_type": None}

    res_paths = get_res_path(params_data["source_language"])

    doc = requests.post(url=f"http://{spring_host}:8080/getDocument/", data={
        "doc_req_params": json.dumps(params_data),
        "doc_req_res_params": json.dumps(res_paths)
    })
    if doc.status_code == 200:
        doc_data = doc.json()
        return doc_data.get('total_word_count')


def merge_source_text_by_text_unit(document_id):
    from ai_workspace_okapi.models import MergedTextUnit
    from ai_workspace_okapi.models import TextUnit, Segment
    text_units = TextUnit.objects.filter(document_id=document_id)

    for text_unit in text_units:
        source_paragraph = ""
        
        segments = text_unit.text_unit_segment_set.all()
        
        if not segments:
            continue  

        for seg in segments:
            if seg.source:
                source_paragraph += seg.source + " "

        MergedTextUnit.objects.create(
            text_unit=text_unit,
            source_para=source_paragraph.strip(),
        )


def re_initiate_failed_batch(task, project):
    from ai_workspace.models import TrackSegmentsBatchStatus
    from ai_workspace_okapi.models import MergedTextUnit
    from ai_workspace.enums import BatchStatus
    from ai_workspace.models import Task
     
    from ai_auth.tasks import adaptive_segment_translation

    try:
        task_id = task.id
        #get_terms_for_task = get_glossary_for_task(project, task)
        failed_task_batches = TrackSegmentsBatchStatus.objects.filter(document=task.document, status=BatchStatus.FAILED)
        task = Task.objects.select_related('job__source_language', 'job__target_language').get(id=task.id)
        source_lang = task.job.source_language.language
        target_lang = task.job.target_language.language

        retry_exceeded_batches = TrackSegmentsBatchStatus.objects.filter(
            document=task.document,
            status=BatchStatus.FAILED,
            retry_count__gte=1
        )

        if retry_exceeded_batches.exists():
            print(f"Skipping re-initiation for batches that have exceeded retry limit: {retry_exceeded_batches.values_list('id', flat=True)}")
            return False
        
        for failed_task_batch in failed_task_batches:
            # check retry count
            # if failed_task_batch.retry_count >= 1:
            #     print(f"Batch {failed_task_batch.id} has reached maximum retry limit. Skipping re-initiation.")
            # else:
            merged_text_units = MergedTextUnit.objects.filter(text_unit__id__range=(failed_task_batch.seg_start_id, failed_task_batch.seg_end_id))
            para = []
            metadata = {}
            for text_unit in merged_text_units:
                para.append(text_unit.source_para)
                metadata[text_unit.text_unit.id] = text_unit.source_para
            
            adaptive_segment_translation.apply_async(
                args=(para, metadata, source_lang, target_lang, task_id, True,),
                kwargs={
                    'failed_batch': True,
                    'celery_task_id': failed_task_batch.celery_task_id,
                    'batch_no': failed_task_batch.celery_task_batch
                },
                queue='high-priority'
            )
            failed_task_batch.status = BatchStatus.ONGOING
            failed_task_batch.progress_percent = 0
            failed_task_batch.retry_count += 1
            failed_task_batch.save()
            cache_key = f"adaptive_progress_{failed_task_batch.id}"
            cache.delete(cache_key)
            return True
    except Exception as e:
        print("Error in re_initiate_failed_batch:", e)
        return False



import os
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from zipfile import BadZipFile


def write_stage_response_in_excel(
    project_id,
    task_id,
    batch_no,
    system_prompt,
    user_message,
    translated_result,
    stage,
    base_dir="Translation_Results",
    input_token=0,
    output_token=0
):
    os.makedirs(base_dir, exist_ok=True)

    project_task_folder = os.path.join(base_dir, f"{project_id}_{task_id}")
    os.makedirs(project_task_folder, exist_ok=True)

    file_path = os.path.join(project_task_folder, f"batch_{batch_no}.xlsx")

    try:
        if os.path.exists(file_path):
            wb = load_workbook(file_path)
        else:
            wb = Workbook()
    except BadZipFile:
        print(f"Warning: {file_path} is not a valid Excel file. Recreating.")
        wb = Workbook()

    if "Sheet" in wb.sheetnames and wb["Sheet"].max_row == 1:
        wb.remove(wb["Sheet"])

    sheet_name = f"batch_{batch_no}"
    if sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
    else:
        ws = wb.create_sheet(title=sheet_name)
        ws.append(["Result", "User_Message", "System_Message", "Stage", "Input_Token", "Output_Token"])

    ws.append([translated_result, user_message, system_prompt, stage, input_token, output_token])

    for column_cells in ws.columns:
        max_length = max(len(str(cell.value)) for cell in column_cells if cell.value)
        col_letter = get_column_letter(column_cells[0].column)
        ws.column_dimensions[col_letter].width = max_length + 2

    wb.save(file_path)
    print(f"Data written to: {file_path}")


