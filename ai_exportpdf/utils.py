import base64
import docx ,json,logging,mimetypes,os ,pdftotext
import re,requests,time,urllib.request
from io import BytesIO
from PyPDF2 import PdfFileReader
from celery import shared_task
from django.contrib.auth import settings
from django.http import HttpResponse
from rest_framework.response import Response
from google.cloud import vision_v1, vision
from google.oauth2 import service_account
from pdf2image import convert_from_path
from tqdm import tqdm
from ai_auth.models import UserCredits
from ai_exportpdf.models import Ai_PdfUpload
from ai_tms.settings import GOOGLE_APPLICATION_CREDENTIALS_OCR, CONVERTIO_API ,OPENAI_API_KEY ,OPENAI_MODEL
from ai_exportpdf.convertio_ocr_lang import lang_code ,lang_codes
from ai_staff.models import Languages
from django.db.models import Q
import math 
import urllib
logger = logging.getLogger('django')
# credentials = service_account.Credentials.from_service_account_file(GOOGLE_APPLICATION_CREDENTIALS_OCR)
client = vision.ImageAnnotatorClient()
google_ocr_indian_language = ['bengali','hindi','kannada','malayalam','marathi','punjabi','tamil','telugu']

def download_file(file_path):
    filename = os.path.basename(file_path)
    fl = open(file_path, 'rb')
    mime_type, _ = mimetypes.guess_type(file_path)
    response = HttpResponse(fl, content_type=mime_type)
    response['Content-Disposition'] = "attachment; filename=%s" % filename
    return response

def direct_download_urlib_docx(url,filename): 
    path , basename = os.path.split(url)
    url = path+"/"+urllib.parse.quote(basename)
    x = urllib.request.urlretrieve(url=url , filename=filename)

def remove_carraige_return(txt):
    with open(f'1.txt', 'w' , encoding="utf-8",newline= '\r\n') as the_file:
            the_file.write(txt)
    final_data = [line.replace('\r\n','\n') if '\r\n' in  line[:2] else line.replace('\r\n','') for line in [x.decode() for x in  open(f'1.txt','rb').readlines()]]
    final_data = "".join(final_data)
    y = final_data.split(".")
    y = "\n".join(y)
    return y

def image_ocr_google_cloud_vision(image , inpaint):
    if inpaint:
        image = vision_v1.types.Image(content=image)
        response = client.text_detection(image = image)
        texts = response.full_text_annotation
        return texts
    else:
        buffer = BytesIO()
        image.save(buffer, format="PNG")
        image = vision_v1.types.Image(content=buffer.getvalue() )
        response = client.text_detection(image = image)
        texts = response.full_text_annotation
        if texts:
            texts = para_creation_from_ocr(texts)
            return texts
        else:
            return ""

@shared_task(serializer='json')
def convertiopdf2docx(id ,language,ocr = None ):
    txt_field_obj = Ai_PdfUpload.objects.get(id = id)
    fp  = txt_field_obj.pdf_file.path
    pdf = PdfFileReader(open(fp,'rb') ,strict=False)
    pdf_len = pdf.getNumPages()
    pdf_file_name = fp.split("/")[-1].split(".pdf")[0]+'.docx'    ## file_name for pdf to sent to convertio
    user_credit = UserCredits.objects.get(Q(user=txt_field_obj.user) & Q(credit_pack_type__icontains="Subscription") & Q(ended_at=None))
    # user_credit =UserCredits.objects.filter(Q(user_id=80) & Q(ended_at=None)).filter(Q(credit_pack_type__icontains="Addon") or Q(credit_pack_type__icontains="Subscription"))[0]
    
    with open(fp, "rb") as pdf_path:
        encoded_string = base64.b64encode(pdf_path.read())
    data = {'apikey': CONVERTIO_API ,'input': 'base64', 'file': encoded_string.decode('utf-8'),'filename':   pdf_file_name,'outputformat': 'docx' }
    # if ocr == "ocr":
        # language = language.split(",")                    #if ocr is True and selecting multiple language
        # language_convertio = [lang_code(i) for i in language]
        # ocr_option =  { "options": { "ocr_enabled": True, "ocr_settings": { "langs": [language_convertio]}}}
        # data = {**data , **ocr_option}     #merge dict
        # txt_field_obj.pdf_api_use = "convertio_ocr"
    #     print("[convertio ocr]")
    # else:
    #     txt_field_obj.pdf_api_use = "convertio"
    #     print("[convertio text]")
    txt_field_obj.pdf_api_use = "convertio"
    txt_field_obj.pdf_no_of_page = int(pdf_len)
    response_status = requests.post(url='https://api.convertio.co/convert' , data=json.dumps(data)).json()
    if response_status['status'] == 'error':
        txt_field_obj.status = "ERROR"
        txt_field_obj.pdf_api_use = "FileCorrupted"
        txt_field_obj.save()
        ###retain cred if error
        file_format,page_length =  file_pdf_check(fp,id)
        # file_format,page_length = pdf_text_check(fp)
        consum_cred = get_consumable_credits_for_pdf_to_docx(page_length ,file_format)
        user_credit.credits_left = user_credit.credits_left + consum_cred
        user_credit.save()
        print({"result":"Error during input file fetching: couldn't connect to host"})
    else:
        get_url = 'https://api.convertio.co/convert/{}/status'.format(str(response_status['data']['id']))
        try:
            while (requests.get(url = get_url).json()['data']['step'] != 'finish'): # \
                # and  (requests.get(url = get_url).json()['status'] != 'ok') \
                #     and (requests.get(url = get_url).json()['code'] != 200):
                txt_field_obj.status = "PENDING"
                txt_field_obj.save()
                time.sleep(2)

            convertio_response_link =  requests.get(url = get_url).json()
            file_link = convertio_response_link['data']['output']['url']  ##after finished get converted file from convertio
            direct_download_urlib_docx(url= file_link , filename= str(settings.MEDIA_ROOT+"/"+ str(txt_field_obj.pdf_file)).split(".pdf")[0] +".docx" )
            txt_field_obj.status = "DONE"
            txt_field_obj.pdf_conversion_sec = int(convertio_response_link['data']['minutes'])*60
            txt_field_obj.docx_url_field = str(settings.MEDIA_URL+str(txt_field_obj.pdf_file)).split(".pdf")[0] +".docx" ##save path to database
            txt_field_obj.docx_file_name = str(txt_field_obj.pdf_file_name).split('.pdf')[0]+ '.docx'
            txt_field_obj.save()
            print({"result":"finished_task" })
        except:
            if "error" in requests.get(url = get_url).json():
                print("OCR Calling")
                response_result = ai_export_pdf.apply_async((id, ),)
            # end = time.time()
            else:
                txt_field_obj.status = "ERROR"
                txt_field_obj.pdf_api_use = "FileCorrupted"
                txt_field_obj.save()
                file_format,page_length =  file_pdf_check(fp,id)
                # file_format,page_length = pdf_text_check(fp)
                consum_cred = get_consumable_credits_for_pdf_to_docx(page_length ,file_format)
                user_credit.credits_left = user_credit.credits_left + consum_cred
                user_credit.save()
                print("pdf_conversion_something went wrong")


import tempfile
#########ocr ######
from celery_progress.backend import ProgressRecorder
@shared_task(serializer='json',bind=True)
def ai_export_pdf(self, id): # , file_language , file_name , file_path
    txt_field_obj = Ai_PdfUpload.objects.get(id = id)
    # user_credit =UserCredits.objects.get(Q(user=txt_field_obj.user) & Q(credit_pack_type__icontains="Subscription") & Q(ended_at=None))
    fp = txt_field_obj.pdf_file.path
    start = time.time()
    pdf = PdfFileReader(open(fp,'rb') ,strict=False)
    pdf_len = pdf.getNumPages()
    try:
        no_of_page_processed_counting = 0
        txt_field_obj.pdf_no_of_page=int(pdf_len)
        doc=docx.Document()
        progress_recorder=ProgressRecorder(self)
        for i in tqdm(range(1,pdf_len+1)):
            with tempfile.TemporaryDirectory() as image:
                image = convert_from_path(fp ,thread_count=8,fmt='png',grayscale=False ,first_page=i,last_page=i ,size=(800, 800) )[0]
                # ocr_pages[i] = pytesseract.image_to_string(image ,lang=language_pair)  tessearct function
                text = image_ocr_google_cloud_vision(image , inpaint=False)
                text = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', text)
                progress_recorder.set_progress(i+1,pdf_len,description='pdf_converting')
                if i == 1:
                    all_text = doc.add_paragraph(text)
                else:
                    all_text.add_run(text)
                # doc.add_paragraph(text)
            end = time.time()
            no_of_page_processed_counting+=1
            txt_field_obj.counter = int(no_of_page_processed_counting)
            txt_field_obj.status = "PENDING"
            txt_field_obj.save()
        progress_recorder.set_progress(pdf_len+1, pdf_len+1, "pdf_convert_completed")
 
        logger.info('finished ocr and saved as docx ,file_name:')
        txt_field_obj.status = "DONE"
        docx_file_path = str(fp).split(".pdf")[0] +".docx"
        doc.save(docx_file_path)
        print("DocxFilePath------------->",docx_file_path)
        html_str = docx_to_html(docx_file_path)
        txt_field_obj.html_data = html_str
        txt_field_obj.docx_url_field = docx_file_path
        txt_field_obj.pdf_conversion_sec = int(round(end-start,2))
        txt_field_obj.pdf_api_use="google-ocr"
        txt_field_obj.docx_file_name=str(txt_field_obj.pdf_file_name).split('.pdf')[0]+ '.docx'
        txt_field_obj.save()
    except:
        end = time.time()
        txt_field_obj.status = "ERROR"
        txt_field_obj.pdf_api_use = "FileCorrupted"
        txt_field_obj.save()
        ###retain cred if error
        # file_format,page_length=file_pdf_check(fp,id) 
        # consum_cred = get_consumable_credits_for_pdf_to_docx(page_length,file_format)
        # user_credit.credits_left=user_credit.credits_left + consum_cred
        # user_credit.save()
        print("pdf_conversion_something went wrong")
 
def google_ocr_pdf():
    pass

def para_creation_from_ocr(texts):
    para_text = []
    for i in  texts.pages:
        for j in i.blocks:
            text_list = []
            for k in j.paragraphs:
                for a in  k.words:
                    text_list.append(" ")
                    for b in a.symbols:
                        # if b.text == ".":
                        text_list.append(b.text)
            para_text.append("".join(text_list))
    para_text = "\n".join(para_text)
    para_text = para_text.replace(" .", ".")
    return para_text

import PyPDF2
from rest_framework import serializers
def file_pdf_check(file_path,pdf_id): 
    try:
        pdfdoc = PyPDF2.PdfReader(file_path)
        pdf_check = {0:'ocr',1:'text'}
        pdf_check_list = []
        for i in tqdm(range(len(pdfdoc.pages))):
            current_page = pdfdoc.pages[i]
            if current_page.extract_text():
                if len(current_page.extract_text()) >=700:
                    pdf_check_list.append(1)
                else:
                    pdf_check_list.append(0)
            else:
                pdf_check_list.append(0)
        return [pdf_check.get(max(pdf_check_list)) , len(pdfdoc.pages)]
    except:
        if pdf_id:
            file_details = Ai_PdfUpload.objects.get(id = pdf_id)
            file_details.status = "ERROR"
            file_details.save()
        #     # return None,None
        #     # file_details.status = "FileCorrupted"
        #     # file_details.save()
        #     raise serializers.ValidationError({'msg':'pdf_corrupted'}, 
        #                                     code =400)
        # else:
        return None,None
    

from ai_workspace.models import Task
def pdf_conversion(id ):
    file_details = Ai_PdfUpload.objects.get(id = id)
    #lang = Languages.objects.get(id=int(file_details.pdf_language)).language.lower()
    #pdf_text_ocr_check = file_pdf_check(file_details.pdf_file.path , id)[0]
    # pdf_text_ocr_check = pdf_text_check(file_details.pdf_file.path)[0]
    #if (pdf_text_ocr_check == 'ocr'): #or (lang in google_ocr_indian_language):
    response_result = ai_export_pdf.apply_async((id, ),)

    file_details.pdf_task_id = response_result.id
    file_details.save()
    logger.info('assigned ocr ,file_name: google indian language'+str(file_details.pdf_file_name))
    return response_result.id

    # elif pdf_text_ocr_check == 'text':
    #     response_result = convertiopdf2docx.apply_async((id,lang ,pdf_text_ocr_check),0)
    #     file_details.pdf_task_id = response_result.id
    #     file_details.save()
    #     logger.info('assigned pdf text ,file_name: convertio'+str(file_details.pdf_file_name))
    #     return response_result.id
    # else:
    #     return "error"


from django.core.files.base import ContentFile
from ai_workspace.api_views import UpdateTaskCreditStatus
def project_pdf_conversion(id):
    task_details = Task.objects.get(id = id)
    user = task_details.job.project.ai_user
    file_obj = ContentFile(task_details.file.file.read(),task_details.file.filename)
    initial_credit = user.credit_balance.get("total_left")
    file_format,page_length = file_pdf_check(task_details.file.file.path,id)

    consumable_credits = get_consumable_credits_for_pdf_to_docx(page_length,file_format)
    if initial_credit > consumable_credits:
        Ai_PdfUpload.objects.create(user= user , file_name = task_details.file.filename, status='YET TO START',
                                   pdf_file_name =task_details.file.filename  ,task = task_details ,pdf_file =file_obj , pdf_language = task_details.job.source_language_id)
        file_details = Ai_PdfUpload.objects.filter(task = task_details).last()
        debit_status, status_code = UpdateTaskCreditStatus.update_credits(user, consumable_credits)
        if file_format:
            response_result = ai_export_pdf.apply_async((file_details.id, ),)
            file_details.pdf_task_id = response_result.id
            file_details.save()
            return response_result.id ,file_details.id
        else:
            return "error"
    else:
        return Response({'msg':'Insufficient Credits'},status=400)

# from django.core.files.base import ContentFile
# from ai_workspace.api_views import UpdateTaskCreditStatus
# def project_pdf_conversion(id):
#     task_details = Task.objects.get(id = id)
#     user = task_details.job.project.ai_user
#     file_obj = ContentFile(task_details.file.file.read(),task_details.file.filename)
#     initial_credit = user.credit_balance.get("total_left")
#     file_format,page_length = file_pdf_check(task_details.file.file.path,id)

#     consumable_credits = get_consumable_credits_for_pdf_to_docx(page_length,file_format)
#     if initial_credit > consumable_credits:
#         Ai_PdfUpload.objects.create(user= user , file_name = task_details.file.filename, status='YET TO START',
#                                    pdf_file_name =task_details.file.filename  ,task = task_details ,pdf_file =file_obj , pdf_language = task_details.job.source_language_id)
#         file_details = Ai_PdfUpload.objects.filter(task = task_details).last()
#         #lang = Languages.objects.get(id=int(file_details.pdf_language)).language.lower()
#         debit_status, status_code = UpdateTaskCreditStatus.update_credits(user, consumable_credits)
#         #if (file_format == 'ocr'): #or (lang in google_ocr_indian_language):
#         if file_format:
#             response_result = ai_export_pdf.apply_async((file_details.id, ),)
#             file_details.pdf_task_id = response_result.id
#             file_details.save()
#             return response_result.id ,file_details.id
#         # elif file_format == 'text':
#         #     response_result = convertiopdf2docx.apply_async((file_details.id,lang ,file_format),0)
#         #     file_details.pdf_task_id = response_result.id
#         #     file_details.save()
#         #     return response_result.id ,file_details.id
#         else:
#             return "error"
#     else:
#         return Response({'msg':'Insufficient Credits'},status=400)

def get_consumable_credits_for_pdf_to_docx(total_pages , formats):
    if formats == 'text':
        return int(total_pages)
    else:
        return int(total_pages)*5

def ceil_round_off(token_len):
    import math
    return math.ceil(len(token_len)/4)
    
import pypandoc
def docx_to_html(docx_file_path):
    print("DocxFilePath------------->",docx_file_path)
    #extra_args = ["--metadata","title= " , "--self-contained","-s"]#,"--standalone"]#,"--css","pandoc.css"]
    output = pypandoc.convert_file(source_file=docx_file_path,
                                   to="html",format='docx')#,
                                   #extra_args=extra_args)
    
    #bootstrap_css = '<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.0.0-beta2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-BmbxuPwQa2lc/FVzBcNJ7UAyJxM6wuqIj61tLrc4wSX0szH/Ev+nYRRuWlolflfl" crossorigin="anonymous">'
    #bootstrap_js = '<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.0.0-beta2/dist/js/bootstrap.bundle.min.js" integrity="sha384-b5kHyXgcpbZJO/tY9Ul7kGkf1S0CWuKcCD38l8YkeH8z8QjE0GmW1gYU5S9FOnJ0" crossorigin="anonymous"></script>'
    return output#bootstrap_css+bootstrap_js+


import pypandoc
def docx_to_html_with_css(docx_file_path):
    print("DocxFilePath------------->",docx_file_path)
    extra_args = ["--metadata","title= " , "--self-contained","--standalone","--css","pandoc.css"]
    output = pypandoc.convert_file(source_file=docx_file_path,
                                   to="html",format='docx',extra_args=extra_args)
    
    bootstrap_css = '<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.0.0-beta2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-BmbxuPwQa2lc/FVzBcNJ7UAyJxM6wuqIj61tLrc4wSX0szH/Ev+nYRRuWlolflfl" crossorigin="anonymous">'
    bootstrap_js = '<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.0.0-beta2/dist/js/bootstrap.bundle.min.js" integrity="sha384-b5kHyXgcpbZJO/tY9Ul7kGkf1S0CWuKcCD38l8YkeH8z8QjE0GmW1gYU5S9FOnJ0" crossorigin="anonymous"></script>'
    return bootstrap_css+bootstrap_js+output

#     # with open("out1226_final.html",'w') as fp:
#     #     fp.write(output)

def remove_duplicate_new_line(text):
    return re.sub(r'\n+', '\n', text)

# def pdf_text_check(file_name ):
#     total_page_area = 0.0
#     total_text_area = 0.0
#     with open(file_name,"rb") as f:
#         doc = fitz.open(f)
#     for page_num, page in enumerate(doc):
#         total_page_area = total_page_area + abs(page.rect)
#         text_area = 0.0
#         for b in page.get_text("blocks"):
#             r = fitz.Rect(b[:4]).get_area()  # rectangle where block text appears
#             text_area = text_area + abs(r )
#         total_text_area = total_text_area + text_area
#     # doc.close()
#     tot = total_text_area / total_page_area
#     len_doc = doc.page_count
#     doc.close()
#     return ["text" if text_perc < 0.01 else "ocr" ,len_doc ]



